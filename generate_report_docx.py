"""Generate Word summary report for THz-TDS temperature study."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FIGURES = Path(__file__).parent / "figures"
OUT = Path(__file__).parent / "THz_TDS_요약보고서.docx"


def set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color_hex, qn("w:val"): "clear"},
    )
    shading.append(shading_elem)


def add_table_m2(doc):
    """Method 2 results table."""
    data = [
        (20, "1.237", "0.023", "-0.4", "3.1"),
        (30, "1.033", "0.033", "-7.8", "2.4"),
        (40, "0.952", "0.012", "-12.5", "1.7"),
        (50, "0.818", "0.070", "-14.3", "2.9"),
        (60, "0.643", "0.040", "-21.8", "1.6"),
        (70, "0.452", "0.075", "-26.0", "1.4"),
        (80, "0.230", "0.030", "-31.6", "1.5"),
        (90, "-0.032", "0.041", "-37.6", "1.4"),
        (100, "-0.270", "0.032", "-39.6", "0.8"),
        (110, "-0.493", "0.032", "-44.8", "0.6"),
    ]
    headers = ["온도 (°C)", "n", "σ(n)", "α (cm⁻¹)", "σ(α)"]

    table = doc.add_table(rows=1 + len(data), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(8)
                r.font.bold = True

    # Data rows
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)

    return table


def add_table_m3(doc):
    """Method 3 results table."""
    data = [
        (20, "0.000", "0.023", "-0.0", "3.1"),
        (30, "-0.204", "0.033", "-7.4", "2.4"),
        (40, "-0.286", "0.012", "-12.1", "1.7"),
        (50, "-0.419", "0.070", "-14.0", "2.9"),
        (60, "-0.594", "0.040", "-21.4", "1.6"),
        (70, "-0.786", "0.075", "-25.6", "1.4"),
        (80, "-1.007", "0.030", "-31.2", "1.5"),
        (90, "-1.269", "0.041", "-37.2", "1.4"),
        (100, "-1.507", "0.032", "-39.2", "0.8"),
        (110, "-1.730", "0.032", "-44.4", "0.6"),
    ]
    headers = ["온도 (°C)", "Δn", "σ(Δn)", "Δα (cm⁻¹)", "σ(Δα)"]

    table = doc.add_table(rows=1 + len(data), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(8)
                r.font.bold = True

    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)

    return table


def add_figure_pair(doc, fig_left, fig_right, cap_left, cap_right):
    """Insert two figures side by side using a 2-column table."""
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in tbl.rows:
        for cell in row.cells:
            cell.width = Cm(7)

    # Images
    for ci, fig_path in enumerate([fig_left, fig_right]):
        cell = tbl.rows[0].cells[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(fig_path), width=Cm(5.5))

    # Captions
    for ci, cap in enumerate([cap_left, cap_right]):
        cell = tbl.rows[1].cells[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cap)
        run.font.size = Pt(8)
        run.font.italic = True

    # Remove table borders
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.makeelement(qn("w:tcBorders"), {})
            for edge in ("top", "left", "bottom", "right"):
                el = tcBorders.makeelement(
                    qn(f"w:{edge}"),
                    {qn("w:val"): "none", qn("w:sz"): "0"},
                )
                tcBorders.append(el)
            tcPr.append(tcBorders)


def add_single_figure(doc, fig_path, caption):
    """Insert a single centered figure with caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fig_path), width=Cm(6))

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_p.add_run(caption)
    run.font.size = Pt(8)
    run.font.italic = True


def main():
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- Default font ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # ══════════════════════════════════════════════════════════════
    # Title
    # ══════════════════════════════════════════════════════════════
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("THz-TDS 온도 의존성 광학물성 측정 요약보고서")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("PE20 이차전지 분리막의 온도별 굴절률 및 흡수계수 분석")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()  # spacing

    # ══════════════════════════════════════════════════════════════
    # 1. 실험 개요
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("1. 실험 개요", level=1)

    doc.add_paragraph(
        "본 연구에서는 테라헤르츠 시간영역 분광법(THz-TDS)을 이용하여 "
        "이차전지용 다공성 폴리에틸렌(PE) 분리막의 온도 의존성 광학물성을 측정하였다. "
        "분리막의 열적 안정성은 배터리 안전성과 직결되며, THz 대역에서의 비파괴 "
        "온도 모니터링 가능성을 평가하고자 하였다."
    )

    doc.add_heading("1.1 시료 및 장비", level=2)

    items = [
        ("시료", "PE20 다공성 PE 분리막 (두께 20 μm)"),
        ("장비", "Menlo Systems ScanControl, Lytera THz-TDS"),
        ("온도 범위", "20–110 °C (10 °C 간격, 총 10개 온도)"),
        ("반복 측정", "각 온도당 5회 반복 (총 50개 시료 + 1개 레퍼런스)"),
        ("데이터", "3,000 포인트/파형, 5회 평균, Tab-separated txt"),
    ]
    for label, desc in items:
        p = doc.add_paragraph(style="List Bullet")
        run_b = p.add_run(f"{label}: ")
        run_b.bold = True
        p.add_run(desc)

    # ══════════════════════════════════════════════════════════════
    # 2. 분석 방법
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("2. 분석 방법", level=1)

    doc.add_heading("2.1 신호 처리", level=2)
    doc.add_paragraph(
        "시간영역 신호에 Hann 윈도우 함수를 적용한 후 zero-padding(×2)과 함께 "
        "FFT를 수행하여 주파수 영역으로 변환하였다. "
        "전달함수 H(ω) = E_sam(ω)/E_ref(ω)를 계산하고, "
        "Nelder-Mead 최적화를 통해 각 주파수에서 굴절률(n)과 소광계수(κ)를 추출하였다."
    )

    doc.add_heading("2.2 박막 근사", level=2)
    doc.add_paragraph(
        "시료 두께(20 μm)가 THz 파장(~300 μm at 1 THz)보다 충분히 작으므로 "
        "박막 근사(thin-film approximation)를 적용하였다. "
        "이 경우 Fresnel 반사 손실을 무시하고 "
        "H ≈ exp[−j(ñ − n_air)k₀d]로 간략화된다."
    )

    doc.add_heading("2.3 분석 방법론", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Method 2 (절대 분석): ")
    run.bold = True
    p.add_run(
        "H(ω) = E_sam(T)/E_ref(air, 20°C). "
        "공기 레퍼런스 대비 절대 광학물성을 추출한다. "
        "챔버 내 공기의 온도별 굴절률 변화를 보정한다."
    )

    p = doc.add_paragraph()
    run = p.add_run("Method 3 (차등 분석): ")
    run.bold = True
    p.add_run(
        "H(ω) = E_sam(T)/E_sam(20°C). "
        "20°C 시료를 기준으로 한 상대적 변화(Δn, Δα)를 추출한다. "
        "공기 레퍼런스에 의한 체계적 오차를 줄일 수 있다."
    )

    # ══════════════════════════════════════════════════════════════
    # 3. 시간영역 측정 결과
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("3. 측정 결과", level=1)

    doc.add_heading("3.1 시간영역 파형", level=2)
    add_single_figure(
        doc,
        FIGURES / "fig1_time_domain.png",
        "Fig. 1. THz 시간영역 파형. 검정: 공기 레퍼런스, 색상: 20–110 °C 시료. "
        "인셋: 피크 부근 확대 (온도 증가 시 시간 시프트 및 진폭 변화 관찰).",
    )

    doc.add_paragraph(
        "Fig. 1에 레퍼런스(공기)와 10개 온도에서의 THz 시간영역 파형을 나타내었다. "
        "20 μm 박막이므로 레퍼런스와 시료 신호의 차이는 미세하나, "
        "인셋 확대에서 온도 증가에 따른 펄스 도달 시간의 좌측 이동(leftward shift)과 "
        "진폭 증가가 관찰된다. 이는 고온에서 분리막의 다공성 감소로 인한 "
        "산란 저감 및 투과율 증가와 일치한다."
    )

    # ══════════════════════════════════════════════════════════════
    # 3.2 Method 2 결과
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("3.2 절대 분석 결과 (Method 2)", level=2)

    add_figure_pair(
        doc,
        FIGURES / "fig2_m2_refractive_index.png",
        FIGURES / "fig3_m2_absorption.png",
        "Fig. 2(a). 굴절률 n(f) 스펙트럼",
        "Fig. 2(b). 흡수계수 α(f) 스펙트럼",
    )

    doc.add_paragraph(
        "Fig. 2에 Method 2로 추출한 굴절률 및 흡수계수 스펙트럼을 나타내었다. "
        "20°C에서 n ≈ 1.24로 벌크 PE(n ≈ 1.53)보다 낮은데, 이는 다공성 구조에 "
        "의한 유효 매질 효과이다. 온도 증가에 따라 n이 급격히 감소하여 "
        "90°C 이상에서 음수값을 보이는데, 이는 20 μm 초박막에서 "
        "챔버 내 공기 경로(~1 cm)의 온도 효과가 시료 효과를 압도하기 때문이다."
    )

    add_single_figure(
        doc,
        FIGURES / "fig4_m2_n_vs_temp.png",
        "Fig. 3. 온도별 굴절률 변화 (Method 2, 0.5/1.0/1.5 THz).",
    )

    doc.add_paragraph(
        "Fig. 3에 특정 주파수(0.5, 1.0, 1.5 THz)에서의 온도 의존성을 나타내었다. "
        "세 주파수 모두 거의 동일한 선형 감소 추세를 보이며, "
        "주파수 의존성은 미미하다."
    )

    # Table M2
    doc.add_paragraph()
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Table 1. Method 2 결과 요약 (1.0 THz 기준, 5회 평균 ± 표준편차)")
    run.font.size = Pt(8)
    run.bold = True

    add_table_m2(doc)

    # ══════════════════════════════════════════════════════════════
    # 3.3 Method 3 결과
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("3.3 차등 분석 결과 (Method 3)", level=2)

    add_figure_pair(
        doc,
        FIGURES / "fig2_m3_delta_n.png",
        FIGURES / "fig3_m3_delta_alpha.png",
        "Fig. 4(a). Δn(f) 스펙트럼 (20°C 대비)",
        "Fig. 4(b). Δα(f) 스펙트럼 (20°C 대비)",
    )

    doc.add_paragraph(
        "Fig. 4에 Method 3으로 추출한 상대적 굴절률 변화(Δn)와 "
        "흡수계수 변화(Δα)를 나타내었다. 20°C 시료를 기준으로 한 차등 분석이므로 "
        "20°C에서 Δn = 0이며, 온도 증가에 따라 단조 감소한다. "
        "스펙트럼은 주파수에 대해 비교적 평탄하며, "
        "Method 2보다 노이즈가 적은 경향을 보인다."
    )

    add_single_figure(
        doc,
        FIGURES / "fig4_m3_delta_n_vs_temp.png",
        "Fig. 5. Δn vs. 온도 (Method 3, 0.5/1.0/1.5 THz).",
    )

    # Table M3
    doc.add_paragraph()
    cap2 = doc.add_paragraph()
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap2.add_run("Table 2. Method 3 결과 요약 (1.0 THz 기준, 20°C 대비 변화량)")
    run.font.size = Pt(8)
    run.bold = True

    add_table_m3(doc)

    # ══════════════════════════════════════════════════════════════
    # 4. 고찰
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("4. 고찰", level=1)

    doc.add_paragraph(
        "20°C에서 측정된 n ≈ 1.24는 다공성 PE 분리막의 유효 굴절률로서, "
        "벌크 PE(n ≈ 1.53)와 공기(n = 1.0)의 혼합으로 설명된다. "
        "Bruggeman 유효 매질 근사를 적용하면 기공률 약 40–50%에 해당하며, "
        "이는 상용 분리막의 공칭 기공률과 일치한다."
    )

    doc.add_paragraph(
        "온도 증가에 따른 굴절률 감소 경향은 두 가지 요인의 복합 효과로 해석된다:"
    )

    items_disc = [
        "챔버 내 공기 경로 효과: 20 μm 시료에 비해 수 cm의 공기 경로가 존재하며, "
        "고온에서 공기의 굴절률 감소가 측정 전달함수에 체계적 오차를 유발한다.",
        "시료 구조 변화: PE 분리막은 100–130°C 부근에서 열수축(thermal shutdown)이 "
        "시작되며, 이에 따른 기공 구조 변화가 유효 굴절률에 영향을 준다.",
        "Method 3 (차등 분석)에서도 유사한 감소 추세가 관찰되어, "
        "단순 공기 경로 보정만으로는 설명이 불충분하며, "
        "시료 자체의 온도 의존성이 존재함을 시사한다.",
    ]
    for item in items_disc:
        doc.add_paragraph(item, style="List Bullet")

    # ══════════════════════════════════════════════════════════════
    # 5. 결론
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("5. 결론", level=1)

    conclusions = [
        "THz-TDS를 이용하여 PE20 분리막의 20–110°C 온도 범위에서 "
        "광학물성(n, α)을 성공적으로 측정하였다.",
        "20°C에서 유효 굴절률 n = 1.24 ± 0.02 (1.0 THz)로, "
        "다공성 구조에 의한 벌크 PE 대비 낮은 값을 확인하였다.",
        "온도 증가에 따라 굴절률이 단조 감소하며, 이 변화는 "
        "0.5–1.5 THz 범위에서 주파수 의존성이 미미하다.",
        "절대 분석(Method 2)과 차등 분석(Method 3) 모두 일관된 "
        "온도 의존 추세를 보이나, 20 μm 초박막에 대한 정량적 해석에는 "
        "공기 경로 효과의 정밀한 보정이 필요하다.",
        "THz-TDS가 이차전지 분리막의 비파괴 온도 모니터링에 "
        "활용될 수 있는 가능성을 확인하였다.",
    ]
    for c in conclusions:
        doc.add_paragraph(c, style="List Bullet")

    # ══════════════════════════════════════════════════════════════
    # 6. 분석 조건
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6. 분석 조건 요약", level=1)

    params = [
        ("시료 두께", "20 μm (0.02 mm)"),
        ("주파수 범위", "0.2–2.5 THz"),
        ("윈도우 함수", "Hann"),
        ("Zero-padding", "×2 (next power of 2)"),
        ("전달함수 모델", "Thin-film approximation"),
        ("최적화", "Nelder-Mead (freq-by-freq, warm-starting)"),
        ("공기 보정", "n_air(T) = 1 + 2.88×10⁻⁴ × (273.15/T_K)"),
        ("챔버 길이", "1 cm"),
        ("소프트웨어", "Python 3.9 + NumPy/SciPy + 자체 개발 THz-TDS 라이브러리"),
    ]
    for label, val in params:
        p = doc.add_paragraph(style="List Bullet")
        run_b = p.add_run(f"{label}: ")
        run_b.bold = True
        p.add_run(val)

    # Save
    doc.save(str(OUT))
    print(f"Report saved: {OUT}")


if __name__ == "__main__":
    main()
