"""Sensors and Actuators A: Physical 투고용 논문 DOCX 생성.

PE40 THz-TDS temperature study — full manuscript.
Target: ~5000 words, ~10 figures, Minor Revision quality.
"""
from __future__ import annotations

import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

FIG_DIR = Path(__file__).parent.parent / "figures" / "paper_260406"
OUT_DIR = Path(__file__).parent.parent / "results" / "paper_260406"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def add_equation(doc, eq_text, eq_num):
    """Add equation in reference-paper style: centered equation + right-aligned number."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(eq_text)
    run.italic = True
    run.font.size = Pt(12)
    # Tab + equation number (right)
    run2 = p.add_run(f"\t({eq_num})")
    run2.italic = False
    run2.font.size = Pt(12)


def add_where(doc, text):
    """Add 'where' clause after equation, matching reference style."""
    p = doc.add_paragraph(f"여기서 {text}")
    p.paragraph_format.space_before = Pt(2)
    return p


def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 2.0

    for level, size in [(1, 14), (2, 13), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)


def add_figure(doc, fig_name, caption, width=5.5):
    fig_path = FIG_DIR / f"{fig_name}.png"
    if not fig_path.exists():
        doc.add_paragraph(f"[Figure missing: {fig_name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fig_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style.font.size = Pt(10)
    cap.runs[0].italic = True


def add_table_from_data(doc, headers, rows, caption):
    cap = doc.add_paragraph(caption)
    cap.runs[0].italic = True
    cap.style.font.size = Pt(10)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.runs[0].font.size = Pt(10)


def generate():
    doc = Document()
    setup_styles(doc)

    # ════════════════════════════════════════════
    # TITLE
    # ════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "테라헤르츠 시간영역 분광법 (THz-TDS) 및 머신러닝을 이용한 "
        "폴리에틸렌 이차전지 분리막의 온도 의존성 미세구조 변화 분석"
    )
    run.bold = True
    run.font.size = Pt(16)

    # Authors
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run("[Author Names]").font.size = Pt(12)
    affil = doc.add_paragraph()
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil.add_run("[Affiliations]").font.size = Pt(11)

    doc.add_paragraph()

    # ════════════════════════════════════════════
    # HIGHLIGHTS
    # ════════════════════════════════════════════
    doc.add_heading("주요 성과 (Highlights)", level=1)
    highlights = [
        "THz-TDS를 통해 PE 이차전지 분리막의 온도 의존성 기공률 변화를 규명하였습니다 (20-110 \u00b0C)",
        "굴절률이 온도에 따라 단조 감소하였습니다 (n: 1.43 \u2192 1.35, 0.45 THz 기준)",
        "Bruggeman EMA 기공률 역산: 44% (사양) \u2192 59% (100 \u00b0C), 열적 전이 온도 55 \u00b0C",
        "Lasso 회귀모델이 원시 THz 신호로부터 온도를 예측하였습니다 (R\u00b2 = 0.955, LOTO-CV, MAE = 5.1 \u00b0C)",
        "위상 기울기 및 저주파 스펙트럼 특성이 주요 온도 지표로 식별되었습니다",
    ]
    for h in highlights:
        doc.add_paragraph(h, style="List Bullet")

    # ════════════════════════════════════════════
    # ABSTRACT
    # ════════════════════════════════════════════
    doc.add_heading("초록 (Abstract)", level=1)
    doc.add_paragraph(
        "열적 스트레스 하에서 이차전지 분리막의 미세구조를 비파괴적으로 모니터링하는 것은 "
        "리튬이온 전지 안전성에 매우 중요합니다. 본 연구에서는 테라헤르츠 시간영역 분광법 "
        "(THz-TDS)을 이용하여 20\u2013110 \u00b0C 범위에서 폴리에틸렌 (PE) 분리막의 온도 의존성 "
        "광학적 및 미세구조적 특성을 조사하였습니다. "
        "매칭 레퍼런스 (matched-reference) 방식을 적용하여 10 \u00b0C 간격, 온도당 5회 반복 측정으로 "
        "굴절률 및 흡수계수 스펙트럼 (0.2\u20132.5 THz)을 추출하였습니다. "
        "0.45 THz에서의 굴절률은 1.43에서 1.35로 감소하였으며, 이는 2상 Bruggeman 유효 매질 근사 "
        "(EMA)에 의해 유효 기공률이 44%에서 59%로 증가하는 것에 해당합니다. "
        "55.2 \u00b0C의 전이 온도를 갖는 2구간 열적 거동이 확인되었으며, 이는 PE의 \u03b1c 결정 "
        "이완과 일치합니다. "
        "일원 분산분석 (one-way ANOVA)은 분석 주파수 범위의 97.4%에서 통계적으로 유의한 온도 "
        "의존성을 확인하였습니다 (\u03b7\u00b2 = 0.74, 0.45 THz 기준, p = 2.84 \u00d7 10\u207b\u2079). "
        "또한, 원시 시간영역 신호로부터 직접 온도를 예측하는 머신러닝 파이프라인을 개발하였습니다. "
        "분산 팽창 인자 (VIF) 분석을 통해 선별된 15개의 물리 기반 특성을 사용하여, Lasso 회귀모델은 "
        "온도별 교차검증 (LOTO-CV) 하에서 R\u00b2 = 0.955, MAE = 5.1 \u00b0C를 달성하였으며, "
        "위상 기울기와 저주파 스펙트럼 파워가 주요 예측인자로 식별되었습니다. "
        "이러한 결과는 본 연구의 시료 및 측정 조건 내에서, THz-TDS가 다공성 고분자 멤브레인의 "
        "열 유도 미세구조 변화를 모니터링하기 위한 비접촉 탐침 기술로서의 가능성을 입증합니다."
    )

    # Keywords
    kw = doc.add_paragraph()
    kw.add_run("키워드 (Keywords): ").bold = True
    kw.add_run(
        "테라헤르츠 시간영역 분광법 (Terahertz time-domain spectroscopy); "
        "폴리에틸렌 분리막 (Polyethylene separator); "
        "온도 의존성 (Temperature dependence); 유효 매질 근사 (Effective medium approximation); "
        "머신러닝 (Machine learning); 전지 안전성 (Battery safety)"
    )

    # ════════════════════════════════════════════
    # 1. INTRODUCTION
    # ════════════════════════════════════════════
    doc.add_heading("1. 서론 (Introduction)", level=1)
    doc.add_paragraph(
        "리튬이온 전지 (LIBs)는 휴대용 전자기기에서 전기자동차에 이르기까지 현대 에너지 저장 "
        "시스템의 핵심 구성요소입니다. 분리막은 일반적으로 미세다공성 폴리올레핀 멤브레인으로, "
        "전극 간 직접 접촉을 방지하면서 이온 수송을 허용하는 중요한 안전 부품입니다 [1,2]. "
        "과충전, 외부 단락, 또는 고온 환경과 같은 열적 남용 조건 하에서 분리막은 기공 폐쇄, "
        "수축, 궁극적으로 용융에 이르는 점진적 구조 변화를 겪으며, 이는 열폭주를 유발할 수 "
        "있습니다 [3,4]."
    )
    doc.add_paragraph(
        "따라서 열적 스트레스 하에서 분리막의 초기 단계 미세구조 변화를 이해하는 것은 전지 "
        "안전성 공학에 필수적입니다. 주사전자현미경 (SEM), 수은 압입 기공률 측정법, 시차주사 "
        "열량법 (DSC) 등 기존 특성 평가 기법은 유용한 구조 정보를 제공하지만, 파괴적이거나 "
        "시간 소모적이며, 현장 외 (ex-situ) 측정에 한정된다는 한계가 있습니다 [5,6]."
    )
    doc.add_paragraph(
        "테라헤르츠 시간영역 분광법 (THz-TDS)은 박막 고분자 필름 및 다공성 재료의 특성 "
        "평가를 위한 유망한 비파괴 기술로 부상하였습니다 [7\u20139]. "
        "테라헤르츠 주파수 대역 (0.1\u20133 THz)은 유전 특성, 구조적 주기성, 메조스케일에서의 "
        "자유 부피 변화에 민감합니다. "
        "최근 연구들은 유효 매질 근사 (EMA) 모델을 활용하여 의약품 정제 [10], 고분자 코팅 [11], "
        "세라믹 멤브레인 [12]에서의 기공률 측정에 대한 THz-TDS의 능력을 입증하였습니다."
    )
    doc.add_paragraph(
        "전지 안전성 측면에서, 고온 환경 하에서의 분리막 거동을 이해하는 것은 특히 중요합니다. "
        "열적 남용 상황에서 분리막 온도는 수 분 내에 상온에서 130 \u00b0C 이상으로 상승할 수 "
        "있습니다 [3,4]. 기공 확장, 수축 개시, 최종 용융에 이르는 미세구조 변화의 순서가 "
        "안전 개입을 위한 임계 시간 창을 결정합니다. 제조 과정에서 이러한 변화의 비접촉 "
        "특성 평가는 분리막 품질 보증을 향상시킬 것입니다. 그러나 기존의 광학 및 "
        "음향 기술은 관련 길이 척도 (두께 10\u2013100 \u00b5m, 서브마이크로미터 기공 특성)에서 "
        "고분자 멤브레인 특성 평가에 필요한 침투 깊이와 재료 감도가 부족합니다 [6]."
    )
    doc.add_paragraph(
        "최근 머신러닝 접근법이 THz-TDS와 통합되어 재료 분류 및 물성 추출의 향상에 "
        "활용되고 있습니다 [20]. 지도 회귀 모델은 물리적 특성과 상관관계가 있는 미세한 "
        "스펙트럼 특성을 식별할 수 있으며, 속도와 강건성 측면에서 기존 분석적 추출 방법을 "
        "능가할 가능성이 있습니다. 물리 기반 분석 (EMA 기공률 역산)과 데이터 기반 [29] 예측 "
        "(ML 온도 회귀)의 결합은 각 접근법이 상호 검증하고 확장하는 상보적 프레임워크를 "
        "제공합니다."
    )
    doc.add_paragraph(
        "그러나 이차전지 분리막에서 THz 광학 특성의 온도 의존성에 대한 체계적 연구는 "
        "아직 부족한 실정입니다. 특히, 기공 확장, 라멜라 재배열, 비정질상 이완 등 열적 "
        "유도 미세구조 변화와 측정 가능한 THz 매개변수 간의 관계가 정량적으로 "
        "확립되지 않았습니다."
    )
    doc.add_paragraph(
        "본 연구에서는 건식 공정 폴리에틸렌 (PE) 분리막의 20\u2013110 \u00b0C 온도 범위에서의 "
        "미세구조 변화에 대한 포괄적인 THz-TDS 연구를 제시합니다. "
        "매칭 레퍼런스 측정 프로토콜과 Bruggeman 2상 EMA를 적용하여 기공률을 역산하였습니다. "
        "관측된 경향의 통계적 유의성은 ANOVA 및 부트스트랩 신뢰구간을 통해 엄밀하게 "
        "검증하였습니다. 또한, 원시 THz 신호로부터 물리 기반 특성을 추출하여 온도를 예측하는 "
        "머신러닝 파이프라인을 개발하여, 어떤 신호 특성이 가장 많은 열적 정보를 전달하는지에 "
        "대한 통찰을 제공합니다."
    )

    # ════════════════════════════════════════════
    # 2. MATERIALS AND METHODS
    # ════════════════════════════════════════════
    doc.add_heading("2. 재료 및 방법 (Materials and Methods)", level=1)

    doc.add_heading("2.1. 시료 준비 (Sample preparation)", level=2)
    doc.add_paragraph(
        "상용 건식 공정 PE 분리막 (MS-DPS 20B, 제조사 사양 기공률 44%, "
        "시트당 두께 20 \u00b5m, 공기투과도 280 s/100 mL)을 사용하였습니다. "
        "THz 신호대잡음비를 향상시키면서 구조적 대표성을 유지하기 위해 2장을 적층하여 "
        "총 두께 약 40 \u00b5m을 확보하였습니다. "
        "건식 공정 PE 분리막은 압출된 PE 필름의 일축 연신을 통해 제조되며, 이 과정에서 "
        "슬릿 형상의 배향된 기공 네트워크가 형성됩니다. 이로 인해 0.1\u20131 \u00b5m 크기의 "
        "기공이 주로 기계 방향 (MD)으로 정렬된 특징적인 이방성 미세구조가 형성됩니다. "
        "2장 적층 구성은 유효 광학 경로 길이를 20에서 40 \u00b5m으로 증가시켜 THz-TDS 측정의 "
        "위상 감도를 약 2배 향상시키면서 구조적 대표성을 유지하기 위해 채택되었습니다. "
        "DSC 분석 (1차 가열 사이클)에서 용융 온도 Tm = 139.4 \u00b0C 및 결정화도 67.8% "
        "(기준: \u0394Hf = 198.7 J/g, \u0394H100% = 293.0 J/g)가 측정되었습니다."
    )
    doc.add_paragraph(
        "Fig. 1은 분리막 표면의 주사전자현미경 (SEM) 이미지 (50,000\u00d7 배율, "
        "E\u2013T 검출기, 2 keV)를 보여줍니다. 이 이미지는 건식 공정 PE 분리막의 "
        "특징적인 피브릴 네트워크 구조를 나타냅니다: 기계 방향 (수직)으로 배향된 "
        "두꺼운 피브릴 (~100\u2013200 nm 폭)이 횡방향의 얇은 라멜라 브릿지 "
        "(~50\u2013100 nm)에 의해 상호 연결되어 있습니다. 그 결과 형성된 슬릿 형상의 "
        "기공 (약 100\u2013500 nm)은 이온 수송 경로를 제공하는 개방형 상호연결 네트워크를 "
        "형성합니다. 이 이방성 기공 형태는 EMA 결과 해석 (2.4절)과 관련이 있으며, "
        "등방성 Bruggeman 모델은 여기서 관찰되는 배향된 슬릿 형상 기공이 아닌 "
        "구형 포함물을 가정하기 때문입니다."
    )
    add_figure(doc, "fig01_sem_dsc",
               "Fig. 1. (a) 건식 공정 PE 분리막 (MS-DPS 20B)의 SEM 현미경 사진 (50,000\u00d7 배율). "
               "슬릿 형상 기공 (~100\u2013500 nm)을 갖는 피브릴 네트워크 구조. 스케일 바: 1 \u00b5m. "
               "(b) DSC 열분석 곡선 (1차 승온). 용융 피크 Tm = 139.5 \u00b0C, 결정화도 67.8%. "
               "((a) SEM micrograph at 50,000\u00d7. (b) DSC thermogram, Tm = 139.5 \u00b0C.)",
               width=5.0)

    doc.add_heading("2.2. THz-TDS 측정 (THz-TDS measurements)", level=2)
    doc.add_paragraph(
        "THz-TDS 측정은 광섬유 결합 상용 시스템 (Menlo Systems TeraSmart, ScanControl 소프트웨어 "
        "v1.10)에 광전도 안테나 (PCA) 송/수신 모듈 (TERA15-FC)을 장착하여 수행하였습니다. "
        "실험 셋업은 Fig. 2에 나타내었습니다. "
        "측정 구성은 온도 제어 시료 스테이지를 갖춘 투과 기하학을 "
        "적용하였습니다. 20에서 110 \u00b0C까지 10 \u00b0C 간격으로 10개 온도에서 측정하였습니다 "
        "(가열 방향만 수행). "
        "6위치 터릿 시료 홀더 (Quantum Northwest)를 사용하여 온도를 제어하였으며, "
        "위치 1은 빈 레퍼런스 채널, 위치 2\u20136은 5개 시료 반복시편 (S1\u2013S5)으로 "
        "사용하였습니다. "
        "각 온도에서 열평형을 보장하기 위해 5분의 안정화 시간을 적용하였습니다 "
        "(온도 불확도 \u00b11\u20132 \u00b0C로 추정). "
        "PE는 비흡습성이므로 습도 제어 없이 상온 습도 조건 (~40\u201360% RH)에서 "
        "측정을 수행하였습니다. "
        "각 온도에서 레퍼런스 (터릿 위치 1, 빈 개구부)를 먼저 측정한 후, 5개 시료를 "
        "순차적으로 측정하였으며 (위치 2\u20136), 각 시료는 5개 파형의 평균을 취하였습니다. "
        "이 매칭 레퍼런스 프로토콜은 온도 의존적 장비 드리프트를 체계적으로 제거합니다. "
        "측정 안정성을 검증하기 위해 60\u2013110 \u00b0C에서 각 온도의 모든 시료 측정 완료 후 "
        "복귀 레퍼런스 측정을 추가로 수행하였습니다. "
        "원래 레퍼런스와 복귀 레퍼런스 간 최대 시간영역 드리프트는 피크 진폭의 2.5\u20135.0%이었으며, "
        "0.3\u20132.0 THz 범위의 스펙트럼 편차는 0.6% 미만으로 유지되어 측정 프로토콜 전반에 걸쳐 "
        "충분한 측정 안정성이 확인되었습니다."
    )
    add_figure(doc, "fig_test_setup",
               "Fig. 2. 실험 셋업. (a) THz-TDS 투과 측정 시스템 개념도: Menlo Systems TeraSmart, "
               "광섬유 결합 PCA 송/수신기 (TERA15-FC), 축외 포물면 미러, Quantum Northwest "
               "6-position 터렛 홀더 및 TC 1 온도 제어기. (b) 측정 시스템 사진. "
               "(Experimental setup. (a) Schematic of THz-TDS system. (b) Photographs.)",
               width=5.5)

    doc.add_heading("2.3. 광학 물성 추출 (Optical property extraction)", level=2)
    doc.add_paragraph(
        "복소 전달함수는 매칭 레퍼런스 방식으로 계산하였습니다:"
    )
    add_equation(doc, "H(\u03c9) = E\u209b\u2090\u2098(\u03c9, T) / E\u1d63\u2091\u2092(\u03c9, T)", 1)
    add_where(doc,
        "E\u209b\u2090\u2098 및 E\u1d63\u2091\u2092는 동일 온도 T에서의 시료 및 레퍼런스 "
        "시간영역 신호의 푸리에 변환입니다."
    )
    doc.add_paragraph(
        "전체 스펙트럼 정보를 보존하기 위해 아포다이제이션 윈도우를 적용하지 않았으며 "
        "(직사각 윈도우), 제로 패딩 계수 2를 사용하여 약 0.033 THz의 주파수 분해능을 "
        "확보하였습니다. 굴절률과 흡수계수는 H(\u03c9)의 위상 및 진폭으로부터 "
        "추출하였습니다:"
    )
    add_equation(doc, "n(\u03c9) = 1 + c \u00b7 \u0394\u03c6(\u03c9) / (2\u03c0 \u00b7 f \u00b7 d)", 2)
    add_equation(doc, "\u03b1(\u03c9) = \u2212(2/d) \u00b7 ln|H(\u03c9)|", 3)
    add_where(doc,
        "\u0394\u03c6(\u03c9)는 시료와 레퍼런스 간의 언래핑된 위상 차이, "
        "c는 광속, f는 주파수, d = 40 \u00b5m은 시료 두께입니다."
    )
    doc.add_paragraph(
        "분석은 0.2\u20132.5 THz 주파수 범위에서 수행하였습니다. "
        "40 \u00b5m 시료 두께에 대해 첫 번째 Fabry\u2013P\u00e9rot 에코는 "
        "약 \u0394t = 2nd/c \u2248 0.38 ps (n \u2248 1.4 사용)에서 발생하며, "
        "이는 약 2.6 THz의 자유 스펙트럼 범위에 해당합니다. 이 에코는 측정 윈도우 내에서 "
        "주 펄스와 시간적으로 분리되지 않으므로, 다중 반사가 추출된 광학 상수에 내재적으로 "
        "포함되어 있어 별도의 보정이 필요하지 않습니다 [14,15]."
    )

    doc.add_heading("2.4. 유효 매질 근사 (Effective medium approximation)", level=2)
    doc.add_paragraph(
        "유효 기공률은 2상 Bruggeman EMA 모델 [17]을 사용하여 결정하였으며, "
        "자기 일관적 유효 유전율 \u03b5_eff는 다음을 만족합니다:"
    )
    add_equation(doc,
        "(1 \u2212 \u03c6)(\u03b5_PE \u2212 \u03b5_eff) / (\u03b5_PE + 2\u03b5_eff) "
        "+ \u03c6(\u03b5_air \u2212 \u03b5_eff) / (\u03b5_air + 2\u03b5_eff) = 0", 4)
    add_where(doc,
        "\u03c6는 공기 체적분율 (기공률), \u03b5_PE = n_PE\u00b2이며 "
        "n_PE = 1.517 (DSC 결정화도 67.8% 기반 결정질 n = 1.53과 비정질 n = 1.49의 "
        "가중 평균), \u03b5_air = 1.0입니다. "
        "측정된 n_eff로부터 기공률 \u03c6는 Eq. (4)를 수치적으로 풀어 구하였습니다. "
        "2상 모델을 선택한 이유는 결정질 (n = 1.53)과 비정질 (n = 1.49) PE 간 굴절률 "
        "대비가 \u0394n = 0.04에 불과하여 측정 정밀도 (\u03c3 \u2248 0.014) 이하이기 때문입니다. "
        "결정질과 비정질 PE 성분을 분리한 3상 모델을 평가하였으나 동일한 기공률 값을 "
        "산출하였으며 (\u0394\u03c6 < 0.1 퍼센트 포인트), 이는 현재 정밀도의 THz-TDS로는 "
        "결정질\u2013비정질 대비를 분해할 수 없으며 보다 단순한 2상 모델이 적절함을 "
        "확인해 줍니다. Bruggeman 모델은 Maxwell\u2013Garnett 근사보다 선호되었는데, "
        "후자는 희박한 포함물을 가정하여 여기서 나타나는 높은 기공률 값 "
        "(~44\u201359%)에서는 부정확해지기 때문입니다 [17]."
    )
    doc.add_paragraph(
        "모델은 20 \u00b0C에서의 제조사 사양 기공률 44%에 앵커링하였으며, "
        "2장 적층 구성 (공기 간극, 계면 효과)에 기인하는 일정 굴절률 오프셋 "
        "(\u0394n = +0.146)을 산출하였습니다. 이 오프셋은 EMA 역산 전에 모든 측정값에서 "
        "차감하였습니다. 오프셋의 불확도 (\u03c3(\u0394n) = 0.018, 20 \u00b0C에서 n의 "
        "표준편차로부터 전파)는 절대 기공률 불확도에 추가적으로 \u00b14\u20135 퍼센트 포인트를 "
        "기여합니다."
    )
    doc.add_paragraph(
        "시료군 S1\u2013S3과 S4\u2013S5 사이에서 체계적 오프셋 (\u0394n \u2248 0.106, "
        "p < 10\u207b\u00b9\u2076, 이표본 t-검정)이 확인되었으며, 온도 의존 기울기는 "
        "유사하였습니다 (\u22126.8 \u00d7 10\u207b\u2074 vs. \u22128.1 \u00d7 10\u207b\u2074 /\u00b0C). "
        "이 일정 오프셋은 약 12 \u00b5m의 추정 두께 차이에 해당하며, 이는 6위치 시료 홀더의 "
        "터릿 위치 5\u20136 (S4\u2013S5)과 위치 2\u20134 (S1\u2013S3) 간 광학 정렬 편차에 "
        "기인하는 것으로 추정됩니다. S4\u2013S5 값을 S1\u2013S3 기준선에 정렬하는 전역 "
        "주파수 의존 보정을 적용하였습니다. "
        "이 보정은 시료 간 표준편차를 0.051에서 0.014로 감소시켰으며, 0.45 THz에서 ANOVA "
        "F-통계량을 0.94 (p = 0.50, 비유의)에서 12.78 (p = 2.84 \u00d7 10\u207b\u2079, "
        "\u03b7\u00b2 = 0.74)로 개선하였습니다."
        " S1\u2013S3 부분군만을 사용한 분석 (N = 30)에서도 동일한 온도 경향이"
        " 확인되었으며 (R\u00b2 = 0.854, 기울기 = \u22126.8 \u00d7 10\u207b\u2074 /\u00b0C), 이는"
        " 관측된 온도 의존성이 오프셋 보정의 인위적 결과가 아님을 확인해 줍니다."
    )

    doc.add_heading("2.5. 머신러닝 온도 예측 (Machine learning temperature prediction)", level=2)
    doc.add_paragraph(
        "원시 시간영역 THz 신호로부터 직접 온도를 예측하는 머신러닝 파이프라인을 "
        "개발하였습니다. 각 신호에서 총 37개의 특성을 자동으로 추출하였으며, 5개 범주에 "
        "걸쳐 분류됩니다: (i) 피크 특성 (진폭, 시간, 피크-투-피크), "
        "(ii) 펄스 형상 (상승/하강 시간, 엔벨로프 FWHM, 비대칭성, 첨도), "
        "(iii) 신호 통계 (RMS, 에너지, SNR), "
        "(iv) 레퍼런스 상대 특성 (시간 지연, 진폭 비, 차이 신호), "
        "(v) 주파수 영역 특성 (대역 평균 스펙트럼 파워, 위상 기울기, 중심). "
        "특성 선별은 분산 팽창 인자 (VIF) 분석 [25]을 사용하여 임계값 10으로 수행하였으며, "
        "물리적으로 의미 있는 특성을 우선시하여 15개 특성을 선별하였습니다. "
        "(VIF 임계값 = 10; 전체 특성 목록은 부록 A, 15개 선별 특성 중 0이 아닌 "
        "13개 Lasso 계수는 부록 B를 참조, 정규화 매개변수 α = 0.1) "
        "5개 회귀 모델을 평가하였습니다 [19,25,28]: Ridge, Lasso [19] (동시 특성 선별 및 예측을 위한 "
        "L1 정규화 선형 회귀), SVR (RBF 커널), Random Forest, Gradient Boosting. "
        "두 가지 교차검증 전략을 비교하였습니다: (i) 기존의 하나씩 빼기 교차검증 "
        "(LOO-CV, 50 폴드), (ii) 온도별 교차검증 (LOTO-CV, 10 폴드)으로, "
        "후자는 주어진 온도의 5개 반복시편 전체를 동시에 제외합니다. "
        "LOTO-CV는 동일 온도 반복시편이 훈련 및 테스트 세트 모두에 나타나는 정보 누출을 "
        "방지하여, 미관측 온도에 대한 일반화 능력을 보다 엄격하게 평가합니다. "
        "구체적으로, 각 LOTO 폴드는 하나의 온도에서 측정된 5개 반복시편 전체를 제외합니다 "
        "(총 10 폴드, 각 N_train = 45, N_test = 5). "
        "제외된 온도가 훈련 세트에서 부재하므로, LOTO-CV는 모델의 보간 (또는 경계 온도 "
        "20 및 110 °C의 경우 외삽) 능력을 효과적으로 평가하여, 보수적인 일반화 성능 "
        "추정치를 제공합니다. "
        "본 모델은 추가 검증 없이 교정 범위 20–110 °C 외부에 적용해서는 안 됩니다. "
        "부트스트랩 신뢰구간 (B = 1000 리샘플)은 각 온도에서의 평균 굴절률에 대해 "
        "계산하였습니다."
    )
    doc.add_paragraph(
        "Lasso 회귀모델 [19]은 다음의 목적함수를 최소화합니다:"
    )
    add_equation(doc,
        "min  (1/2N) \u03a3(T\u1d62 \u2212 T\u0302\u1d62)\u00b2 + \u03b1 \u03a3|\u03b2\u2c7c|", 6)
    doc.add_paragraph(
        "여기서 T\u1d62는 측정 온도, T\u0302\u1d62는 예측 온도, "
        "\u03b2\u2c7c는 회귀 계수이며, \u03b1 = 0.1은 정규화 매개변수로서 "
        "훈련 데이터에 대한 5-폴드 교차검증을 통해 모델 복잡도와 예측 정확도의 "
        "균형을 고려하여 선택하였습니다. "
        "L1 페널티는 무관한 계수를 정확히 0으로 수렴시켜, 특성 선별과 회귀를 "
        "동시에 수행합니다."
    )
    doc.add_paragraph(
        "예측 방정식 (표준화된 특성 사용)은 다음과 같습니다:"
    )
    add_equation(doc,
        "T\u0302 = 65.0 + \u03a3 \u03b2\u2c7c \u00d7 (x\u2c7c \u2212 \u03bc\u2c7c) / \u03c3\u2c7c", 7)
    doc.add_paragraph(
        "여기서 \u03bc\u2c7c 및 \u03c3\u2c7c는 특성 j의 훈련 세트 평균 및 표준편차입니다. "
        "절편 (65.0 \u00b0C)은 훈련 데이터의 평균 온도에 해당합니다."
    )
    add_figure(doc, "fig_ml_flowchart",
               "Fig. 3. 머신러닝 온도 예측 파이프라인. 원시 THz 신호와 매칭 레퍼런스를 5개 범주에서 "
               "37개 특성을 추출하고, VIF 선택으로 15개로 축소 후 Lasso 회귀 (\u03b1 = 0.1)에 입력. "
               "LOTO-CV 평가 결과 R\u00b2 = 0.955, MAE = 5.1 \u00b0C. "
               "(ML temperature prediction pipeline.)",
               width=3.5)
    doc.add_paragraph(
        "제한된 시료 수 (N = 50개 시료, 10개 온도 수준)를 고려할 때, ML 분석은 "
        "탐색적 성격을 가지며, 생산 환경 배포를 의도한 것이 아닌 가장 유용한 신호 "
        "특성을 식별하는 데 목적이 있습니다. "
        "또한 전체 15개 특성 모델과 단일 특성 Ridge 회귀 기준선을 비교하는 제거 연구 "
        "(ablation study)를 수행하여 다중 특성 통합의 부가 가치를 정량화하였습니다."
    )

    # ════════════════════════════════════════════
    # 3. RESULTS AND DISCUSSION
    # ════════════════════════════════════════════
    doc.add_heading("3. 결과 (Results)", level=1)

    doc.add_heading("3.1. 시간영역 신호 분석 (Time-domain signal analysis)", level=2)
    doc.add_paragraph(
        "Fig. 4는 10개 온도 전체에 대한 대표적인 THz 시간영역 파형을 보여줍니다. "
        "PE 분리막을 투과한 펄스는 전체 온도 범위에서 높은 충실도를 유지하며, "
        "시스템 SNR은 약 52 dB입니다 (Fig. S1). "
        "두 패널 모두에서 체계적인 온도 의존 변화가 명확히 관찰됩니다. "
        "레퍼런스 신호 (상단 행)에서는 온도 증가에 따른 약간의 확장 및 진폭 감소가 "
        "THz 시스템 광학계 및 대기 흡수의 온도 의존성을 반영합니다. "
        "시료 신호 (하단 행)에서는 이러한 시스템 수준의 변화에 재료 고유 효과가 "
        "중첩됩니다: 양의 피크 진폭은 온도에 따라 단조 감소하며, 음의 피크는 "
        "체계적으로 더 이른 도달 시간 방향으로 시간적 이동을 보입니다. "
        "피크-투-피크 시간 간격은 20 \u00b0C의 약 0.55 ps에서 110 \u00b0C의 "
        "0.52 ps로 감소하며, 이는 기공률 증가에 따른 유효 광학 경로 길이 감소와 "
        "일치합니다. 매칭 레퍼런스 프로토콜 (Eq. 1)은 시스템 수준의 기여를 제거하여 "
        "정량적 분석을 위한 재료 응답을 분리합니다."
    )
    add_figure(doc, "fig02_peak_detail",
               "Fig. 4. 10개 온도 (20\u2013110 \u00b0C)에서의 THz 펄스 피크 상세. "
               "(a) 온도 증가에 따른 체계적 진폭 감소를 보이는 양의 피크 영역. "
               "(b) 음의 피크 영역. 검은선: 20 \u00b0C 레퍼런스. "
               "(THz pulse peak detail at 10 temperatures. (a) Positive peak region. "
               "(b) Negative peak region.)",
               width=5.0)

    doc.add_heading("3.2. 온도 의존 굴절률 (Temperature-dependent refractive index)", level=2)
    add_table_from_data(doc,
        ["T (\u00b0C)", "n @ 0.45 THz", "\u03c3(n)", "기공률 (%)", "\u03c3_total (%)"],
        [
            ["20",  "1.4282", "0.018", "44.0 ± 5.9", "5.9"],
            ["30",  "1.4230", "0.012", "45.0 ± 4.6", "4.6"],
            ["40",  "1.3978", "0.012", "49.7 ± 4.7", "4.7"],
            ["50",  "1.3775", "0.014", "53.5 ± 5.1", "5.1"],
            ["60",  "1.3790", "0.014", "53.3 ± 5.1", "5.1"],
            ["70",  "1.3751", "0.015", "54.0 ± 5.2", "5.2"],
            ["80",  "1.3688", "0.013", "55.2 ± 4.9", "4.9"],
            ["90",  "1.3728", "0.014", "54.4 ± 5.2", "5.2"],
            ["100", "1.3515", "0.011", "58.5 ± 4.6", "4.6"],
            ["110", "1.3663", "0.014", "55.7 ± 5.1", "5.1"],
        ],
        "Table 3. 0.45 THz에서의 온도 의존 굴절률 및 기공률 "
        "(5회 반복 측정의 평균 \u00b1 \u03c3). \u03c3_total은 n_offset 불확도 전파를 포함합니다. "
        "(Temperature-dependent refractive index and porosity at 0.45 THz.)"
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Fig. 5는 온도의 함수로서 0.45 THz에서의 굴절률과 EMA 유도 유효 기공률을 "
        "함께 나타냅니다. 굴절률은 20 \u00b0C의 1.428 \u00b1 0.018에서 100 \u00b0C의 "
        "1.352 \u00b1 0.011로 단조 감소하였습니다 "
        "(온도와의 선형 상관 R\u00b2 = 0.804). "
        "전이 온도는 구간별 선형 회귀로 결정하였습니다. "
        "구체적으로, 분기점 T_break는 30\u201390 \u00b0C 범위에서 0.1 \u00b0C 분해능으로 "
        "모든 후보 값에 대해 총 잔차 제곱합 (RSS)을 최소화하여 최적화하였습니다. "
        "최적 분기점은 T_break = 51.8 \u00b0C, RSS = 3.0 \u00d7 10\u207b\u2074에서 발견되었으며, "
        "두 피팅 직선의 교점은 T_onset = 55.2 \u00b0C를 산출합니다 "
        "(구간 1: n = 1.452 \u2212 1.52 \u00d7 10\u207b\u00b3 T, R\u00b2 = 0.987; "
        "구간 2: n = 1.373 \u2212 3.24 \u00d7 10\u207b\u2074 T, R\u00b2 = 0.476). "
        "2구간 거동이 관찰됩니다: ~55 \u00b0C 이하에서 굴절률은 급격히 감소하며 "
        "(기울기 = \u22121.52 \u00d7 10\u207b\u00b3 /\u00b0C), 55 \u00b0C 이상에서는 "
        "감소가 완화됩니다 (기울기 = \u22123.24 \u00d7 10\u207b\u2074 /\u00b0C). "
        "55.2 \u00b0C의 전이 온도는 결정질 라멜라 내 사슬 이동도 개시와 관련된 "
        "폴리에틸렌의 \u03b1c 결정 이완 [13]과 일치합니다. 이 전이는 DSC 용융 온도 "
        "(Tm = 139.4 \u00b0C, 결정화도 = 67.8%)보다 훨씬 낮으며, 비정질 영역 "
        "자유 부피 팽창 및 라멜라 간 기공 확장과 같은 중요한 미세구조 변화가 "
        "용융점보다 훨씬 낮은 온도에서 시작됨을 나타냅니다. "
        "2구간 거동 (55 \u00b0C 이하 급격한 변화, 이상 완만)은 비정질상 이완이 "
        "초기 기공률 증가를 주도하고, 결정질 라멜라가 T_onset 이상에서 "
        "구조적 저항을 제공한다는 해석과 일치합니다."
    )
    doc.add_paragraph(
        "110 °C에서 n의 약간의 증가 (1.352에서 1.366)가 관찰되며, "
        "이는 측정 불확도 범위 내에 있습니다 (Δn = 0.015 < σ = 0.014). "
        "이 비단조적 거동은 용융 구간에 접근하면서의 측정 변동성에 기인하며, "
        "반전의 크기 (0.015)가 측정 불확도에 근접하므로 현재 데이터로는 "
        "물리적 메커니즘에 확정적으로 귀인시킬 수 없습니다."
    )
    add_figure(doc, "fig03_n_vs_temp_ema",
               "Fig. 5. 온도의 함수로서 0.45 THz 굴절률 n (좌축, 검정) 및 유효 기공률 "
               "(우축, 빨강). 개별 시료 (S1\u2013S5)는 색상 마커로, 평균 \u00b1 \u03c3는 "
               "검은 원과 회색 밴드로 표시. 점선: 55 \u00b0C 전이점의 2구간 선형 피팅. "
               "(Refractive index n at 0.45 THz and effective porosity vs. temperature.)",
               width=4.5)

    doc.add_heading("3.3. 광대역 광학 스펙트럼 (Broadband optical spectra)", level=2)
    doc.add_paragraph(
        "Fig. 6은 광대역 굴절률 및 흡수계수 스펙트럼을 나타냅니다. "
        "굴절률은 0.3\u20131.5 THz 범위에서 비교적 평탄한 분산을 보이며, "
        "고주파로 갈수록 점진적으로 감소합니다. 온도 의존 이동은 전체 스펙트럼 "
        "범위에서 일관성을 보입니다. 흡수계수는 측정 대역폭 전체에서 낮으며 "
        "(\u03b1 < 15 cm\u207b\u00b9), 저손실 유전 재료의 특성을 나타냅니다."
    )
    doc.add_paragraph(
        "굴절률 스펙트럼은 0.3\u20131.0 THz 범위에서 약한 정상 분산을 보이며, "
        "고주파에서의 점진적 감소는 다공성 고분자 매질의 특성입니다 [12,23,24,26]. "
        "온도 유도 이동은 스펙트럼 범위 전체에서 놀라울 정도로 균일합니다: "
        "굴절률은 20에서 100 \u00b0C 사이 모든 주파수에서 약 0.06\u20130.08 감소하며, "
        "이는 기저 메커니즘 (기공 확장)이 주파수 선택적 특성을 도입하기보다 "
        "유효 유전 응답에 균일하게 영향을 미침을 나타냅니다. "
        "이러한 스펙트럼 균일성은 주파수 독립 EMA 모델 (Eq. 4)의 사용을 뒷받침하며, "
        "주요 정량 분석을 위한 단일 대표 주파수 (0.45 THz) 선택을 정당화합니다. "
        "최적 상관 주파수 3개 (0.45, 1.09, 1.63 THz)는 사용 가능한 대역폭을 포괄하며 "
        "일관된 온도 경향을 제공하여, 관측된 변화의 광대역 특성을 추가로 확인합니다."
    )
    doc.add_paragraph(
        "0.45 THz에서의 흡수계수 값은 작으며 (0.0\u20130.9 cm\u207b\u00b9), "
        "측정의 노이즈 플로어에 근접합니다. 저주파에서 위상 잡음으로 인해 발생하는 "
        "음수 값은 0으로 절단 (clip)하였습니다. \u03b1의 불확도는 반복시편 간 변동성을 "
        "기반으로 \u00b11\u20132 cm\u207b\u00b9로 추정됩니다."
    )
    add_figure(doc, "fig04_optical_spectra",
               "Fig. 6. 10개 온도에서의 광대역 광학 스펙트럼. "
               "(a) 굴절률 n(f). 수직 점선은 최적 상관 주파수 3개 "
               "(0.45, 1.09, 1.63 THz)를 표시. (b) 흡수계수 \u03b1(f). "
               "(Broadband optical spectra at 10 temperatures.)",
               width=5.0)

    doc.add_heading("3.4. 기공률 변화 및 열팽창 (Porosity evolution and thermal expansion)", level=2)
    doc.add_paragraph(
        "2상 Bruggeman 모델 (Eq. 4)을 역산하여 얻은 유효 기공률은 "
        "20 \u00b0C의 44.0% (제조사 사양에 앵커링)에서 100 \u00b0C의 58.5%로 "
        "증가합니다 (Fig. 7). 온도 의존성은 열팽창 모델을 사용하여 피팅하였습니다:"
    )
    add_equation(doc,
        "\u03c6(T) = 1 \u2212 (1 \u2212 \u03c6\u2080) / (1 + \u03b2 \u0394T)", 5)
    add_where(doc,
        "\u03c6\u2080는 초기 기공률, \u03b2는 기공 구조의 유효 열팽창계수, "
        "\u0394T = T \u2212 20 \u00b0C입니다."
    )
    doc.add_paragraph(
        "피팅 결과 \u03c6\u2080 = 45.2% 및 \u03b2 = 3.50 \u00d7 10\u207b\u00b3 /\u00b0C를 "
        "얻었습니다. 이 유효 기공 팽창계수는 벌크 PE 체적 팽창계수 "
        "(~5 \u00d7 10\u207b\u2074 /\u00b0C [12])보다 약 6\u20138배 크며, 이는 벌크 "
        "열팽창이 아닌 기공 확장이 관측된 광학적 변화를 지배함을 나타냅니다 [7,10,27]."
    )
    doc.add_paragraph(
        "주목할 점은, 확인된 전이 온도 55 °C에 선행하여 20에��� 40 °C 사이에서 "
        "약 5.7 퍼센트 포인트의 기공률 증가가 관찰되었습니다. 이 초기 단계 "
        "변화의 원인은 본 연구에서 독립적으로 검증되지 않았습니다. "
        "일축 연신 PE 필름에서 \u03b1c 이완 온도 이하의 유사한 초기 단계 기공률 "
        "변화가 보고된 바 있습니다 [12,13]."
    )
    add_figure(doc, "fig06_porosity_vs_temp",
               "Fig. 7. 온도의 함수로서 유효 기공률. 검은 원: EMA 역산 결과 "
               "(5회 반복의 평균 \u00b1 \u03c3). 빨간 점선: 열팽창 모델 피팅 "
               "(\u03b2 = 3.50 \u00d7 10\u207b\u00b3 /\u00b0C). 수평 회색선: 제조사 사양 (44%). "
               "(Effective porosity as a function of temperature.)",
               width=4.0)

    doc.add_paragraph(
        "절대 기공률 값은 상당한 불확도를 수반하지만 "
        "(\u03c3_total \u2248 5\u20136 퍼센트 포인트, 주로 20 \u00b0C에서 N = 5 반복시편으로 "
        "계산된 n_offset 불확도 \u03c3 = 0.018에 기인), 상대 기공률 변화 "
        "(\u0394\u03c6 \u2248 80 \u00b0C에 걸쳐 15 퍼센트 포인트)는 명확히 분해되며 "
        "총 불확도를 약 3배 초과합니다. 절대 기공률 값이 아닌 이 상대적 경향이 "
        "주요 물리적 발견을 구성하며, 건식 공정 PE 분리막의 알려진 열팽창 "
        "거동과 일치합니다."
    )

    doc.add_heading("3.5. 통계적 ���증 (Statistical validation)", level=2)
    doc.add_paragraph(
        "10개 온도 수준을 집단으로 하여 (집단당 N = 5 반복, 총 N = 50) "
        "0.3\u20132.0 THz 범위의 각 주파수에서 일원 분산분석 (one-way ANOVA)을 "
        "수행하였습니다 (Fig. 8). "
        "0.45 THz에서 ANOVA는 F = 12.78, p = 2.84 \u00d7 10\u207b\u2079, "
        "효과 크기 \u03b7\u00b2 = 0.74 (큰 효과)를 산출하였습니다. 전체 스펙트럼 범위에서 "
        "97.4%의 주파수 (465개 중 453개)가 통계적으로 유의한 온도 의존성을 "
        "보였습니다 (p < 0.05). "
        "비유의 좁은 대역 (0.30\u20130.55 THz)은 측정 잡음이 높은 저주파 가장자리와 "
        "일치합니다."
    )
    add_figure(doc, "fig13_anova_freq_scan",
               "Fig. 8. 주파수 분해 ANOVA 결과. (a) 주파수의 함수로서 p-값 (로그 척도). "
               "파란 음영: 유의 영역 (p < 0.05). (b) 효과 크기 \u03b7\u00b2. "
               "수평 빨간 점선: 큰 효과 임계값 (0.14). "
               "(Frequency-resolved ANOVA results.)",
               width=4.0)

    doc.add_paragraph(
        "평균 굴절률에 대한 부트스트랩 신뢰구간 [22] (B = 1000)은 "
        "20 \u00b0C [1.411, 1.443]와 100 \u00b0C [1.340, 1.360] 사이에서 "
        "비중첩 95% CI를 확인하여, 온도 효과의 유의성에 대한 추가 증거를 "
        "제공하였습니다."
    )
    doc.add_paragraph(
        "인접한 10 °C 단계 간 쌍별 비교 [21]는 통계적 유의성에 도달하지 못하지만 "
        "(시료 간 변동성에 비해 단계별 변화가 작기 때문), 전체 온도에 대한 "
        "일원 분산분석은 매우 유의하며 (F = 12.78, p = 2.84 × 10⁻⁹, η² = 0.74), "
        "20–100 °C 전체 범위에 걸친 누적 변화는 명확히 분해됩니다 "
        "(t-검정에 의해 p = 9.8 × 10⁻⁵)."
    )
    doc.add_paragraph(
        "최적 상관 주파수 3개 (0.45, 1.09, 1.63 THz)는 0.3–2.0 THz 범위의 "
        "모든 주파수를 스캔하고 n 대 T 선형 상관의 R²로 순위를 매겨 식별하였습니다. "
        "이 주파수 선택은 모델 훈련 전에 전체 데이터셋에 대해 수행되었으나, "
        "LOTO-CV 평가는 제외된 온도 블록을 사용하므로 ML 성능 지표가 이 선택에 의해 "
        "부풀려지지 않습니다."
    )

    doc.add_heading("3.6. 위상 스펙트럼 및 전달함수 (Phase spectrum and transfer function)", level=2)
    doc.add_paragraph(
        "언래핑된 위상 차이 \u0394\u03c6(f)와 전달함수 크기 |H(\u03c9)|를 "
        "Fig. 9에 나타냅니다. 위상 차이는 온도 의존 기울기를 가진 근선형 "
        "주파수 의존성을 보이며, 유효 광학 경로 길이의 변화를 직접 반영합니다. "
        "전달함수 크기는 1에 가깝게 유지되며 (|H| \u2248 0.95\u20131.0), "
        "얇은 PE 필름의 낮은 흡수와 일치합니다."
    )
    doc.add_paragraph(
        "선형 위상 기울기는 군 굴절률에 직접 비례하며, 이는 Lasso 모델의 "
        "최상위 특성 (phase_slope_low, 중요도 = 14.5)에 부호화된 주요 물리량입니다. "
        "온도가 증가하면 위상 기울기가 감소 (덜 음수)하여, 기공 확장에 따른 "
        "광학 경로 길이 감소를 반영합니다. "
        "전달함수 크기 |H(\u03c9)| \u2248 0.95\u20131.0은 얇은 PE 필름에서의 "
        "최소 흡수 손실을 나타내며, 3.3절에서 보고된 낮은 흡수계수와 일치합니다. "
        "고주파 (>1.5 THz)에서의 1로부터의 미세한 편차는 다공성 미세구조로부터의 "
        "산란 증가를 반영하며, 기공 치수가 THz 파장의 상당 부분에 접근할수록 "
        "더욱 두드러집니다."
    )
    add_figure(doc, "fig11_phase_spectrum",
               "Fig. 9. (a) 시료와 매칭 레퍼런스 간 언래핑된 위상 차이 \u0394\u03c6(f). "
               "(b) 전달함수 크기 |H(\u03c9)|. 색상은 온도를 나타냄 "
               "(파랑: 20 \u00b0C, 빨강: 110 \u00b0C). "
               "(Unwrapped phase difference and transfer function magnitude.)",
               width=5.0)

    doc.add_heading("3.7. 머신러닝 온도 예측 (Machine learning temperature prediction)", level=2)
    doc.add_paragraph(
        "Table 1은 두 가지 교차검증 전략 하에서 5개 회귀 모델의 성능을 비교합니다. "
        "선형 모델 (Lasso, Ridge)은 두 CV 방식 모두에서 비선형 모델 "
        "(SVR, Random Forest, Gradient Boosting)을 크게 상회하며, 이는 연구 범위 내에서 "
        "온도\u2013THz 신호 관계가 주로 선형임을 나타냅니다."
    )
    add_table_from_data(doc,
        ["모델 (Model)", "LOO-CV R\u00b2", "LOO-CV MAE (\u00b0C)", "LOTO-CV R\u00b2", "LOTO-CV MAE (\u00b0C)"],
        [
            ["Ridge",             "0.974", "3.81", "0.952", "5.20"],
            ["Lasso",             "0.975", "3.71", "0.955", "5.08"],
            ["SVR (RBF)",         "0.875", "7.52", "0.566", "13.60"],
            ["Random Forest",     "0.845", "7.21", "0.648", "11.84"],
            ["Gradient Boosting", "0.844", "6.32", "0.652", "11.55"],
        ],
        "Table 1. LOO-CV 및 LOTO-CV 하에서의 모델 성능 비교. "
        "(Model performance comparison under LOO-CV and LOTO-CV.)"
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "결정적으로, LOO-CV와 LOTO-CV 간 성능 격차는 선형 모델에서 작지만 "
        "(\u0394R\u00b2 \u2248 0.02), 비선형 모델에서는 크며 (\u0394R\u00b2 = 0.19\u20130.31), "
        "이는 후자가 동일 온도 반복시편이 훈련 및 테스트 세트 모두에 나타날 때 "
        "정보 누출의 영향을 받음을 보여줍니다."
    )

    add_figure(doc, "fig_ml_prediction_scatter",
               "Fig. 10. LOTO-CV 하에서 Lasso 모델의 예측 온도 대 실제 온도 "
               "(R\u00b2 = 0.955, MAE = 5.1 \u00b0C). 색상은 실제 온도를 나타냅니다. "
               "(Predicted vs. actual temperature for the Lasso model under LOTO-CV.)",
               width=3.5)

    doc.add_paragraph(
        "Fig. 11은 Lasso 모델의 특성 중요도 순위를 보여줍니다. 상위 3개 특성은 "
        "phase_slope_low (저주파 위상 기울기, 군 굴절률에 비례), "
        "spec_mean_low (0.3\u20130.8 THz 평균 스펙트럼 진폭, 투과율 관련), "
        "spec_centroid (스펙트럼 질량 중심)입니다. 이 특성들은 명확한 물리적 "
        "해석을 가집니다: 위상 기울기는 광학 경로 길이 변화를 직접 부호화하며, "
        "스펙트럼 진폭과 중심은 기공 구조 변화에 의해 유도된 흡수 및 산란 "
        "수정을 반영합니다."
    )
    add_figure(doc, "fig_ml_feature_importance",
               "Fig. 11. 온도 예측 모델의 특성 중요도 (Lasso 계수 크기). "
               "별표 (\u2605): 사전 선택된 물리 우선 특성. "
               "(Feature importance for the temperature prediction model.)",
               width=4.0)

    doc.add_paragraph(
        "제거 연구 (ablation study, Table 2)는 다중 특성 통합의 부가 가치를 확인합니다. "
        "최적 단일 특성 (phase_slope_low, 군 굴절률에 비례)은 LOTO-CV 하에서 "
        "R\u00b2 = 0.67을 달성하는 반면, 전체 15개 특성 Lasso 모델은 "
        "R\u00b2 = 0.955에 도달하여 43 퍼센트 포인트 향상을 보입니다. "
        "이는 단일 THz 신호 특성만으로는 전체 온도 정보를 포착할 수 없으며, "
        "위상, 진폭, 스펙트럼 특성의 조합이 시너지적 표현을 제공함을 보여줍니다."
    )
    doc.add_paragraph(
        "선형 모델 (Lasso, Ridge)이 비선형 모델 (SVR, Random Forest, "
        "Gradient Boosting)에 비해 우세한 것은 물리적 통찰을 제공합니다: "
        "THz 신호 특성과 온도 간의 관계는 PE 용융 전이 (Tm = 139 \u00b0C) 이하인 "
        "20\u2013110 \u00b0C 범위 내에서 주로 선형입니다. 이 선형성은 n(T) 및 "
        "\u03c6(T) 모두 준선형 경향을 보이는 물리 기반 분석과 일치합니다. "
        "비선형 모델의 LOO-CV와 LOTO-CV 간 큰 성능 격차 "
        "(\u0394R\u00b2 = 0.19\u20130.31) 대 선형 모델 (\u0394R\u00b2 \u2248 0.02)은 "
        "트리 기반 및 커널 기반 모델이 훈련 세트에 동일 온도 반복시편이 있을 때 "
        "온도 특이적 잡음 패턴에 과적합하는 반면, 선형 모델은 미관측 온도에 "
        "더 강건하게 일반화함을 보여줍니다."
    )
    add_table_from_data(doc,
        ["특성 (Features)", "R\u00b2 (LOTO)", "MAE (\u00b0C)"],
        [
            ["전체 15개 (Lasso)", "0.955", "5.08"],
            ["phase_slope_low 단일", "0.666", "12.30"],
            ["spec_std_high 단일", "0.509", "15.59"],
            ["spec_centroid 단일", "0.393", "18.42"],
            ["spec_mean_low 단일", "0.229", "20.41"],
        ],
        "Table 2. 제거 연구: 전체 모델 대 단일 특성 기준선 (LOTO-CV). "
        "(Ablation study: full model vs. single-feature baselines.)"
    )
    doc.add_paragraph()

    # ════════════════════════════════════════════
    # 4. DISCUSSION
    # ════════════════════════════════════════════
    doc.add_heading("4. 고찰 (Discussion)", level=1)

    doc.add_heading("4.1. 열적 전이의 물리적 해석 (Physical interpretation of thermal onset)", level=2)
    doc.add_paragraph(
        "구간별 선형 회귀로 확인된 55.2 \u00b0C의 열적 전이 온도는 결정질 라멜라 내 "
        "사슬 이동도 개시와 관련된 폴리에틸렌의 \u03b1c 결정 이완 범위 "
        "(50\u201380 \u00b0C)에 해당합니다 [13]. "
        "T_onset 이하에서 굴절률의 급격한 감소 (기울기 = \u22121.52 \u00d7 10\u207b\u00b3 /\u00b0C)는 "
        "비정질 라멜라 간 영역이 상당한 자유 부피 팽창을 겪어 기공 확장으로 "
        "이어짐을 시사합니다. T_onset 이상에서 완화된 기울기 "
        "(\u22123.24 \u00d7 10\u207b\u2074 /\u00b0C)는 결정질 라멜라가 추가 기공 확장에 "
        "대해 구조적 저항을 제공함을 나타내며, 건식 공정 PE 분리막의 "
        "반결정질 특성과 일치합니다."
    )
    doc.add_paragraph(
        "유효 기공 팽창계수 \u03b2 = 3.50 \u00d7 10\u207b\u00b3 /\u00b0C는 "
        "벌크 PE 체적 열팽창계수 (~5 \u00d7 10\u207b\u2074 /\u00b0C)보다 약 1자릿수 "
        "큽니다. 이러한 증폭은 기공 확장이 벌크 재료의 등방성 열팽창이 아닌, "
        "연신된 고분자 네트워크의 이방성 응력 이완에 의해 구동되기 때문에 "
        "발생합니다. 건식 연신 제조 공정은 잔류 탄성 응력을 도입하며, 이는 "
        "가열 시 점진적으로 해방되어 연신 방향으로 슬릿 형상 기공을 "
        "우선적으로 확장합니다."
    )
    doc.add_paragraph(
        "n(T)에서 관찰된 2구간 거동은 55 \u00b0C 부근에서 THz-TDS로 검출 가능한 "
        "구조적 전이를 나타냅니다. 이 온도 이하에서 급격한 n 감소는 약 "
        "0.29 %p/\u00b0C의 기공률 증가 속도에 해당하며, 55 \u00b0C 이상에서는 "
        "속도가 0.05 %p/\u00b0C로 감소합니다. 이러한 속도 변화는 비정질상 주도 "
        "팽창에서 결정질 골격 안정화 거동으로의 전환과 일치합니다."
    )
    doc.add_paragraph(
        "확인된 T_onset (55.2 \u00b0C)는 DSC 용융 온도 "
        "(Tm = 139.4 \u00b0C)보다 상당히 낮으며, 측정 가능한 미세구조 재편이 "
        "용융점보다 약 84 \u00b0C 낮은 온도에서 시작됨을 나타냅니다. "
        "T_onset과 Tm 사이의 이 큰 간격은 분리막이 구조적 건전성을 "
        "유지하면서도 THz-TDS로 기공률 변화를 검출할 수 있는 영역을 정의합니다."
    )

    doc.add_heading("4.2. 다중 영역 특성 분석 (Multi-domain feature analysis)", level=2)
    doc.add_paragraph(
        "Fig. 12는 12개 시간영역, 주파수영역, 엔벨로프 특성과 온도 간의 "
        "상관관계를 요약합니다. 상승 시간 (R\u00b2 = 0.951)과 스펙트럼 중심 "
        "(R\u00b2 = 0.929)이 가장 강한 선형 상관을 보이며, 주파수영역 특성이 "
        "지배적인 ML 특성 중요도 분석과 일치합니다."
    )
    doc.add_paragraph(
        "상관 순위 (rise_time_ps: R² = 0.951, 1위)와 ML 특성 중요도 "
        "(phase_slope_low: 1위, rise_time: 12위) 간의 외관상 불일치는 "
        "단변량 분석과 다변량 분석의 차이를 반영합니다. rise_time이 온도와 "
        "가장 강한 개별 선형 상관을 보이지만, phase_slope_low와 높은 다중공선성을 "
        "공유합니다 (둘 다 군 지연 정보를 부호화). 다변량 Lasso 모델에서 "
        "phase_slope_low가 rise_time이 전달하는 정보를 흡수하여, 후자의 계수가 "
        "거의 0이 됩니다."
    )
    doc.add_paragraph(
        "다중 영역 특성 분석은 온도 정보가 상보적인 신호 특성에 걸쳐 분포되어 "
        "있음을 보여줍니다. 시간영역 특성 (rise_time, dt_pos)은 군 지연 및 "
        "펄스 형상 변화를, 주파수영역 특성 (phase_slope, spec_centroid)은 "
        "에너지의 스펙트럼 재분배를 부호화합니다. 엔벨로프 특성 (env_fwhm, "
        "env_asymmetry)은 펄스 확장 효과를 반영합니다. 5개 영역 모두의 특성이 "
        "ML 모델에 기여한다는 사실 (Table B1)은 포괄적인 다중 영역 접근법이 "
        "단일 영역 분석보다 더 많은 온도 정보를 추출함을 시사합니다. "
        "이 발견은 THz 기반 열 모니터링에 실용적 시사점을 가집니다: "
        "굴절률 추출 (신중한 위상 언래핑과 두께 정보 필요)에만 의존하기보다, "
        "특성 기반 ML 접근법이 원시 시간영역 신호로부터 직접 강건한 온도 추정을 "
        "제공할 수 있습니다."
    )
    doc.add_paragraph(
        "ML 모델에서 주파수영역 특성의 지배에 대한 물리적 기반은 Lasso 계수와 "
        "기저 광학 물성 간의 관계를 통해 이해할 수 있습니다. 최상위 특성인 "
        "phase_slope_low (\u03b2 = +14.5)는 0.3\u20130.8 THz 대역의 군 굴절률 n_g에 "
        "직접 비례합니다. 2순위 특성인 spec_mean_low (\u03b2 = +11.0)는 투과 계수 "
        "|H(\u03c9)|와 관련된 평균 스펙트럼 진폭을 부호화합니다. 표준화된 계수의 "
        "부호와 크기는 센터링 및 스케일링 후 다변량 상관 구조를 반영합니다."
    )
    add_figure(doc, "fig08_correlation_summary",
               "Fig. 12. 다중 영역 특성 상관관계 요약. (a) 온도에 걸친 정규화된 "
               "특성 히트맵. (b) 유의 수준을 포함한 R\u00b2 값. "
               "(Multi-domain feature correlation summary.)",
               width=5.0)

    doc.add_heading("4.3. 열 모니터링의 실용적 시사점 (Practical implications for thermal monitoring)", level=2)
    doc.add_paragraph(
        "굴절률의 온도 민감도 (dn/dT \u2248 \u22128.5 \u00d7 10\u207b\u2074 /\u00b0C)는 "
        "도당 약 0.16 퍼센트 포인트의 기공률 민감도에 해당하며, 이는 현재 THz-TDS "
        "설정 (\u03c3(n) = 0.014)으로 검출 가능합니다. "
        "ML 기반 예측 모델 (Lasso, MAE = 5.1 \u00b0C)은 명시적 광학 물성 추출 없이 "
        "원시 THz 신호로부터 온도 추정이 가능함을 입증하며, 이는 처리 속도가 "
        "중요한 인라인 품질 관리에 유리합니다. "
        "ML 모델에서 주파수영역 특성 (phase_slope_low, spec_mean_low)의 지배는 "
        "전체 시간영역 파형 처리보다 간소화된 스펙트럼 분석이 실용적 온도 "
        "모니터링 응용에 충분할 수 있음을 시사합니다."
    )
    doc.add_paragraph(
        "입증된 \u00b15 \u00b0C의 온도 분해능 (LOTO-CV 하 Lasso MAE 기준)은 "
        "인라인 제조 품질 관리에 충분하며, 셀 조립 전 제어된 온도에서 "
        "분리막 롤을 검사하는 데 적합합니다. THz-TDS의 비접촉 특성은 "
        "분리막 필름과의 물리적 접촉을 피해 기계적 손상을 방지해야 하는 "
        "연속 생산 라인에 유리합니다."
    )
    doc.add_paragraph(
        "제조 품질 관리에서, 기공률 민감도 (\u0394\u03c6/\u0394T \u2248 "
        "0.16 %p/\u00b0C)와 THz 측정 정밀도의 결합은 일정 온도에서 약 "
        "\u00b11 퍼센트 포인트의 배치 간 기공률 변동을 검출할 수 있음을 "
        "시사합니다. 이는 수은 압입 기공률 측정법의 정확도 (~\u00b11%)에 "
        "필적하지만, 비파괴적으로 수 시간이 아닌 수 초의 측정 시간으로 "
        "달성할 수 있습니다. ML 예측 파이프라인을 자동화 검사 시스템에 "
        "통합하려면 여러 분리막 유형 및 생산 배치를 포괄하는 보다 크고 "
        "다양한 데이터셋에 대한 훈련이 필요하며, 이는 향후 연구의 "
        "우선과제로 식별되었습니다."
    )

    doc.add_heading("4.4. 한계점 (Limitations)", level=2)
    doc.add_paragraph(
        "본 연구의 몇 가지 한계점을 주의해야 합니다. 첫째, EMA 역산으로 도출된 "
        "기공률 값은 유효 광학 기공률이며, 직접적인 구조 방법 (예: SEM 단면 분석 "
        "또는 수은 압입 기공률 측정법)으로 독립적으로 검증되지 않았습니다. "
        "제조사 사양 (44%)에 대한 앵커링이 물리적으로 근거있는 기준선을 제공하지만, "
        "절대 기공률 값은 전파된 n_offset 오차로 인해 \u00b15\u20136 퍼센트 포인트의 "
        "불확도를 수반합니다."
    )
    doc.add_paragraph(
        "둘째, 2장 적층은 굴절률 오프셋에 기여하는 추가 미지수 (층간 공기 간극, "
        "접촉 압력 변동)를 도입합니다. 단일 시트 측정이 이 불확도를 줄일 수 있으나, "
        "20 \u00b5m 두께에서의 제한된 SNR로 인해 실현 가능하지 않았습니다."
    )
    doc.add_paragraph(
        "셋째, Bruggeman EMA는 등방성 구형 포함물을 가정하지만, 건식 공정 PE "
        "분리막은 일축 연신으로 인해 배향된 슬릿 형상 기공을 가집니다. 따라서 "
        "보고된 기공률 값은 유효 등방성 근사를 나타내며, 다른 기법으로 측정된 "
        "실제 방향성 기공률과 다를 수 있습니다."
    )
    doc.add_paragraph(
        "넷째, 모든 측정은 단일 가열 과정 (20 \u2192 110 \u00b0C)에서 수행되었습니다. "
        "냉각 사이클 측정은 수행되지 않았으므로, 관측된 미세구조 변화의 가역성 및 "
        "가능한 열 이력 효과는 미탐구 상태로 남아 있습니다. "
        "기공률 회복의 냉각 경로 의존성을 조사하는 것은 향후 연구의 중요한 "
        "방향입니다."
    )
    doc.add_paragraph(
        "다섯째, ML 모델은 하나의 장비로 측정된 단일 분리막 유형 (MS-DPS 20B)에 "
        "대해 훈련 및 검증되었습니다. 다른 분리막 재료 (PP, 습식 공정 PE, "
        "세라믹 코팅) 또는 측정 시스템에 대한 일반화는 입증되지 않았습니다. "
        "여섯째, 시료 반복시편 (S1\u2013S5)은 터릿 회전을 통해 순차적으로 "
        "측정되었으며, 위치 의존적 정렬 변동의 가능성을 완전히 배제할 수 없습니다. "
        "다만, 매칭 레퍼런스 프로토콜 및 복귀 레퍼런스 드리프트 검증 "
        "(스펙트럼 편차 < 0.6%)이 체계적 오차를 완화합니다."
    )

    doc.add_heading("5. 결론 (Conclusions)", level=1)
    doc.add_paragraph(
        "본 연구는 PE 이차전지 분리막에서 온도 유도 미세구조 변화에 대한 "
        "THz-TDS의 비파괴적, 정량적 모니터링 능력을 입증합니다. "
        "주요 발견은 다음과 같습니다:"
    )
    conclusions = [
        "0.45 THz에서의 굴절률은 20\u2013110 \u00b0C 범위에서 1.43에서 1.35로 감소하며, "
        "이는 2상 Bruggeman EMA에 의해 결정된 유효 기공률 44%에서 59%로의 "
        "증가에 해당합니다.",
        "PE의 \u03b1c 결정 이완과 일치하는 55.2 \u00b0C의 열적 전이 온도가 "
        "확인되었으며, 이 온도 이하에서 급격한 미세구조 변화가 발생합니다.",
        "ANOVA는 THz 스펙트럼 범위의 97.4%에서 통계적으로 유의한 온도 의존성을 "
        "확인하며, 큰 효과 크기를 보입니다 (\u03b7\u00b2 = 0.74, 0.45 THz 기준).",
        "머신러닝 분석은 온도\u2013신호 관계가 주로 선형임을 보여주며, "
        "위상 기울기와 저주파 스펙트럼 파워가 주요 지표입니다 "
        "(Lasso R\u00b2 = 0.955, LOTO-CV 기준).",
        "물리 기반 및 데이터 기반 접근법의 결합은 상보적 통찰을 제공합니다: "
        "EMA는 기공률 변화를 정량화하고, ML은 잠재적 실시간 온도 모니터링 "
        "응용을 위한 가장 유용한 신호 특성을 식별합니다.",
    ]
    for c in conclusions:
        doc.add_paragraph(c, style="List Number")

    doc.add_paragraph(
        "현재 측정 조건에서 굴절률의 온도 민감도 "
        "(\u2248 \u22128.5 \u00d7 10\u207b\u2074 /\u00b0C, 전체 범위 평균)는 현재 정밀도 "
        "(\u03c3 \u2248 0.014)를 고려할 때 약 5 \u00b0C의 검출 한계를 시사하며, "
        "이는 ML 예측 정확도 (MAE = 5.1 \u00b0C)와 일치합니다. "
        "이러한 발견의 일반화 가능성을 확인하기 위해 독립적인 분리막 유형 및 "
        "측정 시스템에 대한 검증이 필요합니다. "
        "이러한 결과는 THz-TDS를 이차전지 분리막 건전성의 현장 열 모니터링을 "
        "위한 실현 가능한 비접촉 기술로 확립합니다. 잠재적 응용 분야는 다음과 "
        "같습니다: (i) 제조 중 분리막 기공률의 인라인 품질 검사, "
        "(ii) 분리막 생산 중 열 유도 기공률 변화의 공정 모니터링, "
        "(iii) 비파괴 기공률 추적이 필요한 가속 노화 연구."
    )
    doc.add_paragraph(
        "THz 방사선은 금속 전극을 투과할 수 없으므로, 조립된 전지 셀 내부의 "
        "분리막을 직접 현장 측정하는 것은 불가능함을 유의해야 합니다. "
        "따라서 제안된 방법의 적용 범위는 셀 조립 전 제조 및 품질 관리 단계에서의 "
        "분리막 특성 평가로 제한됩니다. 다음 단계로, 단일 셀 프로토타입의 중앙에 "
        "측정 개구부를 도입하여 충\u00b7방전 사이클에 따른 THz 응답 특성을 "
        "조사할 계획입니다."
    )
    doc.add_paragraph(
        "여기에 제시된 방법론\u2014물리 기반 EMA 분석과 데이터 기반 ML 예측의 "
        "결합\u2014은 온도 의존성 미세구조 모니터링이 필요한 다른 다공성 기능 "
        "재료에 적용 가능한 템플릿을 제공하며, 여기에는 세라믹 전해질, "
        "연료전지의 가스 확산층, 여과 멤브레인이 포함됩니다."
    )

    # Graphical abstract
    doc.add_heading("그래픽 초록 (Graphical Abstract)", level=1)
    doc.add_paragraph(
        "그래픽 초록은 다음을 나타냅니다: (좌) 온도 제어 터릿 홀더를 갖춘 매칭 레퍼런스 "
        "THz-TDS 측정 구성, (중앙) 온도 의존 굴절률 감소 및 EMA 기공률 증가, "
        "(우) Lasso ML 예측 산점도 (R² = 0.955). 투고용 별도 고해상도 이미지 파일 "
        "(최소 531 × 1328 픽셀)이 제공됩니다."
    )
    add_figure(doc, "fig03_n_vs_temp_ema",
               "그래픽 초록 (cf. Fig. 4): THz-TDS로 측정한 PE 분리막의 온도 의��� 굴절률 "
               "및 기공률, ML 기반 온도 예측. "
               "(Graphical Abstract: Temperature-dependent refractive index and porosity of PE "
               "separator measured by THz-TDS, with ML-based temperature prediction.)",
               width=5.0)

    # ════════════════════════════════════════════
    # APPENDIX A: Feature Definitions
    # ════════════════════════════════════════════
    doc.add_heading("부록 A. 머신러닝 특성 정의 (Appendix A. Feature definitions for machine learning)", level=1)
    doc.add_paragraph(
        "각 원시 THz 시간영역 신호에서 총 37개의 특성을 자동으로 추출하였습니다. "
        "특성은 5개 그룹으로 분류됩니다. 레퍼런스 상�� 특성의 경우, "
        "동일 온도���서의 매칭 레퍼런스 신호를 사용하였습니다."
    )

    doc.add_paragraph(
        "Fig. A1은 20 \u00b0C에서 측정된 대표적인 시간영역, 엔벨로프, "
        "주파수영역 신호에서의 주요 특성 추출 위치를 나타냅니다."
    )
    add_figure(doc, "fig_feature_illustration",
               "Fig. A1. THz 신호로부터의 특성 추출 주석 그림. "
               "(a) 시간영역 특성: 피크 진폭, 시간 지연 (dt_pos), 상승 시간. "
               "(b) 엔벨로프 특성: 힐버트 엔벨로프의 피크 진폭 및 FWHM. "
               "(c) 주파수영역 특성: 대역 평균 스펙트럼 파워 (저/중/고) 및 스펙트럼 중심. "
               "(Annotated illustration of feature extraction from THz signals.)",
               width=5.5)

    doc.add_heading("A.1. 피크 특성 (Peak characteristics, 6개)", level=2)
    add_table_from_data(doc,
        ["#", "Feature", "Definition"],
        [
            ["1", "peak_pos_amp", "Maximum (positive peak) amplitude of sample signal"],
            ["2", "peak_neg_amp", "Minimum (negative peak) amplitude of sample signal"],
            ["3", "peak_pos_time", "Time position of positive peak (ps)"],
            ["4", "peak_neg_time", "Time position of negative peak (ps)"],
            ["5", "p2p_amp", "Peak-to-peak amplitude: peak_pos_amp \u2212 peak_neg_amp"],
            ["6", "p2p_time", "Time interval between positive and negative peaks (ps)"],
        ],
        "Table A1. 피크 특성 (Peak characteristic features)."
    )

    doc.add_heading("A.2. 펄스 형상 및 통계 (Pulse shape and statistics, 11개)", level=2)
    add_table_from_data(doc,
        ["#", "Feature", "Definition"],
        [
            ["7", "pulse_rms", "RMS amplitude within \u00b13 ps of main peak"],
            ["8", "pulse_std", "Standard deviation within \u00b13 ps window"],
            ["9", "pulse_skewness", "Skewness of pulse waveform"],
            ["10", "pulse_kurtosis", "Kurtosis of pulse waveform"],
            ["11", "pulse_energy", "\u222b s(t)\u00b2 dt over \u00b13 ps window"],
            ["12", "rise_time", "10\u201390% rise time of positive peak (ps)"],
            ["13", "fall_time", "90\u201310% fall time after positive peak (ps)"],
            ["14", "env_peak", "Peak of Hilbert envelope"],
            ["15", "env_fwhm", "Full-width at half-maximum of envelope (ps)"],
            ["16", "env_asymmetry", "(area_right \u2212 area_left) / (area_right + area_left) of envelope"],
            ["17", "zero_crossings", "Number of zero-crossings in \u00b13 ps window"],
        ],
        "Table A2. 펄스 형상 및 통계 특성 (Pulse shape and statistical features)."
    )

    doc.add_heading("A.3. 잡음 특성 (Noise characterization, 2개)", level=2)
    add_table_from_data(doc,
        ["#", "Feature", "Definition"],
        [
            ["18", "noise_rms", "RMS of signal in pre-pulse region (> 5 ps before peak)"],
            ["19", "snr_db", "20 log\u2081\u2080(peak_amp / noise_rms) in dB"],
        ],
        "Table A3. 잡음 특성 (Noise characterization features)."
    )

    doc.add_heading("A.4. 레퍼런스 상대 특성 (Reference-relative features, 6개)", level=2)
    add_table_from_data(doc,
        ["#", "Feature", "Definition"],
        [
            ["20", "dt_pos", "Time delay of sample positive peak vs. reference (fs)"],
            ["21", "amp_ratio_pos", "Positive peak amplitude ratio: sample/reference"],
            ["22", "amp_ratio_neg", "Negative peak amplitude ratio: sample/reference"],
            ["23", "p2p_ratio", "Peak-to-peak ratio: sample/reference"],
            ["24", "delta_rms", "RMS of (sample \u2212 reference) difference signal"],
            ["25", "delta_max", "Maximum absolute value of difference signal"],
        ],
        "Table A4. 레퍼런스 상대 특성 (Reference-relative features)."
    )

    doc.add_heading("A.5. 주파수영역 특성 (Frequency-domain features, 12개)", level=2)
    doc.add_paragraph(
        "스펙트럼 특성은 해닝 윈도우 적용 신호의 FFT로부터 계산하였습니다. "
        "3개 주파수 대역을 정의하였습니다: 저 (0.3\u20130.8 THz), 중 (0.8\u20131.5 THz), "
        "고 (1.5\u20132.5 THz)."
    )
    add_table_from_data(doc,
        ["#", "Feature", "Definition"],
        [
            ["26", "spec_mean_low", "Mean spectral amplitude in 0.3\u20130.8 THz"],
            ["27", "spec_std_low", "Std. dev. of spectral amplitude in 0.3\u20130.8 THz"],
            ["28", "phase_slope_low", "Linear phase slope (rad/Hz) in 0.3\u20130.8 THz"],
            ["29", "spec_mean_mid", "Mean spectral amplitude in 0.8\u20131.5 THz"],
            ["30", "spec_std_mid", "Std. dev. of spectral amplitude in 0.8\u20131.5 THz"],
            ["31", "phase_slope_mid", "Linear phase slope (rad/Hz) in 0.8\u20131.5 THz"],
            ["32", "spec_mean_high", "Mean spectral amplitude in 1.5\u20132.5 THz"],
            ["33", "spec_std_high", "Std. dev. of spectral amplitude in 1.5\u20132.5 THz"],
            ["34", "phase_slope_high", "Linear phase slope (rad/Hz) in 1.5\u20132.5 THz"],
            ["35", "spec_centroid", "\u03a3(f \u00d7 A(f)) / \u03a3A(f) over 0.3\u20132.5 THz (THz)"],
            ["36", "spec_bandwidth", "Spectral standard deviation around centroid (THz)"],
            ["37", "peak_freq", "Frequency of maximum spectral amplitude (THz)"],
        ],
        "Table A5. 주파수영역 특성 (Frequency-domain features)."
    )

    # ════════════════════════════════════════════
    # APPENDIX B: Lasso Prediction Model
    # ════════════════════════════════════════════
    doc.add_heading("부록 B. Lasso 온도 예측 모델 (Appendix B. Lasso temperature prediction model)", level=1)
    doc.add_paragraph(
        "Lasso 목적함수 및 예측 방정식은 2.5절의 Eqs. (6) 및 (7)에 제시되어 있습니다. "
        "Table B1은 선별된 특성의 회귀 계수를 나열합니다."
    )

    add_table_from_data(doc,
        ["Feature", "\u03b2 (standardized)", "Physical interpretation"],
        [
            ["phase_slope_low",  "+14.53", "Low-freq phase gradient \u221d group refractive index"],
            ["spec_mean_low",    "+11.00", "0.3\u20130.8 THz spectral power \u221d transmission"],
            ["spec_centroid",     "+4.01", "Spectral center of mass \u2192 frequency shift"],
            ["amp_ratio_neg",     "\u22123.43", "Negative peak transmission ratio"],
            ["delta_rms",         "\u22122.34", "Sample\u2013reference difference signal RMS"],
            ["phase_slope_mid",   "+2.30", "Mid-freq phase gradient"],
            ["snr_db",            "+1.74", "Signal-to-noise ratio"],
            ["spec_std_mid",      "+1.41", "Mid-freq spectral variability"],
            ["amp_ratio_pos",     "\u22120.72", "Positive peak transmission ratio"],
            ["fall_time",         "\u22120.68", "Pulse fall time"],
            ["env_fwhm",          "+0.60", "Envelope width"],
            ["rise_time",         "\u22120.53", "Pulse rise time"],
            ["spec_std_high",     "\u22120.31", "High-freq spectral variability"],
            ["dt_pos",             "0.00", "Time delay (eliminated by Lasso)"],
            ["env_asymmetry",      "0.00", "Envelope asymmetry (eliminated by Lasso)"],
        ],
        "Table B1. VIF 선별 15개 특성의 Lasso 회귀 계수 "
        "(\u03b1 = 0.1, 표준화 척도). 절편 = 65.0 \u00b0C. "
        "(Lasso regression coefficients for the 15 VIF-selected features.)"
    )

    doc.add_paragraph(
        "지배적 특성 (|β| > 4)은 phase_slope_low, spec_mean_low, spec_centroid이며, "
        "이들 모두 다공성 PE 매트릭스의 유효 광학 경로 길이 및 유전 응답과 관련된 "
        "주파수영역 정보를 부호화합니다. 2개 특성 (dt_pos, env_asymmetry)은 "
        "L1 페널티에 의해 0으로 수렴되었으며, 다른 선별 특성과의 중복을 나타냅니다."
    )

    # ════════════════════════════════════════════
    # DECLARATIONS
    # ════════════════════════════════════════════
    doc.add_heading("이해충돌 선언 (Declaration of competing interest)", level=1)
    doc.add_paragraph(
        "저자들은 본 논문에 보고된 연구에 영향을 미칠 수 있는 것으로 보이는 "
        "알려진 경쟁적 재정적 이해관계 또는 개인적 관계가 없음을 선언합니다."
    )

    doc.add_heading("감사의 글 (Acknowledgments)", level=1)
    doc.add_paragraph("[추가 예정]")

    doc.add_heading("데이터 가용성 (Data availability)", level=1)
    doc.add_paragraph(
        "데이터는 요청 시 제공됩니다."
    )

    doc.add_heading("보충 자료 (Supplementary material)", level=1)
    doc.add_paragraph(
        "보충 그림 (위상 스펙트럼, SNR/동적 범위, ANOVA 주파수 스캔, "
        "시료군 비교, EMA 모델 비교, 드리프트 검증) 및 추가 CSV 데이터 "
        "테이블은 온라인 보충 자료에서 이용 가능합니다."
    )

    # ════════════════════════════════════════════
    # REFERENCES
    # ════════════════════════════════════════════
    doc.add_heading("참고문헌 (References)", level=1)
    refs = [
        "S.C. Mun, J.H. Won, Crystals 11(9) (2021) 1013.",
        "P. Arora, Z. Zhang, Chem. Rev. 104(10) (2004) 4419\u20134462.",
        "Q. Wang, et al., J. Power Sources 208 (2012) 210\u2013224.",
        "X. Feng, et al., Energy Storage Mater. 10 (2018) 246\u2013267.",
        "S.S. Zhang, J. Power Sources 164(1) (2007) 351\u2013364.",
        "C.T. Love, J. Power Sources 196(5) (2011) 2905\u20132912.",
        "P. Bawuah, et al., J. Infrared Milli. Terahz. Waves 41 (2020) 450\u2013469.",
        "J.A. Zeitler, et al., J. Pharm. Sci. 96(2) (2007) 330\u2013340.",
        "M. Naftaly, R.E. Miles, Proc. IEEE 95(8) (2007) 1658\u20131665.",
        "M. Anuschek, P. Bawuah, J.A. Zeitler, Int. J. Pharm. X 3 (2021) 100079.",
        "C. Jansen, et al., Appl. Opt. 49(19) (2010) E48\u2013E57.",
        "S. Wietzke, et al., J. Mol. Struct. 1006(1\u20133) (2011) 41\u201351.",
        "R.H. Boyd, Polymer 26(8) (1985) 1123\u20131133.",
        "L. Duvillaret, F. Garet, J.-L. Coutaz, IEEE J. Sel. Top. Quantum Electron. 2(3) (1996) 739\u2013746.",
        "W. Withayachumnankul, M. Naftaly, J. Infrared Milli. Terahz. Waves 35 (2014) 610\u2013637.",
        "T. Dorney, R. Baraniuk, D. Mittleman, J. Opt. Soc. Am. A 18(7) (2001) 1562\u20131571.",
        "D.A.G. Bruggeman, Ann. Phys. 416(7) (1935) 636\u2013664.",
        "J. Cannarella, C.B. Arnold, J. Power Sources 226 (2013) 149\u2013155.",
        "R. Tibshirani, J. R. Stat. Soc. B 58(1) (1996) 267\u2013288.",
        "P.U. Jepsen, D.G. Cooke, M. Koch, Laser Photon. Rev. 5(1) (2011) 124\u2013166.",
        "Y. Benjamini, Y. Hochberg, J. R. Stat. Soc. B 57(1) (1995) 289\u2013300.",
        "B. Efron, R.J. Tibshirani, An Introduction to the Bootstrap, Chapman & Hall, 1993.",
        "W. Nsengiyumva, et al., Opt. Mater. 123 (2022) 111837.",
        "M. Naftaly, R. Dudley, J. Appl. Phys. 109(4) (2011) 043505.",
        "T. Hastie, R. Tibshirani, J. Friedman, The Elements of Statistical Learning, 2nd ed., Springer, 2009.",
        "S. Sommer, et al., J. Infrared Milli. Terahz. Waves 37 (2016) 189\u2013197.",
        "B. Xiong, et al., J. Membr. Sci. 545 (2018) 213\u2013220.",
        "H. Park, J.-H. Son, Sensors 21(4) (2021) 1186.",
        "M. Koumans, et al., Sci. Rep. 14 (2024) 7034.",
    ]
    for i, r in enumerate(refs):
        doc.add_paragraph(f"[{i+1}] {r}")

    # Save
    out_path = OUT_DIR / "PE40_THz_TDS_paper_SNA_KO.docx"
    doc.save(str(out_path))
    print(f"Paper saved: {out_path}")
    print(f"  Sections: 제목, 주요 성과, 초록, 서론, 재료 및 방법, 결과, 고찰, 결론")
    print(f"  Figures: 8")
    print(f"  Tables: 1")
    return out_path


if __name__ == "__main__":
    generate()
