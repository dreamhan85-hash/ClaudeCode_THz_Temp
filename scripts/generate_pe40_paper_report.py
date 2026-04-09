"""PE40 THz-TDS 온도 의존성 실험 결과 보고서 (논문용).

paper_260406 분석 결과 기반. 9개 Figure, EMA 기공률, 샘플별 분석, 레퍼런스 포함.
"""
from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FIG_DIR = Path(__file__).parent.parent / "figures" / "paper_260406"
CSV_DIR = Path(__file__).parent.parent / "results" / "paper_260406"
OUT = Path(__file__).parent.parent / "PE40_THz_TDS_실험결과보고서.docx"


def remove_borders(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            b = tcPr.makeelement(qn("w:tcBorders"), {})
            for e in ("top", "left", "bottom", "right"):
                b.append(b.makeelement(qn(f"w:{e}"), {qn("w:val"): "none", qn("w:sz"): "0"}))
            tcPr.append(b)


def fig_pair(doc, left, right, cap_l, cap_r, w=6.5):
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, (fp, cap) in enumerate([(left, cap_l), (right, cap_r)]):
        cell = tbl.rows[0].cells[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(fp), width=Cm(w))
        cell2 = tbl.rows[1].cells[ci]
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p2.add_run(cap)
        r.font.size = Pt(8)
        r.font.italic = True
    remove_borders(tbl)


def fig_single(doc, fp, cap, w=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(fp), width=Cm(w))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(cap)
    r.font.size = Pt(8)
    r.font.italic = True


def add_table(doc, headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.font.size = Pt(8)
        r.bold = True
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(7.5)
                r.font.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(7.5)


def bullet(doc, label, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(text)


def main():
    df_opt = pd.read_csv(CSV_DIR / "table02_optical_summary.csv")
    df_n = pd.read_csv(CSV_DIR / "table03_per_sample_n.csv")
    df_ema = pd.read_csv(CSV_DIR / "ema_porosity_results.csv")
    df_corr = pd.read_csv(CSV_DIR / "table04_correlation_analysis.csv")

    # Extract primary frequency from EMA results
    primary_freq = df_ema["freq_thz"].iloc[0]
    pf_label = f"{primary_freq:.2f}"
    # Best 3 frequencies from n columns
    best_freq_labels = [c.replace("n@", "").replace("THz", "")
                        for c in df_opt.columns if c.startswith("n@")]

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # ══════════════════════════════════════════════════════════════
    # 표지
    # ══════════════════════════════════════════════════════════════
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("THz-TDS를 이용한 PE 분리막의\n온도 의존성 미세구조 변화 분석")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0, 0, 0)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("건식 PE 분리막(MS-DPS 20B, 40 μm) — Matched Reference 분석\n"
                     "Bruggeman EMA 기공률 역산 + 전이 온도 결정")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(80, 80, 80)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = date.add_run("측정일: 2024-12-25  |  분석일: 2026-04-07")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 1. 실험 개요
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("1. 실험 개요", level=1)

    doc.add_heading("1.1 연구 배경", level=2)
    doc.add_paragraph(
        "이차전지용 PE 분리막은 기공 구조를 통해 리튬 이온 수송을 적절히 제한하도록 "
        "설계된다. 기존 연구는 130°C 이상의 셧다운(기공 폐쇄) 메커니즘에 집중하였으나, "
        "상온~셧다운 사이 중간 영역에서의 미세구조 변화는 거의 보고되지 않았다. "
        "본 연구에서는 THz-TDS를 이용하여 20~110°C 범위에서 PE 분리막의 "
        "온도 의존성 광학물성을 비파괴적으로 측정하고, Bruggeman 유효 매질 근사(EMA)를 "
        "통해 온도별 기공률 변화를 역산하였다.")

    doc.add_heading("1.2 시료 및 측정 조건", level=2)
    bullet(doc, "시료", "MS-DPS 20B 건식(dry process) PE 분리막 × 2장 (총 두께 40 μm)")
    bullet(doc, "사양 기공률", "44% (제조사 사양)")
    bullet(doc, "DSC 결정화도", "67.8% (Cycle 1, ΔH = 198.65 J/g, Tm = 139.4°C)")
    bullet(doc, "장비", "Menlo Systems ScanControl, Lytera THz-TDS")
    bullet(doc, "온도 범위", "20~110°C (10°C 간격, 10개 온도)")
    bullet(doc, "반복 측정", "각 온도당 5회 (S1~S5, 총 50개 시료 측정)")
    bullet(doc, "레퍼런스", "각 온도별 별도 측정 (Matched Reference)")

    doc.add_heading("1.3 전달함수 및 광학 상수 추출", level=2)
    doc.add_paragraph(
        "각 온도에서 해당 온도의 레퍼런스(공기)와 시료를 비교하는 Matched Reference "
        "방식을 적용하였다. 시간 영역 신호에 윈도우 함수를 적용하지 않고(rectangular), "
        "zero-padding(×2) 후 FFT를 수행하여 주파수 영역으로 변환한다.")

    doc.add_paragraph("측정 전달함수는 다음과 같이 정의된다:")
    eq = doc.add_paragraph()
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eq.add_run("H(ω) = E_sam(ω, T) / E_ref(ω, T)")
    r.font.italic = True
    r.font.size = Pt(10)

    doc.add_paragraph(
        "여기서 E_sam과 E_ref는 각각 시료와 레퍼런스의 FFT 스펙트럼이다. "
        "동일 온도 T에서 측정하므로 챔버 내 공기 경로의 온도 효과가 자동으로 상쇄된다.")

    doc.add_paragraph(
        "시료 두께(40 μm)가 THz 파장(~300 μm at 1 THz)보다 충분히 작으므로 "
        "박막 근사(thin-film approximation)를 적용하였다. "
        "이 경우 Fresnel 반사 손실을 무시하고 이론적 전달함수는:")
    eq2 = doc.add_paragraph()
    eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eq2.add_run("H_theory(ω) ≈ exp[−j(ñ − n_air) · k₀ · d]")
    r.font.italic = True
    r.font.size = Pt(10)

    doc.add_paragraph(
        "여기서 ñ = n − jκ (복소 굴절률), k₀ = ω/c (자유 공간 파수), "
        "d = 40 μm (시료 두께), n_air = 1.0이다.")

    doc.add_paragraph(
        "각 주파수에서 |H_meas(ω) − H_theory(ω)|² 를 최소화하는 "
        "Nelder-Mead 최적화를 수행하여 n(ω)과 κ(ω)를 추출하였다. "
        "인접 주파수의 해를 초기값으로 사용하는 warm-starting 기법을 적용하였다. "
        "주파수 범위: 0.2~2.5 THz.")

    doc.add_paragraph("흡수계수와 유전율은 다음과 같이 계산된다:")
    eqs_alpha = doc.add_paragraph()
    eqs_alpha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eqs_alpha.add_run("α(ω) = 2κω/c  [cm⁻¹]")
    r.font.italic = True
    r.font.size = Pt(10)

    eqs_eps = doc.add_paragraph()
    eqs_eps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eqs_eps.add_run("ε' = n² − κ²,    ε'' = 2nκ")
    r.font.italic = True
    r.font.size = Pt(10)

    doc.add_heading("1.4 Bruggeman 3상 EMA 기공률 역산", level=2)
    doc.add_paragraph(
        "다공성 PE 분리막을 3상 복합체(결정 PE + 비정질 PE + 공기)로 모델링하고, "
        "Bruggeman 유효 매질 근사(EMA)를 적용하여 측정된 유효 굴절률 n_eff(T)로부터 "
        "온도별 기공률 f_air(T)를 역산하였다 [1]. "
        "Bruggeman EMA 자기일관(self-consistent) 방정식:")

    eq_ema = doc.add_paragraph()
    eq_ema.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eq_ema.add_run(
        "f_c · (ε_c − ε_eff)/(ε_c + 2ε_eff) + "
        "f_a · (ε_a − ε_eff)/(ε_a + 2ε_eff) + "
        "f_air · (ε_air − ε_eff)/(ε_air + 2ε_eff) = 0")
    r.font.italic = True
    r.font.size = Pt(9)

    doc.add_paragraph("여기서:")
    params_ema = [
        "ε_c = n_c² = 1.53² = 2.341 (결정 PE의 유전율)",
        "ε_a = n_a² = 1.49² = 2.220 (비정질 PE의 유전율)",
        "ε_air = 1.0 (공기의 유전율)",
        "f_c = (1 − f_air) × X_c, f_a = (1 − f_air) × (1 − X_c) (체적 분율)",
        "X_c = 0.678 (DSC 결정화도, ΔH_f/ΔH_f^100% = 198.65/293.0)",
        "ε_eff = n_eff² (유효 유전율, 비자성 근사 μ_r = 1)",
    ]
    for p in params_ema:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_paragraph(
        "위 방정식에서 n_eff(T)를 측정값으로 대입하고, f_air를 미지수로 하여 "
        "Brent 방법(scipy.optimize.brentq)으로 수치적으로 역산한다.")

    doc.add_heading("1.5 2장 겹침 보정 (Spec Anchoring)", level=2)
    doc.add_paragraph(
        "2장 정전기 밀착에 의한 표면 기공 부분 폐쇄로 측정 n이 단일 시트보다 높게 나타난다. "
        "제조사 사양 기공률(44%)을 기준으로 보정한다:")

    eq_offset = doc.add_paragraph()
    eq_offset.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eq_offset.add_run(
        "n_offset = n_measured(20°C) − n_EMA(f_air = 0.44)\n"
        "       = 1.3415 − 1.2818 = +0.0598")
    r.font.italic = True
    r.font.size = Pt(9.5)

    eq_corr = doc.add_paragraph()
    eq_corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eq_corr.add_run("n_corrected(T) = n_measured(T) − n_offset")
    r.font.italic = True
    r.font.size = Pt(10)

    doc.add_paragraph(
        "보정된 n_corrected(T)를 Bruggeman 역산에 대입하여 f_air(T)를 계산한다. "
        "이 방법은 2장 겹침이 n에 온도 무관한 상수 오프셋을 추가한다고 가정한다.")

    doc.add_heading("1.6 전이 온도 결정 (Two-Regime Fitting)", level=2)
    doc.add_paragraph(
        "n(T) 데이터를 두 개의 선형 구간(regime)으로 분할하는 piecewise linear fitting을 "
        "수행하여 전이 온도 T_onset를 결정하였다:")

    eq_regime = doc.add_paragraph()
    eq_regime.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eq_regime.add_run(
        "n(T) = { a₁·T + b₁,  T ≤ T_break\n"
        "       { a₂·T + b₂,  T > T_break")
    r.font.italic = True
    r.font.size = Pt(9.5)

    doc.add_paragraph(
        "T_break를 30~90°C 범위에서 스캔하며, 양쪽 잔차 제곱합 "
        "Σ(n_i − a₁T_i − b₁)² + Σ(n_j − a₂T_j − b₂)²를 최소화하는 "
        "최적 분할점을 scipy.optimize.minimize_scalar로 결정한다. "
        "두 직선의 교점 T_onset = (b₂ − b₁)/(a₁ − a₂)가 전이 온도이다.")

    doc.add_heading("1.7 열팽창 모델", level=2)
    doc.add_paragraph(
        "기공률의 온도 의존성을 체적 열팽창으로 모델링하였다 [3]:")

    eq_therm = doc.add_paragraph()
    eq_therm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eq_therm.add_run("f_air(T) = 1 − (1 − f_air,0) / (1 + β · (T − T_ref))")
    r.font.italic = True
    r.font.size = Pt(10)

    doc.add_paragraph(
        "여기서 f_air,0은 초기 기공률(T_ref = 20°C), β는 유효 체적 열팽창계수이다. "
        "β는 PE 매트릭스의 열팽창에 의한 기공 확대를 반영하며, "
        "scipy.optimize.curve_fit으로 측정 데이터에 fitting하여 결정한다.")

    # ══════════════════════════════════════════════════════════════
    # 2. 시간 영역 분석
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("2. 시간 영역 분석", level=1)

    doc.add_heading("2.1 온도별 THz 파형", level=2)
    fig_single(doc, FIG_DIR / "fig01_time_domain_zoom.png",
               "Fig. 1. 온도별 THz 파형 (피크 영역 확대, −1.5~2.0 ps). "
               "검정: Ref(T), 색상: S1~S5.", w=15)
    doc.add_paragraph(
        "각 온도에서 해당 온도 레퍼런스(검정)와 5개 시료 신호를 비교하였다. "
        "40 μm 시료에 의한 시간 지연과 진폭 감쇠가 명확히 관찰된다.")

    doc.add_heading("2.2 피크 상세 — 감쇠 및 지연", level=2)
    fig_single(doc, FIG_DIR / "fig02_peak_detail.png",
               "Fig. 2. 전 온도 오버레이. 좌: 양의 피크, 우: 음의 피크. "
               "검정: 20°C Ref. 온도 증가 시 진폭 증가 + 시간 좌측 이동.", w=13)
    doc.add_paragraph(
        "온도 증가에 따라 (1) 시료 펄스의 진폭이 레퍼런스에 가까워지고(투과율 증가), "
        "(2) 피크 위치가 좌측으로 이동하는(굴절률 감소) 경향이 관찰된다.")

    doc.add_heading("2.3 샘플별 시간 특성 vs 온도", level=2)
    fig_single(doc, FIG_DIR / "fig05_time_features_vs_temp.png",
               "Fig. 3. 좌: 샘플별 시간 지연(Δt_avg) vs 온도 (R²=0.773***). "
               "우: P2P ratio vs 온도. 검정: 평균±σ.", w=13)
    doc.add_paragraph(
        "시간 지연은 20°C(~45 fs) → 100°C(~37 fs)로 감소하며 "
        "강한 온도 상관(R² = 0.77, p < 0.001)을 보인다. "
        "P2P 비율은 0.98~1.00 범위로 온도 의존성이 약하다(R² = 0.10, p = 0.37).")

    # ══════════════════════════════════════════════════════════════
    # 3. 광학 상수
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("3. 광학 상수 (Matched Reference)", level=1)

    doc.add_heading("3.1 굴절률 및 흡수계수 스펙트럼", level=2)
    fig_single(doc, FIG_DIR / "fig04_optical_spectra.png",
               "Fig. 4. 좌: n(f), 우: α(f). 색상: 20°C(파랑) → 110°C(빨강). "
               "주파수 범위 0.3~2.0 THz.", w=13)
    doc.add_paragraph(
        "온도별 굴절률 스펙트럼에서 20~30°C가 가장 높고, "
        "40°C 이후 급격히 감소한 뒤 50~90°C에서 안정화된다. "
        f"최적 분석 주파수 3개({', '.join(f'{f} THz' for f in best_freq_labels)})를 "
        "수직 점선으로 표시하였다.")

    doc.add_heading("3.2 n(T) 및 EMA 기공률 (★ 핵심)", level=2)

    # Read actual values from EMA results
    n_20 = df_ema.iloc[0][df_ema.columns[1]]  # n measured at 20°C
    n_100 = df_ema.iloc[8][df_ema.columns[1]]  # n measured at 100°C

    fig_single(doc, FIG_DIR / "fig03_n_vs_temp_ema.png",
               f"Fig. 5. 좌축: n @ {pf_label} THz (샘플별 + Mean±σ). "
               f"우축: EMA 역산 기공률(%).", w=8)
    doc.add_paragraph(
        f"n @ {pf_label} THz는 20°C에서 {n_20:.3f} → 100°C에서 {n_100:.3f}로 감소한다. "
        f"Two-regime piecewise linear fitting 결과 전이 온도 T_onset가 결정되었다.")

    # Table: optical summary — dynamically read columns
    n_cols = [c for c in df_opt.columns if c.startswith("n@")]
    alpha_col = [c for c in df_opt.columns if c.startswith("alpha@")][0]
    eps_col = [c for c in df_opt.columns if c.startswith("eps'@")][0]
    opt_headers = ["온도(°C)", "Δt_avg(fs)"] + n_cols + [alpha_col, eps_col, "f_air(%)"]
    opt_rows = []
    for _, row in df_opt.iterrows():
        r = [f"{int(row['Temperature (C)'])}", str(row["dt_avg (fs)"])]
        for nc in n_cols:
            r.append(str(row[nc]))
        r.append(str(row[alpha_col]))
        r.append(str(row[eps_col]))
        r.append(str(row["f_air (%)"]))
        opt_rows.append(r)
    add_table(doc, opt_headers, opt_rows,
              "Table 1. 온도별 광학 상수 및 EMA 기공률 요약")
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 4. EMA 기공률 분석
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("4. EMA 기공률 분석", level=1)

    fig_single(doc, FIG_DIR / "fig06_porosity_vs_temp.png",
               "Fig. 6. 온도별 유효 기공률 (Bruggeman EMA, 사양 앵커링). "
               "빨간 파선: 열팽창 모델 (β = 2.22×10⁻³/°C).", w=8)

    doc.add_paragraph(
        "제조사 사양 기공률(44%)로 앵커링한 Bruggeman 3상 EMA 역산 결과, "
        "온도 증가에 따라 기공률이 44%(20°C) → 53~55%(100°C)로 "
        "약 10%p 증가한다.")

    doc.add_paragraph("주요 발견:")
    findings = [
        "20~30°C: f_air ≈ 44% (안정, 사양값과 일치)",
        "40°C: f_air ≈ 49% (기공률 증가 시작)",
        "50~90°C: f_air ≈ 51~53% (plateau, 네트워크 연화/기계적 평형)",
        "100°C: f_air ≈ 55% (최대, 셧다운 전 최대 기공 확대)",
    ]
    for f in findings:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_paragraph(
        "열팽창 모델 f_air(T) = 1 − (1−f₀)/(1+β·ΔT)로 fitting한 결과 "
        "[Wei et al., 2013], 초기 기공률 f₀ = 45.8%, "
        "유효 체적 열팽창계수 β = 2.22×10⁻³/°C가 결정되었다. "
        "이는 벌크 PE의 열팽창계수(~2×10⁻⁴/°C)보다 약 10배 높은 값으로, "
        "다공성 구조의 기공 확대 효과가 열팽창을 증폭시킴을 나타낸다.")

    # ══════════════════════════════════════════════════════════════
    # 5. 유전율 분석
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("5. 유전율 분석", level=1)
    fig_single(doc, FIG_DIR / "fig07_dielectric.png",
               "Fig. 7. 좌: ε'(f) 온도별 스펙트럼. 우: ε' @ 1.0 THz 샘플별 온도 추세.",
               w=13)
    doc.add_paragraph(
        "유전율 실수부 ε' = n² − κ²는 유효 분석 주파수에서 1.6~2.0 범위이며, "
        "굴절률과 동일한 온도 의존성을 보인다. "
        "전기화학 관점에서 ε'의 감소는 기공률 증가에 의한 "
        "유효 유전 환경 변화를 반영한다.")

    # ══════════════════════════════════════════════════════════════
    # 6. 샘플별 분석 및 이상 샘플 검출
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6. 샘플별 분석 및 이상 샘플 검출", level=1)

    doc.add_heading("6.1 샘플별 광학 특성 비교", level=2)
    fig_single(doc, FIG_DIR / "fig09_per_sample_overview.png",
               "Fig. 8. 샘플별(S1~S5) 온도 추세. "
               "좌상: n, 우상: α, 좌하: ε', 우하: f_air.", w=14)

    # Per-sample n table
    n_headers = ["온도(°C)", "S1", "S2", "S3", "S4", "S5", "Mean", "σ"]
    n_rows = []
    for _, row in df_n.iterrows():
        n_rows.append([
            f"{int(row['Temperature (C)'])}",
            row["S1"], row["S2"], row["S3"], row["S4"], row["S5"],
            row["Mean"], row["sigma"],
        ])
    add_table(doc, n_headers, n_rows,
              "Table 2. 샘플별 n @ 1.0 THz")
    doc.add_paragraph()

    doc.add_heading("6.2 이상 샘플 검출 결과", level=2)
    outlier_headers = ["샘플", "<n>", "R²", "기울기(/°C)", "<|편차|>", "판정"]
    outlier_rows = [
        ["S1", "1.307", "0.14", "−1.23×10⁻⁴", "0.011", "OK"],
        ["S2", "1.315", "0.28", "−2.41×10⁻⁴", "0.012", "OK"],
        ["S3", "1.321", "0.64**", "−6.70×10⁻⁴", "0.014", "OK"],
        ["S4", "1.306", "0.45*", "−5.58×10⁻⁴", "0.007", "OK (최소 편차)"],
        ["S5", "1.286", "0.65**", "−9.00×10⁻⁴", "0.023", "OK (최대 편차)"],
    ]
    add_table(doc, outlier_headers, outlier_rows,
              "Table 3. 샘플별 이상 검출 분석 (그룹 평균 σ = 0.016)")
    doc.add_paragraph()

    doc.add_paragraph(
        "모든 샘플이 1.5σ 이내로 통계적 이상치는 없다. "
        "다만 S5가 전 온도에서 일관되게 낮은 n(평균 편차 0.023)을 보이며, "
        "S1은 온도 상관성이 거의 없다(R² = 0.14). "
        "S3, S4, S5가 가장 명확한 온도 추세(R² > 0.45)를 보인다.")

    doc.add_paragraph(
        "현재 모든 샘플을 포함하여 분석하였으며, "
        "필요 시 S5(편차 최대) 또는 S1(상관 최소)을 "
        "제외한 분석도 가능하다.")

    # ══════════════════════════════════════════════════════════════
    # 7. 온도 상관 분석
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("7. 온도 상관 분석 (시료 고유 특성)", level=1)

    fig_single(doc, FIG_DIR / "fig08_correlation_summary.png",
               "Fig. 9. 좌: 시료 특성 온도별 정규화 히트맵. "
               "우: R² 바 차트 (파랑: p<0.05, 회색: 비유의).", w=14)

    doc.add_paragraph(
        "시료 고유 특성만을 대상으로 온도 상관 분석을 수행하였다. "
        "레퍼런스 자체의 절대 진폭 변화(amp_pos, p2p 등)는 "
        "시스템 특성이므로 제외하였다.")

    corr_headers = ["순위", "파라미터", "도메인", "R²", "p-value", "유의성"]
    corr_rows = []
    for _, row in df_corr.iterrows():
        if int(row["Rank"]) > 10:
            break
        corr_rows.append([
            str(int(row["Rank"])), row["Feature"], row["Domain"],
            row["R2"], row["p-value"], row["Significance"],
        ])
    add_table(doc, corr_headers, corr_rows,
              "Table 4. 시료 특성 온도 상관 순위 — Top 10")
    doc.add_paragraph()

    doc.add_paragraph(
        "Rise time(R² = 0.95***)과 스펙트럼 중심(R² = 0.93***)이 "
        "가장 높은 온도 상관성을 보이며, 시간 지연(dt_avg: R² = 0.77***), "
        "엔벨로프 특성(env_fwhm: R² = 0.79***, env_asymmetry: R² = 0.78***)이 "
        "그 다음이다. 진폭 기반 특성(P2P ratio, amp ratio)은 "
        "유의한 온도 상관이 없다(p > 0.05).")

    # ══════════════════════════════════════════════════════════════
    # 8. 고찰
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("8. 고찰", level=1)

    doc.add_heading("8.1 기공 구조 변화 메커니즘", level=2)
    doc.add_paragraph(
        "온도 증가에 따른 굴절률 감소(→ 기공률 증가)는 두 가지 구간으로 구분된다:")
    items_mech = [
        "Regime 1 (20~59°C): PE 매트릭스의 열팽창에 의한 기공 확대. "
        "건식 분리막의 슬릿형 기공이 연신 방향으로 확장되며 기공률이 급격히 증가한다. "
        "이 온도 범위는 급속 충전 시 배터리 내부 온도(40~60°C)에 해당하며, "
        "설계 의도보다 빠른 시점에 기공 구조 변화가 시작됨을 시사한다.",
        "Regime 2 (59~110°C): 기공률 plateau. PE 네트워크의 연화로 인한 "
        "기계적 평형 상태 도달. 기공 확대와 네트워크 이완이 균형을 이루어 "
        "유효 굴절률이 거의 일정하게 유지된다.",
    ]
    for item in items_mech:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("8.2 2장 겹침 효과", level=2)
    doc.add_paragraph(
        "PE20 2장을 정전기로 밀착시킨 PE40 시료에서 측정된 n(1.342)은 "
        "단일 시트 기대값(n ≈ 1.282, f_air = 44%)보다 높다. "
        "이는 정전기 밀착에 의한 표면 기공의 부분적 폐쇄에 기인하며, "
        "n_offset = +0.060으로 보정하였다. "
        "유사한 현상이 분리막 압축 연구에서도 보고되었다 [Xu et al., 2023].")

    doc.add_heading("8.3 시간 영역 vs 주파수 영역", level=2)
    doc.add_paragraph(
        "시간 영역 직접 분석(rise time: R² = 0.95, dt_avg: R² = 0.77)이 "
        "주파수 영역 광학 상수 추출(n vs T: R² ≈ 0.14~0.65 샘플 의존)보다 "
        "높은 온도 민감도를 보인다. 이는 40 μm 박막에서 Nelder-Mead 최적화의 "
        "수렴 불안정성을 회피하면서도 온도 변화를 민감하게 검출할 수 있음을 시사하며, "
        "실시간 모니터링 응용에 유리하다.")

    # ══════════════════════════════════════════════════════════════
    # 9. 결론
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("9. 결론", level=1)

    conclusions = [
        "THz-TDS로 건식 PE 분리막(MS-DPS 20B, 40 μm)의 20~110°C "
        "온도 의존성 광학물성을 Matched Reference 방식으로 측정하였다.",
        f"n @ {pf_label} THz: {n_20:.3f}(20°C) → {n_100:.3f}(100°C). "
        "Bruggeman 3상 EMA로 기공률 역산 시 44%(20°C) → 55~59%(100°C).",
        "Two-regime fitting으로 전이 온도 T_onset ≈ 59°C 결정. "
        "40°C부터 기공 구조 변화 시작, 59°C 이후 plateau.",
        "열팽창 모델 fitting 결과 유효 체적 열팽창계수 β = 2.22×10⁻³/°C "
        "(벌크 PE 대비 ~10배, 다공성 구조 증폭 효과).",
        "급속 충전 온도(40~60°C)에서 이미 구조 변화가 시작됨 — "
        "비균일 이온 플럭스 위험 시사.",
        "시간 영역 직접 분석(rise time, dt_avg)이 광학 상수 추출보다 "
        "높은 온도 민감도를 보여, 실시간 모니터링에 적합.",
        "5개 샘플 모두 1.5σ 이내로 통계적 이상치 없음. "
        "S5(편차 최대), S4(최소 편차)의 차이는 시료 간 변동 범위 내.",
    ]
    for c in conclusions:
        doc.add_paragraph(c, style="List Bullet")

    # ══════════════════════════════════════════════════════════════
    # 10. 분석 조건
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("10. 분석 조건", level=1)
    params = [
        ("시료 두께", "40 μm (MS-DPS 20B × 2장, 정전기 밀착)"),
        ("분석 방법", "Matched Reference — H(ω) = E_sam(T)/E_ref(T)"),
        ("주파수 범위", "0.2~2.5 THz"),
        ("윈도우 함수", "없음 (rectangular)"),
        ("Zero-padding", "×2"),
        ("모델", "Thin-film approximation, Nelder-Mead 최적화"),
        ("EMA", "Bruggeman 3상 (결정PE 1.53 + 비정질PE 1.49 + 공기 1.0)"),
        ("기공률 앵커링", "20°C에서 제조사 사양 44%로 고정, n_offset = +0.060"),
        ("전이 온도", "Piecewise linear fitting (two-regime)"),
        ("열팽창 모델", "f_air(T) = 1 − (1−f₀)/(1+β·ΔT)"),
    ]
    for label, val in params:
        bullet(doc, label, val)

    # ══════════════════════════════════════════════════════════════
    # 11. 참고문헌
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("11. 참고문헌", level=1)
    refs_list = [
        "[1] Bawuah et al., \"Terahertz-Based Porosity Measurement of Pharmaceutical "
        "Tablets: a Tutorial,\" J. Infrared Milli. Terahz. Waves, vol. 41, "
        "pp. 450-469, 2020. doi:10.1007/s10762-019-00659-0",
        "[2] Bawuah et al., \"THz-TDS for powder compact porosity and pore shape "
        "measurements: An error analysis of the anisotropic Bruggeman model,\" "
        "Int. J. Pharm. X, vol. 3, p. 100078, 2021. doi:10.1016/j.ijpx.2021.100078",
        "[3] Wei et al., \"Effect of Porosity on the Thermal Expansion Coefficient "
        "of Porous Materials,\" Proc. Eng. Mech. Inst. Conf., 2013. "
        "doi:10.1061/9780784412992.220",
        "[4] Deimede et al., \"Manufacturing Processes of Microporous Polyolefin "
        "Separators for Lithium-Ion Batteries,\" Crystals, vol. 11(9), p. 1013, "
        "2021. doi:10.3390/cryst11091013",
        "[5] Xu et al., \"Effect of a compressed separator on the electrochemical "
        "performance of Li-ion battery,\" J. Power Sources, 2023. "
        "doi:10.1016/j.jpowsour.2023.232900",
    ]
    for ref in refs_list:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(8.5)

    doc.save(str(OUT))
    print(f"Report saved: {OUT}")


if __name__ == "__main__":
    main()
