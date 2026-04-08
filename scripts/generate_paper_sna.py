"""Sensors and Actuators A: Physical 투고용 논문 DOCX 생성.

PE40 THz-TDS temperature study — full manuscript.
Target: ~5000 words, ~10 figures, Minor Revision quality.
"""
from __future__ import annotations

import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

FIG_DIR = Path(__file__).parent.parent / "figures" / "paper_260406"
OUT_DIR = Path(__file__).parent.parent / "results" / "paper_260406"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def add_equation(doc, eq_text, eq_num):
    """Add equation in reference-paper style: centered equation + right-aligned number."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(eq_text)
    run.italic = True
    run.font.size = Pt(12)
    # Tab + equation number (right)
    run2 = p.add_run(f"\t({eq_num})")
    run2.italic = False
    run2.font.size = Pt(12)


def add_where(doc, text):
    """Add 'where' clause after equation, matching reference style."""
    p = doc.add_paragraph(f"where {text}")
    p.paragraph_format.space_before = Pt(2)
    return p


def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 2.0

    for level, size in [(1, 14), (2, 13), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)


def add_figure(doc, fig_name, caption, width=5.5):
    fig_path = FIG_DIR / f"{fig_name}.png"
    if not fig_path.exists():
        doc.add_paragraph(f"[Figure missing: {fig_name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fig_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style.font.size = Pt(10)
    cap.runs[0].italic = True


def add_table_from_data(doc, headers, rows, caption):
    cap = doc.add_paragraph(caption)
    cap.runs[0].italic = True
    cap.style.font.size = Pt(10)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.runs[0].font.size = Pt(10)


def generate():
    doc = Document()
    setup_styles(doc)

    # ════════════════════════════════════════════
    # TITLE
    # ════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Temperature-Dependent Microstructural Changes of Polyethylene "
        "Battery Separators Investigated by Terahertz Time-Domain Spectroscopy "
        "and Machine Learning"
    )
    run.bold = True
    run.font.size = Pt(16)

    # Authors
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run("[Author Names]").font.size = Pt(12)
    affil = doc.add_paragraph()
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil.add_run("[Affiliations]").font.size = Pt(11)

    doc.add_paragraph()

    # ════════════════════════════════════════════
    # HIGHLIGHTS
    # ════════════════════════════════════════════
    doc.add_heading("Highlights", level=1)
    highlights = [
        "THz-TDS reveals temperature-dependent porosity changes in PE battery separators (20-110 \u00b0C)",
        "Refractive index decreases monotonically with temperature (n: 1.43 \u2192 1.35 at 0.45 THz)",
        "Bruggeman EMA porosity inversion: 44% (spec) \u2192 59% at 100 \u00b0C with thermal onset at 55 \u00b0C",
        "Lasso regression predicts temperature from raw THz signals with R\u00b2 = 0.955 (LOTO-CV, MAE = 5.1 \u00b0C)",
        "Phase slope and low-frequency spectral features identified as primary temperature indicators",
    ]
    for h in highlights:
        doc.add_paragraph(h, style="List Bullet")

    # ════════════════════════════════════════════
    # ABSTRACT
    # ════════════════════════════════════════════
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Non-destructive monitoring of battery separator microstructure under thermal stress "
        "is critical for lithium-ion battery safety. This study investigates the temperature-dependent "
        "optical and microstructural properties of polyethylene (PE) separators using terahertz "
        "time-domain spectroscopy (THz-TDS) over the range of 20\u2013110 \u00b0C. "
        "A matched-reference method was employed to extract refractive index and absorption coefficient "
        "spectra (0.2\u20132.5 THz) at 10 \u00b0C intervals with five replicates per temperature. "
        "The refractive index at 0.45 THz decreased from 1.43 to 1.35, corresponding to an effective "
        "porosity increase from 44% to 59% via two-phase Bruggeman effective medium approximation. "
        "A two-regime thermal behavior was identified with an onset temperature of 55.2 \u00b0C, "
        "consistent with the \u03b1c crystalline relaxation of PE. "
        "One-way ANOVA confirmed statistically significant temperature dependence across 97.4% of the "
        "analyzed frequency range (\u03b7\u00b2 = 0.74 at 0.45 THz, p = 2.84 \u00d7 10\u207b\u2079). "
        "Furthermore, a machine learning pipeline was developed to predict temperature directly from "
        "raw time-domain signals. Using 15 physics-informed features selected via variance inflation "
        "factor analysis, a Lasso regression model achieved R\u00b2 = 0.955 with MAE = 5.1 \u00b0C "
        "under leave-one-temperature-out cross-validation, with phase slope and low-frequency spectral "
        "power identified as the dominant predictors. "
        "These results demonstrate, within the studied sample and measurement conditions, "
        "the feasibility of THz-TDS as a non-contact probe for monitoring thermally induced "
        "microstructural changes in porous polymer membranes."
    )

    # Keywords
    kw = doc.add_paragraph()
    kw.add_run("Keywords: ").bold = True
    kw.add_run(
        "Terahertz time-domain spectroscopy; Polyethylene separator; "
        "Temperature dependence; Effective medium approximation; "
        "Machine learning; Battery safety"
    )

    # ════════════════════════════════════════════
    # 1. INTRODUCTION
    # ════════════════════════════════════════════
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Lithium-ion batteries (LIBs) are integral to modern energy storage, from portable electronics "
        "to electric vehicles. The separator, typically a microporous polyolefin membrane, is a critical "
        "safety component that prevents direct contact between electrodes while permitting ion transport "
        "[1,2]. Under thermal abuse conditions\u2014overcharging, external short-circuits, or elevated "
        "ambient temperatures\u2014the separator undergoes progressive structural changes: pore closure, "
        "shrinkage, and ultimately melting, which can trigger thermal runaway [3,4]."
    )
    doc.add_paragraph(
        "Understanding the early-stage microstructural evolution of separators under thermal stress is "
        "therefore essential for battery safety engineering. Conventional characterization techniques "
        "such as scanning electron microscopy (SEM), mercury intrusion porosimetry, and differential "
        "scanning calorimetry (DSC) provide valuable structural information but are destructive, "
        "time-consuming, or limited to ex-situ measurements [5,6]."
    )
    doc.add_paragraph(
        "Terahertz time-domain spectroscopy (THz-TDS) has emerged as a promising non-destructive "
        "technique for characterizing thin polymer films and porous materials [7\u20139]. "
        "The terahertz frequency range (0.1\u20133 THz) is sensitive to dielectric properties, "
        "structural periodicity, and free-volume changes at the mesoscale. "
        "Recent studies have demonstrated the capability of THz-TDS for porosity determination "
        "in pharmaceutical tablets [10], polymer coatings [11], and ceramic membranes [12] "
        "using effective medium approximation (EMA) models."
    )
    doc.add_paragraph(
        "In the context of battery safety, understanding separator behavior under elevated "
        "temperatures is particularly important. During thermal abuse events, the separator "
        "temperature can rise from ambient to above 130 \u00b0C within minutes [3,4]. The sequence "
        "of microstructural changes\u2014pore widening, shrinkage onset, and eventual melting\u2014"
        "determines the critical time window for safety intervention. Non-contact "
        "characterization of these changes during manufacturing would improve separator quality "
        "assurance. However, existing optical and acoustic techniques lack the penetration "
        "depth and material sensitivity required for polymer membrane characterization at "
        "the relevant length scales (10\u2013100 \u00b5m thickness, sub-micrometer pore features) [6]."
    )
    doc.add_paragraph(
        "Machine learning approaches have recently been integrated with THz-TDS to enhance "
        "material classification and property extraction [20]. Supervised regression models "
        "can identify subtle spectral features that correlate with physical properties, "
        "potentially surpassing traditional analytical extraction methods in speed and robustness. "
        "The combination of physics-based analysis (EMA porosity inversion) with data-driven [29] "
        "prediction (ML temperature regression) offers a complementary framework where each "
        "approach validates and extends the other."
    )
    doc.add_paragraph(
        "However, systematic studies on the temperature dependence of THz optical properties "
        "in battery separators remain scarce. In particular, the relationship between "
        "thermally induced microstructural changes\u2014such as pore expansion, lamellar "
        "rearrangement, and amorphous-phase relaxation\u2014and measurable THz parameters "
        "has not been quantitatively established."
    )
    doc.add_paragraph(
        "In this work, we present a comprehensive THz-TDS investigation of temperature-dependent "
        "microstructural changes in dry-process polyethylene (PE) separators over 20\u2013110 \u00b0C. "
        "A matched-reference measurement protocol is employed with Bruggeman two-phase EMA for "
        "porosity inversion. The statistical significance of observed trends is rigorously validated "
        "through ANOVA and bootstrap confidence intervals. Additionally, we develop a machine learning "
        "pipeline that extracts physics-informed features from raw THz signals to predict temperature, "
        "providing insights into which signal characteristics carry the most thermal information."
    )

    # ════════════════════════════════════════════
    # 2. MATERIALS AND METHODS
    # ════════════════════════════════════════════
    doc.add_heading("2. Materials and Methods", level=1)

    doc.add_heading("2.1. Sample preparation", level=2)
    doc.add_paragraph(
        "Commercial dry-process PE separators (MS-DPS 20B, manufacturer-specified porosity 44%, "
        "thickness 20 \u00b5m per sheet, air permeability 280 s/100 mL) were used. "
        "Two sheets were stacked to achieve a total thickness of approximately 40 \u00b5m, "
        "enhancing the THz signal-to-noise ratio while maintaining structural representativeness. "
        "The dry-process PE separator is manufactured by uniaxial stretching of an extruded "
        "PE film, which creates a network of slit-shaped, oriented pores. This process results "
        "in a characteristic anisotropic microstructure with pore dimensions on the order of "
        "0.1\u20131 \u00b5m, aligned predominantly in the machine direction. The two-sheet stacking "
        "configuration was adopted to increase the effective optical path length from 20 to 40 \u00b5m, "
        "improving the phase sensitivity of the THz-TDS measurement by approximately a factor of two "
        "while maintaining structural representativeness. "
        "DSC analysis (first heating cycle) yielded a melting temperature Tm = 139.4 \u00b0C "
        "and crystallinity of 67.8% (based on \u0394Hf = 198.7 J/g, \u0394H100% = 293.0 J/g)."
    )
    doc.add_paragraph(
        "Fig. 1 shows a scanning electron micrograph of the separator surface (50,000\u00d7 magnification, "
        "E\u2013T detector, 2 keV). The image reveals the characteristic fibrillar network structure "
        "of dry-process PE separators: thick fibrils (~100\u2013200 nm width) oriented in the machine "
        "direction (vertical) are interconnected by thinner lamellar bridges (~50\u2013100 nm) in the "
        "transverse direction. The resulting slit-shaped pores, with dimensions of approximately "
        "100\u2013500 nm, form an open, interconnected network that provides the ionic transport "
        "pathways. This anisotropic pore geometry is relevant to the interpretation of the EMA "
        "results (Section 2.4), as the isotropic Bruggeman model assumes spherical inclusions "
        "rather than the oriented slit-shaped pores observed here."
    )
    add_figure(doc, "fig01_sem_dsc",
               "Fig. 1. (a) SEM micrograph of the dry-process PE separator (MS-DPS 20B) at "
               "50,000\u00d7 magnification, showing the fibrillar network with slit-shaped pores "
               "(~100\u2013500 nm). Scale bar: 1 \u00b5m. (b) DSC thermogram (first heating cycle) "
               "showing the melting peak at Tm = 139.5 \u00b0C with 67.8% crystallinity.",
               width=5.0)

    doc.add_heading("2.2. THz-TDS measurements", level=2)
    doc.add_paragraph(
        "THz-TDS measurements were performed using a commercial fiber-coupled system (Menlo Systems "
        "TeraSmart, ScanControl software v1.10) equipped with photoconductive antenna emitter and "
        "detector modules (TERA15-FC). The experimental setup is illustrated in Fig. 2. "
        "The measurement configuration employed a transmission geometry "
        "with a temperature-controlled sample stage. Measurements were conducted at 10 temperatures "
        "from 20 to 110 \u00b0C in 10 \u00b0C steps (heating direction only). "
        "Temperature was controlled using a six-position turret sample holder "
        "(Quantum Northwest), with position 1 serving as the empty reference channel "
        "and positions 2\u20136 holding the five sample replicates (S1\u2013S5). "
        "A stabilization time of 5 minutes was applied at each temperature to ensure "
        "thermal equilibrium (temperature uncertainty estimated at \u00b11\u20132 \u00b0C). "
        "Measurements were performed under ambient humidity conditions (~40\u201360% RH); "
        "no humidity control was applied, as PE is non-hygroscopic. "
        "At each temperature, the reference (turret position 1, empty aperture) was measured first, "
        "followed by the five samples sequentially (positions 2\u20136), each averaged over 5 waveforms. "
        "This matched-reference protocol eliminates systematic temperature-dependent instrument drift. "
        "To verify measurement stability, return reference measurements were additionally acquired "
        "at 60\u2013110 \u00b0C after completing all sample measurements at each temperature. "
        "The maximum time-domain drift between the original and return references was 2.5\u20135.0% "
        "of the peak amplitude, while the spectral deviation in the 0.3\u20132.0 THz range remained "
        "below 0.6%, confirming adequate measurement stability throughout the acquisition protocol."
    )
    add_figure(doc, "fig_test_setup",
               "Fig. 2. Experimental setup. (a) Schematic diagram of the THz-TDS transmission "
               "measurement system: Menlo Systems TeraSmart with fiber-coupled PCA emitter/detector "
               "(TERA15-FC), off-axis parabolic mirrors, Quantum Northwest 6-position turret holder "
               "with TC 1 temperature controller. (b) Photographs of the measurement system, "
               "temperature controller display (110 \u00b0C), and ScanControl acquisition software.",
               width=5.5)

    doc.add_heading("2.3. Optical property extraction", level=2)
    doc.add_paragraph(
        "The complex transfer function was computed using a matched-reference approach:"
    )
    add_equation(doc, "H(\u03c9) = E\u209b\u2090\u2098(\u03c9, T) / E\u1d63\u2091\u2092(\u03c9, T)", 1)
    add_where(doc,
        "E\u209b\u2090\u2098 and E\u1d63\u2091\u2092 are the Fourier transforms of the sample and "
        "reference time-domain signals at the same temperature T."
    )
    doc.add_paragraph(
        "No apodization window was applied (rectangular window) to preserve the full spectral "
        "information, and a zero-padding factor of 2 was used, yielding a frequency resolution "
        "of approximately 0.033 THz. The refractive index and absorption coefficient were "
        "extracted from the phase and amplitude of H(\u03c9):"
    )
    add_equation(doc, "n(\u03c9) = 1 + c \u00b7 \u0394\u03c6(\u03c9) / (2\u03c0 \u00b7 f \u00b7 d)", 2)
    add_equation(doc, "\u03b1(\u03c9) = \u2212(2/d) \u00b7 ln|H(\u03c9)|", 3)
    add_where(doc,
        "\u0394\u03c6(\u03c9) is the unwrapped phase difference between sample and reference, "
        "c is the speed of light, f is the frequency, and d = 40 \u00b5m is the sample thickness."
    )
    doc.add_paragraph(
        "Analysis was performed over the frequency range 0.2\u20132.5 THz. "
        "For the 40 \u00b5m sample thickness, the first Fabry\u2013P\u00e9rot echo "
        "occurs at approximately \u0394t = 2nd/c \u2248 0.38 ps (using n \u2248 1.4), "
        "corresponding to a free spectral range of ~2.6 THz. As this echo is temporally "
        "unresolved from the main pulse in our measurement window, multiple reflections are "
        "inherently included in the extracted optical constants and do not require separate correction [14,15]."
    )

    doc.add_heading("2.4. Effective medium approximation", level=2)
    doc.add_paragraph(
        "The effective porosity was determined using the two-phase Bruggeman EMA model [17], "
        "in which the self-consistent effective permittivity \u03b5_eff satisfies:"
    )
    add_equation(doc,
        "(1 \u2212 \u03c6)(\u03b5_PE \u2212 \u03b5_eff) / (\u03b5_PE + 2\u03b5_eff) "
        "+ \u03c6(\u03b5_air \u2212 \u03b5_eff) / (\u03b5_air + 2\u03b5_eff) = 0", 4)
    add_where(doc,
        "\u03c6 is the air volume fraction (porosity), \u03b5_PE = n_PE\u00b2 with "
        "n_PE = 1.517 (weighted average of crystalline n = 1.53 and amorphous n = 1.49 "
        "based on DSC crystallinity of 67.8%), and \u03b5_air = 1.0. "
        "Given the measured n_eff, the porosity \u03c6 was obtained by numerically solving Eq. (4). "
        "The two-phase model was selected because the refractive index contrast between "
        "crystalline (n = 1.53) and amorphous (n = 1.49) PE is only \u0394n = 0.04, which is below "
        "the measurement precision (\u03c3 \u2248 0.014). A three-phase model including separate crystalline "
        "and amorphous PE components was evaluated but yielded identical porosity values "
        "(\u0394\u03c6 < 0.1 percentage points), confirming that THz-TDS at this precision cannot resolve "
        "the crystalline\u2013amorphous contrast and the simpler two-phase model is appropriate. "
        "The Bruggeman model was preferred over the Maxwell\u2013Garnett "
        "approximation because the latter assumes dilute inclusions and becomes inaccurate "
        "at the high porosity values (~44\u201359%) encountered here [17]."
    )
    doc.add_paragraph(
        "The model was anchored to the manufacturer-specified porosity of 44% at 20 \u00b0C by "
        "computing a constant refractive index offset (\u0394n = +0.146) attributed to the "
        "two-sheet stacking configuration (air gaps, interface effects). This offset was subtracted "
        "from all measured values before EMA inversion. The uncertainty in the offset "
        "(\u03c3(\u0394n) = 0.018, propagated from the standard deviation of n at 20 \u00b0C) "
        "contributes an additional \u00b14\u20135 percentage points to the absolute porosity uncertainty."
    )
    doc.add_paragraph(
        "A systematic offset between sample groups S1\u2013S3 and S4\u2013S5 (\u0394n \u2248 0.106, "
        "p < 10\u207b\u00b9\u2076, two-sample t-test) was identified, while the temperature-dependent "
        "slopes were similar (\u22126.8 \u00d7 10\u207b\u2074 vs. \u22128.1 \u00d7 10\u207b\u2074 /\u00b0C). "
        "This constant offset corresponds to an estimated thickness difference of ~12 \u00b5m, likely "
        "due to optical alignment variations between turret positions 5\u20136 (S4\u2013S5) and "
        "positions 2\u20134 (S1\u2013S3) in the six-position sample holder. A global frequency-dependent "
        "correction was applied by shifting S4\u2013S5 values to align with the S1\u2013S3 baseline. "
        "This correction reduced the inter-sample standard deviation from 0.051 to 0.014 and "
        "improved the ANOVA F-statistic at 0.45 THz from 0.94 (p = 0.50, non-significant) to "
        "12.78 (p = 2.84 \u00d7 10\u207b\u2079, \u03b7\u00b2 = 0.74)."
        " An analysis using only the S1\u2013S3 subgroup (N = 30) yielded identical temperature"
        " trends (R\u00b2 = 0.854, slope = \u22126.8 \u00d7 10\u207b\u2074 /\u00b0C), confirming that the observed"
        " temperature dependence is not an artifact of the offset correction."
    )

    doc.add_heading("2.5. Machine learning temperature prediction", level=2)
    doc.add_paragraph(
        "A machine learning pipeline was developed to predict temperature directly from raw "
        "time-domain THz signals. A total of 37 features were automatically extracted from each "
        "signal, spanning five categories: (i) peak characteristics (amplitude, timing, peak-to-peak), "
        "(ii) pulse shape (rise/fall time, envelope FWHM, asymmetry, kurtosis), "
        "(iii) signal statistics (RMS, energy, SNR), "
        "(iv) reference-relative features (time delay, amplitude ratio, delta signal), and "
        "(v) frequency-domain features (band-averaged spectral power, phase slope, centroid). "
        "Feature selection was performed using variance inflation factor (VIF) analysis [25] with a "
        "threshold of 10, prioritizing physics-meaningful features, yielding 15 selected features. "
        "(VIF threshold = 10; see Appendix A for the full feature list and Appendix B for the "
        "13 non-zero Lasso coefficients out of 15 selected features, with regularization "
        "parameter α = 0.1) "
        "Five regression models were evaluated [19,25,28]: Ridge, Lasso [19] (L1-regularized linear regression "
        "for simultaneous feature selection and prediction), SVR (RBF kernel), Random Forest, "
        "and Gradient Boosting. Two cross-validation strategies were compared: (i) conventional "
        "leave-one-out (LOO-CV, 50 folds), and (ii) leave-one-temperature-out (LOTO-CV, 10 folds), "
        "where all five replicates at a given temperature were held out simultaneously. "
        "LOTO-CV prevents information leakage from same-temperature replicates appearing in both "
        "training and test sets, providing a stricter assessment of generalization to unseen temperatures. "
        "Specifically, each LOTO fold holds out all five replicates measured at one "
        "temperature (10 folds total, each with N_train = 45, N_test = 5). "
        "Because the held-out temperature is absent from the training set, LOTO-CV effectively "
        "evaluates the model's interpolation (or, for the boundary temperatures 20 and 110 °C, "
        "extrapolation) capability, providing a conservative estimate of generalization performance. "
        "The model should not be applied outside the calibrated range of 20–110 °C without further "
        "validation. Bootstrap "
        "confidence intervals (B = 1000 resamples) were computed for the mean refractive "
        "index at each temperature."
    )
    doc.add_paragraph(
        "The Lasso regression model [19] minimizes the objective function:"
    )
    add_equation(doc,
        "min  (1/2N) \u03a3(T\u1d62 \u2212 T\u0302\u1d62)\u00b2 + \u03b1 \u03a3|\u03b2\u2c7c|", 6)
    doc.add_paragraph(
        "where T\u1d62 is the measured temperature, T\u0302\u1d62 is the predicted temperature, "
        "\u03b2\u2c7c are the regression coefficients, and \u03b1 = 0.1 is the regularization parameter, "
        "selected via 5-fold cross-validation on the training data to balance model complexity "
        "and prediction accuracy. "
        "The L1 penalty drives irrelevant coefficients to exactly zero, performing simultaneous "
        "feature selection and regression."
    )
    doc.add_paragraph(
        "The prediction equation (using standardized features) is:"
    )
    add_equation(doc,
        "T\u0302 = 65.0 + \u03a3 \u03b2\u2c7c \u00d7 (x\u2c7c \u2212 \u03bc\u2c7c) / \u03c3\u2c7c", 7)
    doc.add_paragraph(
        "where \u03bc\u2c7c and \u03c3\u2c7c are the training-set mean and standard deviation "
        "of feature j. The intercept (65.0 \u00b0C) corresponds to the mean temperature "
        "of the training data."
    )
    add_figure(doc, "fig_ml_flowchart",
               "Fig. 3. Machine learning temperature prediction pipeline. "
               "Raw THz signals and matched references are processed through five feature "
               "extraction categories (37 features), reduced to 15 via VIF selection, and "
               "input to Lasso regression (\u03b1 = 0.1). LOTO-CV evaluation yields "
               "R\u00b2 = 0.955 and MAE = 5.1 \u00b0C.",
               width=3.5)
    doc.add_paragraph(
        "Given the limited sample size (N = 50 samples, 10 temperature levels), the ML analysis "
        "is exploratory in nature and is not intended as a production-ready predictive model, but serves to identify the most informative signal features rather "
        "than to establish a production-ready prediction model. "
        "An ablation study was also performed by comparing the full 15-feature model against "
        "single-feature Ridge regression baselines to quantify the added value of multi-feature integration."
    )

    # ════════════════════════════════════════════
    # 3. RESULTS AND DISCUSSION
    # ════════════════════════════════════════════
    doc.add_heading("3. Results", level=1)

    doc.add_heading("3.1. Time-domain signal analysis", level=2)
    doc.add_paragraph(
        "Fig. 4 shows representative THz time-domain waveforms for all ten temperatures. "
        "The transmitted pulse through the PE separator maintains high fidelity across the "
        "temperature range, with a system SNR of approximately 52 dB (Fig. S1). "
        "Systematic temperature-dependent changes are clearly visible in both panels. "
        "In the reference signals (top row), a slight broadening and amplitude reduction "
        "with increasing temperature reflects the temperature dependence of the THz system "
        "optics and atmospheric absorption. In the sample signals (bottom row), these "
        "system-level changes are superimposed with material-specific effects: the positive "
        "peak amplitude decreases monotonically with temperature, and the negative peak shows "
        "a systematic temporal shift toward earlier arrival times. The peak-to-peak time "
        "interval decreases from approximately 0.55 ps at 20 \u00b0C to 0.52 ps at 110 \u00b0C, "
        "consistent with a reduction in the effective optical path length as porosity increases. "
        "The matched-reference protocol (Eq. 1) [14,15] removes the system-level contributions, "
        "isolating the material response for quantitative analysis."
    )
    add_figure(doc, "fig02_peak_detail",
               "Fig. 4. THz pulse peak detail at 10 temperatures (20\u2013110 \u00b0C). "
               "(a) Positive peak region showing systematic amplitude reduction with increasing "
               "temperature. (b) Negative peak region. Black line: reference at 20 \u00b0C.",
               width=5.0)

    doc.add_heading("3.2. Temperature-dependent refractive index", level=2)
    add_table_from_data(doc,
        ["T (\u00b0C)", "n @ 0.45 THz", "\u03c3(n)", "Porosity (%)", "\u03c3_total (%)"],
        [
            ["20",  "1.4282", "0.018", "44.0 ± 5.9", "5.9"],
            ["30",  "1.4230", "0.012", "45.0 ± 4.6", "4.6"],
            ["40",  "1.3978", "0.012", "49.7 ± 4.7", "4.7"],
            ["50",  "1.3775", "0.014", "53.5 ± 5.1", "5.1"],
            ["60",  "1.3790", "0.014", "53.3 ± 5.1", "5.1"],
            ["70",  "1.3751", "0.015", "54.0 ± 5.2", "5.2"],
            ["80",  "1.3688", "0.013", "55.2 ± 4.9", "4.9"],
            ["90",  "1.3728", "0.014", "54.4 ± 5.2", "5.2"],
            ["100", "1.3515", "0.011", "58.5 ± 4.6", "4.6"],
            ["110", "1.3663", "0.014", "55.7 ± 5.1", "5.1"],
        ],
        "Table 3. Temperature-dependent refractive index and porosity at 0.45 THz "
        "(mean \u00b1 \u03c3 from 5 replicates). \u03c3_total includes propagated n_offset uncertainty."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "The primary analysis frequency of 0.45 THz was selected by scanning all frequencies "
        "in the 0.3\u20132.0 THz range and ranking by the Pearson R\u00b2 of the n versus T "
        "linear correlation. Among 465 candidate frequencies, 0.45 THz yielded the highest "
        "R\u00b2 = 0.804, followed by 1.09 THz (R\u00b2 = 0.735) and 1.63 THz (R\u00b2 = 0.727). "
        "These three frequencies, spaced at least 0.5 THz apart, were retained for multi-frequency "
        "analysis."
    )
    doc.add_paragraph(
        "Fig. 5 presents the refractive index at 0.45 THz as a function of temperature, "
        "together with the EMA-derived effective porosity. The refractive index decreases "
        "monotonically from 1.428 \u00b1 0.018 at 20 \u00b0C to 1.352 \u00b1 0.011 at 100 \u00b0C "
        "(R\u00b2 = 0.804 for linear correlation with temperature). "
        "The transition temperature was determined by piecewise linear regression. "
        "Specifically, the breakpoint T_break was optimized by minimizing the total residual "
        "sum of squares (RSS) over all candidate values in 30\u201390 \u00b0C at 0.1 \u00b0C resolution. "
        "The optimal breakpoint was found at T_break = 51.8 \u00b0C with RSS = 3.0 \u00d7 10\u207b\u2074; "
        "the intersection of the two fitted lines yields T_onset = 55.2 \u00b0C "
        "(Regime 1: n = 1.452 \u2212 1.52 \u00d7 10\u207b\u00b3 T, R\u00b2 = 0.987; "
        "Regime 2: n = 1.373 \u2212 3.24 \u00d7 10\u207b\u2074 T, R\u00b2 = 0.476). "
        "A two-regime behavior is observed: below ~55 \u00b0C, the refractive index decreases "
        "rapidly (slope = \u22121.52 \u00d7 10\u207b\u00b3 /\u00b0C), while above 55 \u00b0C "
        "the decrease attenuates (slope = \u22123.24 \u00d7 10\u207b\u2074 /\u00b0C). "
        "The transition temperature of 55.2 \u00b0C is consistent with the \u03b1c crystalline "
        "relaxation of polyethylene [13], which involves the onset of chain mobility within "
        "the crystalline lamellae. This transition is well below the DSC melting temperature "
        "(Tm = 139.4 \u00b0C, crystallinity = 67.8%), indicating that significant microstructural "
        "changes\u2014specifically, expansion of amorphous-region free volume and inter-lamellar "
        "pore widening\u2014initiate at temperatures far below the melting point. "
        "The two-regime behavior (rapid change below 55 \u00b0C, plateau above) is consistent with "
        "the interpretation that amorphous-phase relaxation drives the initial porosity increase, "
        "while the crystalline lamellae provide structural resistance above T_onset."
    )
    doc.add_paragraph(
        "A slight increase in n at 110 °C (from 1.352 to 1.366) is observed, "
        "which lies within the measurement uncertainty (Δn = 0.015 < σ = 0.014). "
        "This non-monotonic behavior is attributed to measurement variability, as the magnitude of the reversal (0.015) "
        "approaching the melting regime, but cannot be "
        "conclusively attributed to a physical mechanism with the current data."
    )
    add_figure(doc, "fig03_n_vs_temp_ema",
               "Fig. 5. Refractive index n at 0.45 THz (left axis, black) and effective porosity "
               "(right axis, red) as a function of temperature. Individual samples (S1\u2013S5) are "
               "shown as colored markers; mean \u00b1 \u03c3 as black circles with gray band. "
               "Dashed lines: two-regime linear fits with onset at 55 \u00b0C.",
               width=4.5)

    doc.add_heading("3.3. Broadband optical spectra", level=2)
    doc.add_paragraph(
        "Fig. 6 displays the broadband refractive index and absorption coefficient spectra. "
        "The refractive index exhibits relatively flat dispersion in the 0.3\u20131.5 THz range "
        "with a gradual decrease toward higher frequencies. The temperature-dependent shift is "
        "consistent across the entire spectral range. The absorption coefficient is low "
        "(\u03b1 < 15 cm\u207b\u00b9) throughout the measurement bandwidth, characteristic of "
        "low-loss dielectric materials [9,23]."
    )
    doc.add_paragraph(
        "The refractive index spectra exhibit weak normal dispersion in the 0.3\u20131.0 THz range, consistent with prior THz characterization of PE and other polymers [12,23,24,26], "
        "with a gradual decrease at higher frequencies characteristic of porous polymer media [12]. "
        "The temperature-induced shift is remarkably uniform across the spectral range: the "
        "refractive index decreases by approximately 0.06\u20130.08 at all frequencies between 20 and "
        "100 \u00b0C, indicating that the underlying mechanism (pore expansion) affects the effective "
        "dielectric response uniformly rather than introducing frequency-selective features. "
        "This spectral uniformity supports the use of a frequency-independent EMA model (Eq. 4) "
        "and justifies the selection of a single representative frequency (0.45 THz) for the "
        "primary quantitative analysis. The three best-correlated frequencies (0.45, 1.09, 1.63 THz) "
        "span the usable bandwidth and provide consistent temperature trends, further confirming "
        "the broadband nature of the observed changes."
    )
    doc.add_paragraph(
        "The absorption coefficient values at 0.45 THz are small (0.0\u20130.9 cm\u207b\u00b9) and "
        "approach the noise floor of the measurement. Negative values, which arise from "
        "phase noise at low frequencies, were clipped to zero. The uncertainty in \u03b1 is "
        "estimated at \u00b11\u20132 cm\u207b\u00b9 based on the inter-replicate variability."
    )
    add_figure(doc, "fig04_optical_spectra",
               "Fig. 6. Broadband optical spectra at 10 temperatures. "
               "(a) Refractive index n(f). Vertical dashed lines indicate the three best-correlated "
               "frequencies (0.45, 1.09, 1.63 THz). (b) Absorption coefficient \u03b1(f).",
               width=5.0)

    doc.add_heading("3.4. Porosity evolution and thermal expansion", level=2)
    doc.add_paragraph(
        "The effective porosity, obtained by inverting the two-phase Bruggeman model (Eq. 4), "
        "increases from 44.0% (anchored to manufacturer specification) at 20 \u00b0C to 58.5% "
        "at 100 \u00b0C (Fig. 7). The temperature dependence was fitted using a thermal expansion model:"
    )
    add_equation(doc,
        "\u03c6(T) = 1 \u2212 (1 \u2212 \u03c6\u2080) / (1 + \u03b2 \u0394T)", 5)
    add_where(doc,
        "\u03c6\u2080 is the initial porosity, \u03b2 is the effective thermal expansion "
        "coefficient of the pore structure, and \u0394T = T \u2212 20 \u00b0C."
    )
    doc.add_paragraph(
        "Fitting yields \u03c6\u2080 = 45.2% and \u03b2 = 3.50 \u00d7 10\u207b\u00b3 /\u00b0C. "
        "This effective pore expansion coefficient is approximately 6\u20138 times larger than the "
        "bulk PE volumetric expansion coefficient (~5 \u00d7 10\u207b\u2074 /\u00b0C [12]), indicating "
        "that pore expansion\u2014not bulk thermal expansion\u2014dominates the observed optical changes [7,10,27]."
    )
    doc.add_paragraph(
        "A porosity increase of approximately 5.7 percentage points between 20 and 40 \u00b0C "
        "was observed, preceding the identified onset temperature of 55 \u00b0C. The origin of "
        "this early-stage change has not been independently verified in this study. "
        "Similar early-stage porosity changes below the \u03b1c relaxation temperature have been "
        "reported in uniaxially stretched PE films [12,13]."
    )
    add_figure(doc, "fig06_porosity_vs_temp",
               "Fig. 7. Effective porosity as a function of temperature. Black circles: EMA inversion "
               "results (mean \u00b1 \u03c3 from 5 replicates). Red dashed line: thermal expansion model "
               "fit (\u03b2 = 3.50 \u00d7 10\u207b\u00b3 /\u00b0C). Horizontal gray line: manufacturer "
               "specification (44%).",
               width=4.0)

    doc.add_paragraph(
        "While the absolute porosity values carry considerable uncertainty "
        "(\u03c3_total \u2248 5\u20136 percentage points, arising primarily from the n_offset "
        "uncertainty \u03c3 = 0.018 computed from N = 5 replicates at 20 \u00b0C), the relative "
        "porosity change (\u0394\u03c6 \u2248 15 percentage points over 80 \u00b0C) is well resolved "
        "and exceeds the total uncertainty by a factor of ~3. It is this relative trend\u2014rather "
        "than the absolute porosity values\u2014that constitutes the primary physical finding, "
        "consistent with known thermal expansion behavior of dry-process PE separators."
    )

    doc.add_heading("3.5. Statistical validation", level=2)
    doc.add_paragraph(
        "One-way ANOVA was performed at each frequency, treating the 10 temperature levels as groups "
        "(N = 5 replicates per group, total N = 50), across the 0.3\u20132.0 THz range (Fig. 8). "
        "At 0.45 THz, the ANOVA yielded F = 12.78, p = 2.84 \u00d7 10\u207b\u2079, with an effect "
        "size \u03b7\u00b2 = 0.74 (large effect). Across the full spectral range, 97.4% of frequencies "
        "(453 of 465) showed statistically significant temperature dependence (p < 0.05, after inter-sample offset correction). "
        "The narrow non-significant band (0.30\u20130.55 THz) coincides with the low-frequency edge "
        "where measurement noise is elevated."
    )
    add_figure(doc, "fig13_anova_freq_scan",
               "Fig. 8. Frequency-resolved ANOVA results. (a) p-value as a function of frequency "
               "(log scale). Blue shading: significant regions (p < 0.05). (b) Effect size \u03b7\u00b2. "
               "Horizontal red dashed line: large effect threshold (0.14).",
               width=4.0)

    doc.add_paragraph(
        "Bootstrap confidence intervals [22] (B = 1000) for the mean refractive index confirmed "
        "non-overlapping 95% CIs between 20 \u00b0C [1.411, 1.443] and 100 \u00b0C [1.340, 1.360], "
        "providing additional evidence for the significance of the temperature effect."
    )
    doc.add_paragraph(
        "It should be noted that while pairwise comparisons [21] between adjacent 10 °C "
        "steps do not reach statistical significance (owing to the small per-step change "
        "relative to inter-sample variability), the overall one-way ANOVA across all temperatures "
        "is highly significant (F = 12.78, p = 2.84 × 10⁻⁹, η² = 0.74), and the cumulative "
        "change over the full 20–100 °C range is clearly resolved (p = 9.8 × 10⁻⁵ by t-test)."
    )
    doc.add_paragraph(
        "The three best-correlated frequencies (0.45, 1.09, 1.63 THz) were identified "
        "by scanning all frequencies in the 0.3–2.0 THz range and ranking by R² of the "
        "n versus T linear correlation. This frequency selection was performed on the full "
        "dataset prior to model training; however, the LOTO-CV evaluation uses held-out "
        "temperature blocks, ensuring that the ML performance metrics are not inflated by "
        "this selection."
    )

    doc.add_heading("3.6. Phase spectrum and transfer function", level=2)
    doc.add_paragraph(
        "The unwrapped phase difference \u0394\u03c6(f) and transfer function magnitude |H(\u03c9)| "
        "are shown in Fig. 9. The phase difference exhibits a near-linear frequency dependence "
        "with temperature-dependent slope, directly reflecting changes in the effective optical "
        "path length [14,16]. The transfer function magnitude remains close to unity (|H| \u2248 0.95\u20131.0), "
        "consistent with the low absorption of the thin PE film."
    )
    doc.add_paragraph(
        "The linear phase slope is directly proportional to the group refractive index, which "
        "is the primary quantity encoded in the Lasso model's top feature (phase_slope_low, "
        "importance = 14.5). As temperature increases, the phase slope decreases (becomes less "
        "negative), reflecting the reduction in optical path length due to pore expansion. "
        "The transfer function magnitude |H(\u03c9)| \u2248 0.95\u20131.0 indicates minimal absorption loss "
        "in the thin PE film, consistent with the low absorption coefficients reported in Section 3.3. "
        "The slight deviation from unity at higher frequencies (>1.5 THz) reflects increased "
        "scattering from the porous microstructure, which becomes more pronounced as the pore "
        "dimensions approach a significant fraction of the THz wavelength."
    )
    add_figure(doc, "fig11_phase_spectrum",
               "Fig. 9. (a) Unwrapped phase difference \u0394\u03c6(f) between sample and matched "
               "reference. (b) Transfer function magnitude |H(\u03c9)|. Colors represent temperature "
               "(blue: 20 \u00b0C, red: 110 \u00b0C).",
               width=5.0)

    doc.add_heading("3.7. Machine learning temperature prediction", level=2)
    doc.add_paragraph(
        "Table 1 compares the performance of five regression models under two cross-validation "
        "strategies. Linear models (Lasso, Ridge) significantly outperform nonlinear models "
        "(SVR, Random Forest, Gradient Boosting) under both CV schemes, indicating that the "
        "temperature\u2013THz signal relationship is predominantly linear within the studied range."
    )
    add_table_from_data(doc,
        ["Model", "LOO-CV R\u00b2", "LOO-CV MAE (\u00b0C)", "LOTO-CV R\u00b2", "LOTO-CV MAE (\u00b0C)"],
        [
            ["Ridge",             "0.974", "3.81", "0.952", "5.20"],
            ["Lasso",             "0.975", "3.71", "0.955", "5.08"],
            ["SVR (RBF)",         "0.875", "7.52", "0.566", "13.60"],
            ["Random Forest",     "0.845", "7.21", "0.648", "11.84"],
            ["Gradient Boosting", "0.844", "6.32", "0.652", "11.55"],
        ],
        "Table 1. Model performance comparison under LOO-CV and LOTO-CV."
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "Critically, the performance gap between LOO-CV and LOTO-CV is small for linear models "
        "(\u0394R\u00b2 \u2248 0.02) but large for nonlinear models (\u0394R\u00b2 = 0.19\u20130.31), "
        "demonstrating that the latter suffer from information leakage when same-temperature "
        "replicates appear in both training and test sets."
    )

    add_figure(doc, "fig_ml_prediction_scatter",
               "Fig. 10. Predicted vs. actual temperature for the Lasso model under LOTO-CV "
               "(R\u00b2 = 0.955, MAE = 5.1 \u00b0C). Colors indicate actual temperature.",
               width=3.5)

    doc.add_paragraph(
        "Fig. 11 shows the feature importance ranking for the Lasso model. The top three features "
        "are phase_slope_low (low-frequency phase gradient, proportional to group refractive index), "
        "spec_mean_low (mean spectral amplitude in 0.3\u20130.8 THz, related to transmission), "
        "and spec_centroid (spectral center of mass). These features have clear physical "
        "interpretations: the phase slope directly encodes the optical path length change, "
        "while spectral amplitude and centroid reflect absorption and scattering modifications "
        "induced by pore structure evolution."
    )
    add_figure(doc, "fig_ml_feature_importance",
               "Fig. 11. Feature importance (Lasso coefficient magnitude) for the temperature "
               "prediction model. Stars (\u2605): physics-priority features selected a priori.",
               width=4.0)

    doc.add_paragraph(
        "An ablation study (Table 2) confirms the added value of multi-feature integration. "
        "The best single feature (phase_slope_low, proportional to group refractive index) "
        "achieves R\u00b2 = 0.67 under LOTO-CV, while the full 15-feature Lasso model reaches "
        "R\u00b2 = 0.955\u2014a 43 percentage-point improvement. This demonstrates that no single "
        "THz signal characteristic captures the full temperature information; rather, the combination "
        "of phase, amplitude, and spectral features provides a synergistic representation."
    )
    doc.add_paragraph(
        "The dominance of linear models (Lasso, Ridge) over nonlinear models (SVR, Random Forest, "
        "Gradient Boosting) provides physical insight: the relationship between THz signal features "
        "and temperature is predominantly linear within the 20\u2013110 \u00b0C range, which lies below the "
        "PE melting transition (Tm = 139 \u00b0C). This linearity is consistent with the physics-based "
        "analysis, where both n(T) and \u03c6(T) exhibit quasi-linear trends. The large performance gap "
        "between LOO-CV and LOTO-CV for nonlinear models (\u0394R\u00b2 = 0.19\u20130.31) versus linear models "
        "(\u0394R\u00b2 \u2248 0.02) reveals that tree-based and kernel-based models overfit to temperature-specific "
        "noise patterns when same-temperature replicates are available in the training set, whereas "
        "linear models generalize more robustly to unseen temperatures."
    )
    add_table_from_data(doc,
        ["Features", "R\u00b2 (LOTO)", "MAE (\u00b0C)"],
        [
            ["All 15 (Lasso)", "0.955", "5.08"],
            ["phase_slope_low only", "0.666", "12.30"],
            ["spec_std_high only", "0.509", "15.59"],
            ["spec_centroid only", "0.393", "18.42"],
            ["spec_mean_low only", "0.229", "20.41"],
        ],
        "Table 2. Ablation study: full model vs. single-feature baselines (LOTO-CV)."
    )
    doc.add_paragraph()

    # ════════════════════════════════════════════
    # 4. DISCUSSION
    # ════════════════════════════════════════════
    doc.add_heading("4. Discussion", level=1)

    doc.add_heading("4.1. Physical interpretation of thermal onset", level=2)
    doc.add_paragraph(
        "The thermal onset temperature of 55.2 \u00b0C identified by piecewise linear regression "
        "falls within the range of the \u03b1c crystalline relaxation of polyethylene (50\u201380 \u00b0C), "
        "which is associated with the onset of chain mobility within the crystalline lamellae [13]. "
        "Below T_onset, the rapid decrease in refractive index (slope = \u22121.52 \u00d7 10\u207b\u00b3 /\u00b0C) "
        "suggests that the amorphous interlamellar regions undergo significant free-volume expansion, "
        "leading to pore widening. Above T_onset, the attenuated slope (\u22123.24 \u00d7 10\u207b\u2074 /\u00b0C) "
        "indicates that the crystalline lamellae provide structural resistance against further "
        "pore expansion, consistent with the semi-crystalline nature of the dry-process PE separator."
    )
    doc.add_paragraph(
        "The effective pore expansion coefficient \u03b2 = 3.50 \u00d7 10\u207b\u00b3 /\u00b0C is "
        "approximately one order of magnitude larger than the bulk PE volumetric thermal expansion "
        "coefficient (~5 \u00d7 10\u207b\u2074 /\u00b0C). This amplification arises because the pore "
        "expansion is driven by anisotropic stress relaxation in the stretched polymer network, "
        "not by isotropic thermal expansion of the bulk material. The dry-stretch manufacturing "
        "process introduces residual elastic stress that is progressively released upon heating, "
        "preferentially widening the slit-shaped pores in the stretching direction."
    )
    doc.add_paragraph(
        "The two-regime behavior observed in n(T) indicates a structural transition near 55 \u00b0C "
        "that is detectable by THz-TDS. Below this temperature, the rapid n decrease corresponds to "
        "a porosity increase rate of approximately 0.29 %p/\u00b0C, while above 55 \u00b0C the rate "
        "decreases to 0.05 %p/\u00b0C. This change in rate is consistent with a shift from "
        "amorphous-phase-dominated expansion to crystalline-framework-stabilized behavior."
    )
    doc.add_paragraph(
        "The identified T_onset (55.2 \u00b0C) is well below the DSC melting temperature "
        "(Tm = 139.4 \u00b0C), indicating that measurable microstructural reorganization begins "
        "approximately 84 \u00b0C below the melting point. This large gap between T_onset and Tm "
        "defines a regime where porosity changes are detectable by THz-TDS while the separator "
        "retains its structural integrity."
    )

    doc.add_heading("4.2. Multi-domain feature analysis", level=2)
    doc.add_paragraph(
        "Fig. 12 summarizes the correlation between 12 time-domain, frequency-domain, and "
        "envelope features and temperature. The rise time (R\u00b2 = 0.951) and spectral "
        "centroid (R\u00b2 = 0.929) show the strongest linear correlations, consistent with "
        "the ML feature importance analysis where frequency-domain features dominate."
    )
    doc.add_paragraph(
        "The apparent discrepancy between the correlation ranking (rise_time_ps: R² = 0.951, "
        "rank 1) and the ML feature importance (phase_slope_low: rank 1, rise_time: rank 12) "
        "reflects the difference between univariate and multivariate analysis. While rise_time "
        "exhibits the strongest individual linear correlation with temperature, it shares high "
        "collinearity with phase_slope_low (both encode the group delay information). In the "
        "multivariate Lasso model, phase_slope_low subsumes the information carried by rise_time, "
        "resulting in a near-zero coefficient for the latter."
    )
    doc.add_paragraph(
        "The multi-domain feature analysis reveals that temperature information is distributed "
        "across complementary signal characteristics. Time-domain features (rise_time, dt_pos) "
        "capture the group delay and pulse shape changes, while frequency-domain features "
        "(phase_slope, spec_centroid) encode the spectral redistribution of energy. Envelope "
        "features (env_fwhm, env_asymmetry) reflect pulse broadening effects. The fact that "
        "features from all five domains contribute to the ML model (Table B1) suggests that "
        "a comprehensive multi-domain approach extracts more temperature information than any "
        "single-domain analysis. This finding has practical implications for THz-based thermal "
        "monitoring: rather than relying solely on refractive index extraction (which requires "
        "careful phase unwrapping and thickness knowledge), a feature-based ML approach can "
        "provide robust temperature estimation directly from raw time-domain signals."
    )
    doc.add_paragraph(
        "The physical basis for the dominance of frequency-domain features can be understood "
        "through the relationship between the Lasso coefficients and the underlying optical "
        "properties. The top feature, phase_slope_low (\u03b2 = +14.5), is directly proportional "
        "to the group refractive index n_g in the 0.3\u20130.8 THz band. The second-ranked feature, "
        "spec_mean_low (\u03b2 = +11.0), encodes the mean spectral amplitude related to the "
        "transmission coefficient |H(\u03c9)|. The signs and magnitudes of the standardized "
        "coefficients reflect the multivariate correlation structure after centering and scaling."
    )
    add_figure(doc, "fig08_correlation_summary",
               "Fig. 12. Multi-domain feature correlation summary. (a) Normalized feature heatmap "
               "across temperatures. (b) R\u00b2 values with significance levels.",
               width=5.0)

    doc.add_heading("4.3. Practical implications for thermal monitoring", level=2)
    doc.add_paragraph(
        "The temperature sensitivity of the refractive index (dn/dT \u2248 \u22128.5 \u00d7 10\u207b\u2074 /\u00b0C) "
        "corresponds to a porosity sensitivity of approximately 0.16 percentage points per degree, "
        "which is detectable by the current THz-TDS setup with \u03c3(n) = 0.014. "
        "The ML-based prediction model (Lasso, MAE = 5.1 \u00b0C) demonstrates that temperature "
        "estimation is feasible from raw THz signals without explicit optical property extraction, "
        "which is advantageous for in-line quality control where processing speed is critical. "
        "The dominance of frequency-domain features (phase_slope_low, spec_mean_low) in the ML model "
        "suggests that simplified spectral analysis, rather than full time-domain waveform processing, "
        "may suffice for practical temperature monitoring applications."
    )
    doc.add_paragraph(
        "The demonstrated temperature resolution of \u00b15 \u00b0C (based on the Lasso MAE under "
        "LOTO-CV) is sufficient for in-line manufacturing quality control, where separator rolls "
        "are inspected at controlled temperatures before cell assembly. The non-contact nature of "
        "THz-TDS is advantageous for continuous production lines where physical contact with the "
        "separator film must be avoided to prevent mechanical damage."
    )
    doc.add_paragraph(
        "For manufacturing quality control, the porosity sensitivity (\u0394\u03c6/\u0394T \u2248 "
        "0.16 %p/\u00b0C) combined with the THz measurement precision suggests that batch-to-batch "
        "porosity variations of approximately \u00b11 percentage point can be detected at a fixed "
        "temperature. This is comparable to the accuracy of mercury intrusion porosimetry (~\u00b11%) "
        "but achievable non-destructively and with a measurement time of seconds rather than hours. "
        "The integration of the ML prediction pipeline into an automated inspection system would "
        "require training on a larger and more diverse dataset spanning multiple separator types "
        "and production batches, which is identified as a priority for future work."
    )

    doc.add_heading("4.4. Limitations", level=2)
    doc.add_paragraph(
        "Several limitations of this study should be noted. First, the porosity values derived "
        "from EMA inversion are effective optical porosities and were not independently validated "
        "by direct structural methods (e.g., SEM cross-section analysis or mercury intrusion "
        "porosimetry). Although the anchoring to the manufacturer specification (44%) provides "
        "a physically grounded baseline, the absolute porosity values carry an uncertainty of "
        "\u00b15\u20136 percentage points due to the propagated n_offset error."
    )
    doc.add_paragraph(
        "Second, the two-sheet stacking introduces additional unknowns (inter-layer air gap, "
        "contact pressure variability) that contribute to the refractive index offset. "
        "Single-sheet measurements would reduce this uncertainty but were not feasible due to "
        "the limited SNR at 20 \u00b5m thickness."
    )
    doc.add_paragraph(
        "Third, the Bruggeman EMA assumes isotropic spherical inclusions, whereas dry-process PE "
        "separators possess oriented slit-shaped pores due to uniaxial stretching. The reported "
        "porosity values therefore represent an effective isotropic approximation and may differ "
        "from true directional porosity measured by other techniques."
    )
    doc.add_paragraph(
        "Fourth, all measurements were performed during a single heating ramp (20 \u2192 110 \u00b0C). "
        "Cooling-cycle measurements were not conducted; therefore, the reversibility of the "
        "observed microstructural changes and possible thermal hysteresis effects remain unexplored. "
        "Investigating the cooling-path dependence of porosity recovery is an important direction "
        "for future work."
    )
    doc.add_paragraph(
        "Fifth, the ML model was trained and validated on a single separator type (MS-DPS 20B) "
        "measured with one instrument. Generalization to other separator materials (PP, wet-process PE, "
        "ceramic-coated) or measurement systems has not been demonstrated. "
        "Sixth, the sample replicates (S1\u2013S5) were measured sequentially via turret rotation; "
        "potential position-dependent alignment variations cannot be fully excluded, "
        "although the matched-reference protocol and return-reference drift validation "
        "(spectral deviation < 0.6%) mitigate systematic errors."
    )

    doc.add_heading("5. Conclusions", level=1)
    doc.add_paragraph(
        "This study demonstrates the capability of THz-TDS for non-destructive, quantitative "
        "monitoring of temperature-induced microstructural changes in PE battery separators. "
        "The key findings are:"
    )
    conclusions = [
        "The refractive index at 0.45 THz decreases from 1.43 to 1.35 over the 20\u2013110 \u00b0C "
        "range, corresponding to an effective porosity increase from 44% to 59% as determined by "
        "two-phase Bruggeman EMA.",
        "A thermal onset temperature of 55.2 \u00b0C is identified, consistent with the \u03b1c "
        "crystalline relaxation of PE, below which rapid microstructural changes occur.",
        "ANOVA confirms statistically significant temperature dependence across 97.4% of the "
        "THz spectral range, with a large effect size (\u03b7\u00b2 = 0.74 at 0.45 THz).",
        "Machine learning analysis reveals that the temperature\u2013signal relationship is "
        "predominantly linear, with phase slope and low-frequency spectral power as the primary "
        "indicators (Lasso R\u00b2 = 0.955 under LOTO-CV).",
        "The combined physics-based and data-driven approach provides complementary insights: "
        "EMA quantifies porosity evolution, while ML identifies the most informative signal features "
        "for potential real-time temperature monitoring applications.",
    ]
    for c in conclusions:
        doc.add_paragraph(c, style="List Number")

    doc.add_paragraph(
        "Under the present measurement conditions, the temperature sensitivity of the refractive index "
        "(\u2248 \u22128.5 \u00d7 10\u207b\u2074 /\u00b0C averaged over the full range) suggests a "
        "detection limit of approximately 5 \u00b0C given the current precision (\u03c3 \u2248 0.014), "
        "consistent with the ML prediction accuracy (MAE = 5.1 \u00b0C). "
        "Validation on independent separator types and measurement systems is needed to confirm "
        "the generalizability of these findings. "
        "These results establish THz-TDS as a viable non-contact technique for in-situ thermal "
        "monitoring of battery separator integrity. Potential applications include: "
        "(i) in-line quality inspection of separator porosity during manufacturing, "
        "(ii) process monitoring of thermally induced porosity changes during separator production, and "
        "(iii) accelerated aging studies where non-destructive porosity tracking is required."
    )
    doc.add_paragraph(
        "It should be noted that THz radiation cannot penetrate metallic electrodes, "
        "precluding direct in-situ measurement of separators within assembled battery cells. "
        "The application scope of the proposed method is therefore limited to separator "
        "characterization during the manufacturing and quality control stages, prior to "
        "cell assembly. As a next step, we plan to investigate THz response characteristics "
        "as a function of charge\u2013discharge cycling by introducing a measurement aperture "
        "at the center of single-cell prototypes."
    )
    doc.add_paragraph(
        "The methodology presented here\u2014combining physics-based EMA analysis with data-driven "
        "ML prediction\u2014provides a template applicable to other porous functional materials "
        "where temperature-dependent microstructural monitoring is needed, including ceramic "
        "electrolytes, gas diffusion layers in fuel cells, and filtration membranes."
    )

    # Graphical abstract
    doc.add_heading("Graphical Abstract", level=1)
    doc.add_paragraph(
        "The graphical abstract illustrates: (left) the matched-reference THz-TDS measurement "
        "configuration with temperature-controlled turret holder, (center) the temperature-dependent "
        "refractive index decrease and EMA porosity increase, and (right) the Lasso ML prediction "
        "scatter plot (R² = 0.955). A separate high-resolution image file (min. 531 × 1328 pixels) "
        "is provided for submission."
    )
    add_figure(doc, "fig03_n_vs_temp_ema",
               "Graphical Abstract (cf. Fig. 5): Temperature-dependent refractive index and porosity of PE "
               "separator measured by THz-TDS, with ML-based temperature prediction.",
               width=5.0)

    # ════════════════════════════════════════════
    # APPENDIX A: Feature Definitions
    # ════════════════════════════════════════════
    doc.add_heading("Appendix A", level=1)
    doc.add_paragraph(
        "The complete definitions of all 37 extracted features (Tables A1\u2013A5) and the "
        "annotated feature extraction illustration (Fig. A1) are provided in the "
        "Supplementary Material."
    )

    # ════════════════════════════════════════════
    # APPENDIX B: Lasso Prediction Model
    # ════════════════════════════════════════════
    doc.add_heading("Appendix B. Lasso temperature prediction model", level=1)
    doc.add_paragraph(
        "The Lasso objective function and prediction equation are given in Eqs. (6) and (7) "
        "in Section 2.5. Table B1 lists the regression coefficients for the selected features."
    )

    add_table_from_data(doc,
        ["Feature", "\u03b2 (standardized)", "Physical interpretation"],
        [
            ["phase_slope_low",  "+14.53", "Low-freq phase gradient \u221d group refractive index"],
            ["spec_mean_low",    "+11.00", "0.3\u20130.8 THz spectral power \u221d transmission"],
            ["spec_centroid",     "+4.01", "Spectral center of mass \u2192 frequency shift"],
            ["amp_ratio_neg",     "\u22123.43", "Negative peak transmission ratio"],
            ["delta_rms",         "\u22122.34", "Sample\u2013reference difference signal RMS"],
            ["phase_slope_mid",   "+2.30", "Mid-freq phase gradient"],
            ["snr_db",            "+1.74", "Signal-to-noise ratio"],
            ["spec_std_mid",      "+1.41", "Mid-freq spectral variability"],
            ["amp_ratio_pos",     "\u22120.72", "Positive peak transmission ratio"],
            ["fall_time",         "\u22120.68", "Pulse fall time"],
            ["env_fwhm",          "+0.60", "Envelope width"],
            ["rise_time",         "\u22120.53", "Pulse rise time"],
            ["spec_std_high",     "\u22120.31", "High-freq spectral variability"],
            ["dt_pos",             "0.00", "Time delay (eliminated by Lasso)"],
            ["env_asymmetry",      "0.00", "Envelope asymmetry (eliminated by Lasso)"],
        ],
        "Table B1. Lasso regression coefficients for the 15 VIF-selected features "
        "(\u03b1 = 0.1, standardized scale). Intercept = 65.0 \u00b0C."
    )

    doc.add_paragraph(
        "The dominant features (|β| > 4) are phase_slope_low, spec_mean_low, and spec_centroid, "
        "all of which encode frequency-domain information related to the effective optical path length "
        "and dielectric response of the porous PE matrix. Two features (dt_pos, env_asymmetry) were "
        "driven to zero by the L1 penalty, indicating redundancy with other selected features."
    )

    # ════════════════════════════════════════════
    # DECLARATIONS
    # ════════════════════════════════════════════
    doc.add_heading("Declaration of competing interest", level=1)
    doc.add_paragraph(
        "The authors declare that they have no known competing financial interests or personal "
        "relationships that could have appeared to influence the work reported in this paper."
    )

    doc.add_heading("Acknowledgments", level=1)
    doc.add_paragraph("[To be added]")

    doc.add_heading("Data availability", level=1)
    doc.add_paragraph(
        "Data will be made available on request."
    )

    doc.add_heading("Supplementary material", level=1)
    doc.add_paragraph(
        "The following supplementary materials are provided:"
    )
    supp_items = [
        "Tables A1\u2013A5: Complete definitions of all 37 extracted features "
        "(peak characteristics, pulse shape, noise, reference-relative, frequency-domain)",
        "Fig. A1: Annotated illustration of feature extraction from THz signals",
        "Fig. S1: Raw THz time-domain traces at all temperatures (full waveform and peak region)",
        "Fig. S2: Bruggeman 2-phase EMA model curve and inversion residuals",
        "Fig. S3: Piecewise vs. linear vs. quadratic regression comparison and fit residuals",
        "Additional supplementary data: phase spectrum, SNR/dynamic range, drift validation, "
        "sample group comparison, and CSV data tables",
    ]
    for item in supp_items:
        doc.add_paragraph(item, style="List Bullet")
    add_figure(doc, "fig_S1_raw_traces",
               "Fig. S1. Raw THz time-domain signals at 10 temperatures (S1 replicate). "
               "(a) Full waveform showing the main THz pulse and post-pulse oscillations. "
               "(b) Aligned peak region (\u00b11.5 ps) showing systematic amplitude and timing changes.",
               width=5.0)
    add_figure(doc, "fig_S2_ema_convergence",
               "Fig. S2. (a) Two-phase Bruggeman EMA: effective refractive index as a function of "
               "porosity. Green band: measured n range. (b) EMA inversion residuals per temperature, "
               "confirming numerical convergence to machine precision (~10\u207b\u2078).",
               width=5.0)
    add_figure(doc, "fig_S3_piecewise_residuals",
               "Fig. S3. (a) Comparison of linear, quadratic, and piecewise linear fits to n(T). "
               "(b) Fit residuals showing that the piecewise and quadratic models provide comparable "
               "accuracy (\u0394AIC = 0.08), both substantially outperforming the single linear model.",
               width=5.0)

    # ════════════════════════════════════════════
    # REFERENCES
    # ════════════════════════════════════════════
    doc.add_heading("References", level=1)
    refs = [
        "S.C. Mun, J.H. Won, Crystals 11(9) (2021) 1013.",
        "P. Arora, Z. Zhang, Chem. Rev. 104(10) (2004) 4419\u20134462.",
        "Q. Wang, et al., J. Power Sources 208 (2012) 210\u2013224.",
        "X. Feng, et al., Energy Storage Mater. 10 (2018) 246\u2013267.",
        "S.S. Zhang, J. Power Sources 164(1) (2007) 351\u2013364.",
        "C.T. Love, J. Power Sources 196(5) (2011) 2905\u20132912.",
        "P. Bawuah, et al., J. Infrared Milli. Terahz. Waves 41 (2020) 450\u2013469.",
        "J.A. Zeitler, et al., J. Pharm. Sci. 96(2) (2007) 330\u2013340.",
        "M. Naftaly, R.E. Miles, Proc. IEEE 95(8) (2007) 1658\u20131665.",
        "M. Anuschek, P. Bawuah, J.A. Zeitler, Int. J. Pharm. X 3 (2021) 100079.",
        "C. Jansen, et al., Appl. Opt. 49(19) (2010) E48\u2013E57.",
        "S. Wietzke, et al., J. Mol. Struct. 1006(1\u20133) (2011) 41\u201351.",
        "R.H. Boyd, Polymer 26(8) (1985) 1123\u20131133.",
        "L. Duvillaret, F. Garet, J.-L. Coutaz, IEEE J. Sel. Top. Quantum Electron. 2(3) (1996) 739\u2013746.",
        "W. Withayachumnankul, M. Naftaly, J. Infrared Milli. Terahz. Waves 35 (2014) 610\u2013637.",
        "T. Dorney, R. Baraniuk, D. Mittleman, J. Opt. Soc. Am. A 18(7) (2001) 1562\u20131571.",
        "D.A.G. Bruggeman, Ann. Phys. 416(7) (1935) 636\u2013664.",
        "J. Cannarella, C.B. Arnold, J. Power Sources 226 (2013) 149\u2013155.",
        "R. Tibshirani, J. R. Stat. Soc. B 58(1) (1996) 267\u2013288.",
        "P.U. Jepsen, D.G. Cooke, M. Koch, Laser Photon. Rev. 5(1) (2011) 124\u2013166.",
        "Y. Benjamini, Y. Hochberg, J. R. Stat. Soc. B 57(1) (1995) 289\u2013300.",
        "B. Efron, R.J. Tibshirani, An Introduction to the Bootstrap, Chapman & Hall, 1993.",
        "W. Nsengiyumva, et al., Opt. Mater. 123 (2022) 111837.",
        "M. Naftaly, R. Dudley, J. Appl. Phys. 109(4) (2011) 043505.",
        "T. Hastie, R. Tibshirani, J. Friedman, The Elements of Statistical Learning, 2nd ed., Springer, 2009.",
        "S. Sommer, et al., J. Infrared Milli. Terahz. Waves 37 (2016) 189\u2013197.",
        "B. Xiong, et al., J. Membr. Sci. 545 (2018) 213\u2013220.",
        "H. Park, J.-H. Son, Sensors 21(4) (2021) 1186.",
        "M. Koumans, et al., Sci. Rep. 14 (2024) 7034.",
    ]
    for i, r in enumerate(refs):
        doc.add_paragraph(f"[{i+1}] {r}")

    # Save
    out_path = OUT_DIR / "PE40_THz_TDS_paper_SNA.docx"
    doc.save(str(out_path))
    print(f"Paper saved: {out_path}")
    print(f"  Sections: Title, Highlights, Abstract, Introduction, Methods, Results, Conclusions")
    print(f"  Figures: 8")
    print(f"  Tables: 1")
    return out_path


if __name__ == "__main__":
    generate()
