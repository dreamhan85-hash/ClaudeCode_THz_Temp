"""Generate Word report for PE40 (260406_Temp) THz-TDS temperature study.

PE20 분리막 2장 겹침 (40 um), Matched Reference 분석 + 샘플별 온도 특성.
"""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FIG_DIR = Path(__file__).parent.parent / "figures" / "report_260406"
CSV_DIR = Path(__file__).parent.parent / "results" / "260406_Temp"
OUT = Path(__file__).parent.parent / "PE40_THz_TDS_분석보고서.docx"


def remove_table_borders(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.makeelement(qn("w:tcBorders"), {})
            for edge in ("top", "left", "bottom", "right"):
                el = tcBorders.makeelement(
                    qn(f"w:{edge}"), {qn("w:val"): "none", qn("w:sz"): "0"})
                tcBorders.append(el)
            tcPr.append(tcBorders)


def add_figure_pair(doc, fig_left, fig_right, cap_left, cap_right):
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = Cm(7)
    for ci, fp in enumerate([fig_left, fig_right]):
        cell = tbl.rows[0].cells[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(fp), width=Cm(6.5))
    for ci, cap in enumerate([cap_left, cap_right]):
        cell = tbl.rows[1].cells[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cap)
        run.font.size = Pt(8)
        run.font.italic = True
    remove_table_borders(tbl)


def add_single_figure(doc, fig_path, caption, width_cm=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(fig_path), width=Cm(width_cm))
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_p.add_run(caption)
    run.font.size = Pt(8)
    run.font.italic = True


def add_data_table(doc, headers, data_rows, caption=None):
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_p.add_run(caption)
        run.font.size = Pt(8)
        run.bold = True
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(7.5)
                r.font.bold = True
    for ri, row_data in enumerate(data_rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(7.5)


def main():
    df = pd.read_csv(CSV_DIR / "pe40_summary_stats.csv")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # ── Title ──
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PE40 THz-TDS 온도 의존성 분석 보고서")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "PE20 이차전지 분리막 2장 겹침(40 μm) — 온도별 Matched Reference 분석")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("측정일: 2024-12-25  |  분석일: 2026-04-06")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 1. 실험 개요
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("1. 실험 개요", level=1)
    doc.add_paragraph(
        "THz-TDS를 이용하여 이차전지용 다공성 PE 분리막의 온도 의존성 광학물성을 "
        "측정하였다. 각 온도에서 해당 온도의 레퍼런스(공기)와 시료를 비교하는 "
        "Matched Reference 방식을 적용하여, 공기 경로의 온도 효과를 자동으로 "
        "상쇄하고 시료 고유의 온도 의존성만을 추출하였다.")

    doc.add_heading("1.1 시료 및 측정", level=2)
    items = [
        ("시료", "PE20 다공성 PE 분리막 × 2장 (총 두께 40 μm)"),
        ("장비", "Menlo Systems ScanControl, Lytera THz-TDS"),
        ("온도 범위", "20–110 °C (10 °C 간격, 총 10개 온도)"),
        ("반복 측정", "각 온도당 5회 반복 (S1–S5, 총 50개 시료 측정)"),
        ("레퍼런스", "각 온도별 별도 측정 (Ref_20.txt ~ Ref_110.txt)"),
        ("분석 방법", "H(ω) = E_sam(T) / E_ref(T) — Matched Reference"),
    ]
    for label, desc in items:
        p = doc.add_paragraph(style="List Bullet")
        run_b = p.add_run(f"{label}: ")
        run_b.bold = True
        p.add_run(desc)

    # ══════════════════════════════════════════════════════════════
    # 2. 시간영역 분석
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("2. 시간영역 분석", level=1)

    doc.add_heading("2.1 피크 영역 확대", level=2)
    add_single_figure(
        doc, FIG_DIR / "fig01_time_domain_zoom.png",
        "Fig. 1. 온도별 THz 파형 (피크 확대). "
        "검정: Ref(T), 색상: 5회 반복 시료.", width_cm=14)
    doc.add_paragraph(
        "각 온도에서 해당 온도의 레퍼런스와 5개 시료 파형을 비교하였다. "
        "40 μm 시료에 의한 시간 지연과 진폭 감쇠가 관찰된다.")

    doc.add_heading("2.2 신호 감쇠 및 지연 상세", level=2)
    add_single_figure(
        doc, FIG_DIR / "fig02_peak_detail.png",
        "Fig. 2. 전 온도 오버레이. 좌: 양의 피크, 우: 음의 피크. "
        "검정: 20°C Ref. 온도 증가 → 진폭 증가(투과율↑) + 시간 좌측 이동(지연 감소).",
        width_cm=12)

    doc.add_heading("2.3 샘플별 시간 지연 vs 온도", level=2)
    add_single_figure(
        doc, FIG_DIR / "fig04_per_sample_dt_vs_temp.png",
        "Fig. 3. 샘플별(S1–S5) 시간 지연 vs 온도. "
        "좌: Δt+(양의 피크), 우: Δt−(음의 피크). 검정: 평균±σ.",
        width_cm=12)
    doc.add_paragraph(
        "5개 샘플 모두 온도 증가에 따라 시간 지연이 감소하는 일관된 추세를 보인다. "
        "Δt_avg: 20°C(~45 fs) → 100°C(~37 fs), R² = 0.77, p < 0.001. "
        "S5가 일관되게 낮은 지연을 보여 시료 간 약간의 변동이 존재한다.")

    doc.add_heading("2.4 샘플별 진폭 비율 vs 온도", level=2)
    add_single_figure(
        doc, FIG_DIR / "fig05_per_sample_amplitude_vs_temp.png",
        "Fig. 4. 샘플별 P2P ratio, Amp+ ratio vs 온도.",
        width_cm=12)
    doc.add_paragraph(
        "P2P 비율은 0.980–0.995 범위로 온도 의존성이 약하다 (R² = 0.10, p = 0.37). "
        "이는 투과 손실 변화가 ~1.5% 이내의 미세한 수준임을 나타낸다.")

    # Table 1
    td_headers = ["온도(°C)", "Δt+(fs)", "σ", "Δt−(fs)", "σ", "P2P ratio"]
    td_rows = []
    for _, row in df.iterrows():
        td_rows.append([
            f"{int(row['Temperature (C)'])}",
            f"{row['dt_pos']:+.1f}", f"{row['dt_pos_std']:.1f}",
            f"{row['dt_neg']:+.1f}", f"{row['dt_neg_std']:.1f}",
            f"{row['p2p_ratio']:.4f}",
        ])
    add_data_table(doc, td_headers, td_rows,
                   "Table 1. 온도별 시간 지연 및 P2P 비율 (5회 평균 ± σ)")
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 3. 광학 상수
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("3. 광학 상수 (Matched Reference)", level=1)

    doc.add_heading("3.1 굴절률 및 흡수계수 스펙트럼", level=2)
    add_single_figure(
        doc, FIG_DIR / "fig_mr_n_alpha.png",
        "Fig. 5. n(f) 및 α(f). 색상: 20°C(파랑) → 110°C(빨강). "
        "H = Sam(T)/Ref(T).", width_cm=12)
    doc.add_paragraph(
        "1.0 THz 기준 n ≈ 1.28–1.34 범위를 보인다. "
        "20–30°C에서 가장 높고 40°C 이후 감소한 뒤 50–90°C에서 안정화된다.")

    doc.add_heading("3.2 샘플별 n vs 온도", level=2)
    add_single_figure(
        doc, FIG_DIR / "fig_Matched_Ref_per_sample_n_vs_temp.png",
        "Fig. 6. 샘플별(S1–S5) n vs 온도 (0.5, 1.0, 1.5 THz).",
        width_cm=14)
    doc.add_paragraph(
        "개별 샘플의 온도 추세를 비교하면, S3·S4가 높은 n, S5가 낮은 n을 보인다. "
        "샘플 간 ±0.03의 변동이 존재하나, 전체적인 감소 추세는 일관적이다. "
        "개별 샘플 상관: S5(R² = 0.66**), S3(R² = 0.64**), S4(R² = 0.46*), "
        "S2(R² = 0.31), S1(R² = 0.17).")

    # Table 2: Per-sample n
    ns_headers = ["온도(°C)", "S1", "S2", "S3", "S4", "S5", "Mean", "σ"]
    ns_rows = []
    for _, row in df.iterrows():
        t = int(row["Temperature (C)"])
        n_mean = row.get("n@1.0THz", 0)
        n_std = row.get("n_std@1.0THz", 0)
        s_vals = []
        for si in range(1, 6):
            col = f"dt_pos_S{si}"
            s_vals.append("")  # placeholder
        ns_rows.append([
            f"{t}", "", "", "", "", "",
            f"{n_mean:.4f}", f"{n_std:.4f}",
        ])

    # Better: read per-sample from optical properties CSV
    df_opt = pd.read_csv(CSV_DIR / "pe40_optical_properties.csv")
    ns_rows = []
    for t in sorted(df["Temperature (C)"].unique()):
        t = int(t)
        vals = []
        for rep in range(1, 6):
            sub = df_opt[(df_opt["temperature_c"] == t) & (df_opt["replicate"] == rep)]
            if len(sub) > 0:
                idx = (sub["freq_thz"] - 1.0).abs().idxmin()
                vals.append(sub.loc[idx, "n"])
            else:
                vals.append(float("nan"))
        import numpy as np
        m, s = np.nanmean(vals), np.nanstd(vals)
        ns_rows.append([f"{t}"] + [f"{v:.4f}" for v in vals] + [f"{m:.4f}", f"{s:.4f}"])

    add_data_table(doc, ns_headers, ns_rows,
                   "Table 2. 샘플별 n @ 1.0 THz (Matched Reference)")
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 4. 유전율
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("4. 유전율", level=1)
    add_single_figure(
        doc, FIG_DIR / "fig_mr_dielectric.png",
        "Fig. 7. 유전율 스펙트럼. 좌: ε' = n²−κ², 우: ε'' = 2nκ.",
        width_cm=12)
    doc.add_paragraph(
        "ε'은 1.0 THz에서 1.64–1.80 범위이며 굴절률과 동일한 온도 의존성을 보인다.")

    # ══════════════════════════════════════════════════════════════
    # 5. 온도 상관 분석
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("5. 온도 상관 분석 (시료 특성)", level=1)

    doc.add_paragraph(
        "시료 고유 특성(비율, 지연, 차이 신호 등)만을 대상으로 온도 상관 분석을 "
        "수행하였다. 레퍼런스 자체의 절대 진폭 변화는 시스템 특성이므로 제외하였다.")

    doc.add_heading("5.1 상관 히트맵", level=2)
    add_single_figure(
        doc, FIG_DIR / "fig_correlation_heatmap.png",
        "Fig. 8. 시료 특성의 온도별 정규화 히트맵.")

    doc.add_heading("5.2 주요 특성 온도 추세", level=2)
    add_figure_pair(
        doc,
        FIG_DIR / "fig_feat_dt_avg.png",
        FIG_DIR / "fig_feat_rise_time_ps.png",
        "Fig. 9(a). Δt_avg vs 온도 (R²=0.773***)",
        "Fig. 9(b). Rise time vs 온도 (R²=0.951***)",
    )
    add_figure_pair(
        doc,
        FIG_DIR / "fig_feat_group_delay_ps.png",
        FIG_DIR / "fig_feat_delta_rms.png",
        "Fig. 9(c). Group delay vs 온도 (R²=0.622**)",
        "Fig. 9(d). Delta RMS vs 온도 (R²=0.445*)",
    )

    doc.add_heading("5.3 온도 상관 순위", level=2)
    corr_headers = ["순위", "파라미터", "도메인", "R²", "p-value"]
    corr_rows = [
        ["1", "rise_time_ps", "시간 영역", "0.951", "1.6×10⁻⁶ ***"],
        ["2", "centroid_sam", "주파수 영역", "0.929", "7.3×10⁻⁶ ***"],
        ["3", "env_asymmetry", "엔벨로프", "0.776", "7.6×10⁻⁴ ***"],
        ["4", "dt_avg", "시간 영역", "0.773", "8.0×10⁻⁴ ***"],
        ["5", "dt_pos", "시간 영역", "0.772", "8.2×10⁻⁴ ***"],
        ["6", "dt_neg", "시간 영역", "0.709", "2.3×10⁻³ **"],
        ["7", "group_delay_ps", "주파수 영역", "0.622", "6.7×10⁻³ **"],
        ["8", "delta_centroid", "차이 신호", "0.553", "1.4×10⁻² *"],
        ["9", "delta_area", "차이 신호", "0.501", "2.2×10⁻² *"],
        ["10", "dt_env", "엔벨로프", "0.482", "2.6×10⁻² *"],
    ]
    add_data_table(doc, corr_headers, corr_rows,
                   "Table 3. 시료 특성 온도 상관 순위 — Top 10 (p < 0.05)")
    doc.add_paragraph()

    doc.add_paragraph(
        "Rise time(R² = 0.95)과 스펙트럼 중심(R² = 0.93)이 가장 높은 상관성을 보이며, "
        "시간 지연(dt_avg: R² = 0.77)이 그 다음이다. "
        "P2P ratio, 진폭 비율 등 진폭 기반 특성은 유의한 온도 상관이 없다 (p > 0.05).")

    # ══════════════════════════════════════════════════════════════
    # 6. 고찰
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6. 고찰", level=1)

    doc.add_heading("6.1 시간 영역 vs 주파수 영역", level=2)
    doc.add_paragraph(
        "시간 영역 직접 분석(rise time, dt_avg)이 광학 상수(n) 추출보다 "
        "더 높은 온도 상관성을 보인다. 이는 40 μm 박막에서 Nelder-Mead 최적화의 "
        "불안정성을 회피하면서도 온도 변화를 민감하게 검출할 수 있음을 시사한다.")

    doc.add_heading("6.2 샘플 간 변동", level=2)
    doc.add_paragraph(
        "5개 샘플 중 S5가 일관되게 낮은 n과 시간 지연을 보이며, "
        "S3가 가장 높은 값을 나타낸다. 이는 2장 겹침 시 층간 공극 상태의 "
        "불균일성에 기인한 것으로 판단되며, 개별 샘플의 온도 상관성도 "
        "S5(R² = 0.66) > S3(R² = 0.64) > S4(R² = 0.46) 순으로 차이를 보인다.")

    doc.add_heading("6.3 PE20 vs PE40", level=2)
    doc.add_paragraph(
        "PE40(n ≈ 1.34)이 PE20(n ≈ 1.28)보다 약 5% 높다. "
        "2장 겹침에 의한 층간 밀착으로 유효 공극률이 감소한 것으로 해석된다.")

    # ══════════════════════════════════════════════════════════════
    # 7. 결론
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("7. 결론", level=1)

    conclusions = [
        "PE40(40 μm)에 대해 온도별 Matched Reference 분석을 수행하여 "
        "각 온도의 Ref(T)와 Sample(T)만을 비교, 시료 고유 온도 특성을 추출하였다.",
        "n @ 1.0 THz: 1.34 (20°C) → 1.28 (100°C). "
        "기울기 −5.3×10⁻⁴/°C.",
        "시료 특성 중 rise time(R² = 0.95***), 스펙트럼 중심(R² = 0.93***), "
        "시간 지연(R² = 0.77***)이 가장 유의한 온도 지표이다.",
        "진폭 기반 특성(P2P ratio, amp ratio)은 온도 상관이 "
        "유의하지 않다 (p > 0.05).",
        "5개 샘플 간 n 변동폭 ±0.03이 존재하나, "
        "온도 감소 추세는 모든 샘플에서 일관적이다.",
        "시간 영역 직접 분석이 광학 상수 추출보다 높은 온도 민감도를 보여, "
        "40 μm급 박막의 온도 모니터링에 더 적합하다.",
    ]
    for c in conclusions:
        doc.add_paragraph(c, style="List Bullet")

    # ══════════════════════════════════════════════════════════════
    # 8. 분석 조건
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("8. 분석 조건", level=1)
    params = [
        ("시료 두께", "40 μm (PE20 × 2장)"),
        ("분석 방법", "Matched Reference — H(ω) = E_sam(T)/E_ref(T)"),
        ("주파수 범위", "0.2–2.5 THz"),
        ("윈도우", "Hann + Zero-padding ×2"),
        ("모델", "Thin-film approximation"),
        ("최적화", "Nelder-Mead (freq-by-freq, warm-starting)"),
        ("피크 검출", "Cubic spline 10x oversampling"),
        ("특성 도메인", "시간 피크 / 엔벨로프 / 차이 신호 / 주파수 / 스펙트럼"),
    ]
    for label, val in params:
        p = doc.add_paragraph(style="List Bullet")
        run_b = p.add_run(f"{label}: ")
        run_b.bold = True
        p.add_run(val)

    doc.save(str(OUT))
    print(f"Report saved: {OUT}")


if __name__ == "__main__":
    main()
